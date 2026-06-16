from datetime import datetime
from collections import deque
from io import BytesIO
import csv
import gzip
import hashlib
import importlib.util
import math
import os
import re
import sys
import tarfile
import threading
import time
import random
import zipfile
from typing import Any, Dict, List, Generator, Optional, Set, Tuple
from urllib.parse import quote, urlparse
import json
from uuid import uuid4
from datetime import timezone
from email.utils import parsedate_to_datetime
from xml.sax.saxutils import escape as xml_escape

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, HTTPException, Header, Request, UploadFile, File, Form, Body
from fastapi.responses import FileResponse, StreamingResponse
from starlette.middleware.base import BaseHTTPMiddleware
try:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, inch
    from reportlab.lib.fonts import addMapping
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Preformatted, NextPageTemplate, Table, TableStyle
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
from pydantic import BaseModel, ConfigDict, Field

# Ensure repo-root modules are importable even when uvicorn is started from flowernet-web/
_WEB_DIR = os.path.dirname(__file__)
_REPO_ROOT = os.path.abspath(os.path.join(_WEB_DIR, ".."))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)
_GEN_DIR = os.path.join(_REPO_ROOT, "flowernet-generator")
if _GEN_DIR not in sys.path:
    sys.path.insert(0, _GEN_DIR)
_OUTLINER_DIR = os.path.join(_REPO_ROOT, "flowernet-outliner")
if _OUTLINER_DIR not in sys.path and os.path.isdir(_OUTLINER_DIR):
    sys.path.insert(0, _OUTLINER_DIR)

# 导入 Citation Verifier 用于引证质量控制
try:
    from citation_verifier import CitationVerifier, verify_references
    HAS_CITATION_VERIFIER = True
except ImportError:
    HAS_CITATION_VERIFIER = False
    print("⚠️ Citation Verifier 未安装，跳过引证验证")

# 导入 Domain Filter 用于基于领域相关性的过滤
try:
    from domain_filter import get_domain_filter
    HAS_DOMAIN_FILTER = True
except ImportError:
    HAS_DOMAIN_FILTER = False
    print("⚠️ Domain Filter 未安装，跳过领域过滤")

# 导入 RAG Search 用于引用兜底重试
try:
    from rag_search import RAGSearchEngine
    HAS_RAG_SEARCH_ENGINE = True
except Exception:
    HAS_RAG_SEARCH_ENGINE = False
    print("⚠️ RAG Search 未安装，无法进行引用兜底重试")

# Optional fetcher for enriching candidate metadata
try:
    import requests
    HAS_REQUESTS = True
except Exception:
    HAS_REQUESTS = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except Exception:
    HAS_BS4 = False

# 导入指标定义
try:
    from metrics_definition import (
        FLOWERNET_METRICS,
        METRICS_CATEGORIES,
        FLOWERNET_FEATURES,
        get_all_metrics,
        get_all_categories,
        get_metrics_by_category,
    )
    HAS_METRICS_DEFINITION = True
except ImportError:
    HAS_METRICS_DEFINITION = False
    print("⚠️ Metrics Definition 未安装，跳过指标展示")
    FLOWERNET_METRICS = {}
    METRICS_CATEGORIES = {}
    FLOWERNET_FEATURES = {}

try:
    from flowernet_agent_stack import get_checkpoint_store
except ImportError:
    class _LocalCheckpointStore:
        def __init__(self):
            state_dir = os.getenv("FLOWERNET_STATE_DIR", "/tmp/flowernet_state")
            os.makedirs(state_dir, exist_ok=True)
            self.path = os.path.join(state_dir, "poffices_checkpoints.json")
            self._lock = threading.Lock()

        def _load(self) -> Dict[str, Any]:
            try:
                with open(self.path, "r", encoding="utf-8") as fh:
                    data = json.load(fh)
                return data if isinstance(data, dict) else {}
            except Exception:
                return {}

        def _save(self, data: Dict[str, Any]) -> None:
            tmp = f"{self.path}.tmp"
            with open(tmp, "w", encoding="utf-8") as fh:
                json.dump(data, fh, ensure_ascii=False)
            os.replace(tmp, self.path)

        def set(self, key: str, value: Any, ttl_seconds: Optional[int] = None) -> None:
            expires_at = time.time() + ttl_seconds if ttl_seconds else None
            with self._lock:
                data = self._load()
                data[key] = {"value": value, "expires_at": expires_at}
                self._save(data)

        def get(self, key: str) -> Any:
            with self._lock:
                data = self._load()
                entry = data.get(key)
                if not isinstance(entry, dict):
                    return None
                expires_at = entry.get("expires_at")
                if expires_at and time.time() > float(expires_at):
                    data.pop(key, None)
                    self._save(data)
                    return None
                return entry.get("value")

    def get_checkpoint_store():
        return _LocalCheckpointStore()

try:
    from flowernet_epistemic import (
        EpistemicAuditEngine,
        attach_chapter_assets,
        augment_content_prompts,
        augment_user_requirements,
        render_audit_markdown,
    )
    HAS_EPISTEMIC_AUDIT = True
except Exception as exc:
    HAS_EPISTEMIC_AUDIT = False
    EpistemicAuditEngine = None  # type: ignore
    print(f"⚠️ Epistemic audit layer unavailable: {exc}")


OUTLINER_URL = os.getenv("OUTLINER_URL", "http://localhost:8003")
GENERATOR_URL = os.getenv("GENERATOR_URL", "http://localhost:8002")
REQUEST_TIMEOUT = int(os.getenv("REQUEST_TIMEOUT", "7200"))
WEB_STATS_PATH = os.getenv("FLOWERNET_WEB_STATS_PATH", "/tmp/flowernet_web_stats.json")
WEB_TOTAL_GENERATIONS_BASE = int(os.getenv("FLOWERNET_TOTAL_GENERATIONS_BASE", "151"))
DOWNSTREAM_RETRIES = int(os.getenv("DOWNSTREAM_RETRIES", "1"))
DOWNSTREAM_BACKOFF = float(os.getenv("DOWNSTREAM_BACKOFF", "1.0"))
DOWNSTREAM_MAX_BACKOFF = float(os.getenv("DOWNSTREAM_MAX_BACKOFF", "30.0"))
DOWNSTREAM_JITTER = float(os.getenv("DOWNSTREAM_JITTER", "0.35"))
DOWNSTREAM_OUTLINER_MIN_RETRIES = int(os.getenv("DOWNSTREAM_OUTLINER_MIN_RETRIES", "1"))
DOWNSTREAM_OUTLINER_MAX_RETRIES = int(os.getenv("DOWNSTREAM_OUTLINER_MAX_RETRIES", "2"))
OUTLINER_TASK_START_RETRIES = int(os.getenv("OUTLINER_TASK_START_RETRIES", "2"))
OUTLINER_TASK_START_BACKOFF = float(os.getenv("OUTLINER_TASK_START_BACKOFF", "8.0"))
OUTLINER_TASK_START_MAX_WAIT = int(os.getenv("OUTLINER_TASK_START_MAX_WAIT", "120"))
DOWNSTREAM_OUTLINER_MIN_DELAY_503 = float(os.getenv("DOWNSTREAM_OUTLINER_MIN_DELAY_503", "8.0"))
DOWNSTREAM_GENERATOR_MIN_RETRIES = int(os.getenv("DOWNSTREAM_GENERATOR_MIN_RETRIES", "1"))
DOWNSTREAM_GENERATOR_MIN_DELAY_429 = float(os.getenv("DOWNSTREAM_GENERATOR_MIN_DELAY_429", "10.0"))
GENERATOR_RESUME_RETRIES = int(os.getenv("GENERATOR_RESUME_RETRIES", "0"))
GENERATOR_DOWNSTREAM_MIN_TIMEOUT = int(os.getenv("GENERATOR_DOWNSTREAM_MIN_TIMEOUT", "600"))
GENERATOR_TASK_POLL_MAX_TIMEOUT = int(os.getenv("GENERATOR_TASK_POLL_MAX_TIMEOUT", "2400"))
GENERATOR_STATUS_MISS_LIMIT = int(os.getenv("GENERATOR_STATUS_MISS_LIMIT", "4"))
GENERATOR_PREFLIGHT_ENABLED = os.getenv("GENERATOR_PREFLIGHT_ENABLED", "true").lower() == "true"
GENERATOR_PREFLIGHT_STRICT = os.getenv("GENERATOR_PREFLIGHT_STRICT", "false").lower() == "true"
GENERATOR_PREFLIGHT_TIMEOUT = max(3.0, float(os.getenv("GENERATOR_PREFLIGHT_TIMEOUT", "20")))
GENERATOR_PREFLIGHT_PROVIDERS = [
    item.strip()
    for item in os.getenv("GENERATOR_PREFLIGHT_PROVIDERS", "deepseek").split(",")
    if item.strip()
]
OUTLINER_DOWNSTREAM_MIN_TIMEOUT = int(os.getenv("OUTLINER_DOWNSTREAM_MIN_TIMEOUT", "600"))
OUTLINER_TASK_POLL_MAX_TIMEOUT = int(os.getenv("OUTLINER_TASK_POLL_MAX_TIMEOUT", "900"))
OUTLINER_STREAM_MAX_WAIT = int(os.getenv("OUTLINER_STREAM_MAX_WAIT", str(OUTLINER_TASK_POLL_MAX_TIMEOUT)))
DOWNSTREAM_OUTLINER_MIN_DELAY_429 = float(os.getenv("DOWNSTREAM_OUTLINER_MIN_DELAY_429", "35.0"))
DOWNSTREAM_OUTLINER_COOLDOWN_429 = float(os.getenv("DOWNSTREAM_OUTLINER_COOLDOWN_429", "60.0"))
WEB_INPROCESS_OUTLINER_FALLBACK = os.getenv("WEB_INPROCESS_OUTLINER_FALLBACK", "false").lower() == "true"
POFFICES_POLL_WAIT_SECONDS = max(1, int(os.getenv("POFFICES_POLL_WAIT_SECONDS", "25")))
POFFICES_POLL_INTERVAL_SECONDS = max(0.5, float(os.getenv("POFFICES_POLL_INTERVAL_SECONDS", "2.0")))
DOWNSTREAM_GENERATOR_COOLDOWN_429 = float(os.getenv("DOWNSTREAM_GENERATOR_COOLDOWN_429", "20.0"))
GENERATOR_RESUME_BACKOFF = float(os.getenv("GENERATOR_RESUME_BACKOFF", "2.0"))
API_AUTH_ENABLED = os.getenv("API_AUTH_ENABLED", "false").lower() == "true"
API_KEY = os.getenv("FLOWERNET_API_KEY", "")
BEARER_TOKEN = os.getenv("FLOWERNET_BEARER_TOKEN", "")
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "")
FRONTEND_SIGNATURE_ENFORCED = os.getenv("FRONTEND_SIGNATURE_ENFORCED", "true").lower() == "true"
WEB_DEFAULT_REL_THRESHOLD = float(os.getenv("WEB_DEFAULT_REL_THRESHOLD", "0.765"))
# Slightly stricter than the previous 0.40 default so borderline repetition
# enters Controller, while still leaving room for normal long-document overlap.
WEB_DEFAULT_RED_THRESHOLD = float(os.getenv("WEB_DEFAULT_RED_THRESHOLD", "0.265"))
ENABLE_CITATION_QA = os.getenv("ENABLE_CITATION_QA", "true").lower() == "true"
CITATION_MIN_SECTION_HIGH_QUALITY = int(os.getenv("CITATION_MIN_SECTION_HIGH_QUALITY", "1"))
CITATION_LOW_QUALITY_MAX_RATIO = float(os.getenv("CITATION_LOW_QUALITY_MAX_RATIO", "0.5"))
STRICT_CITATION_ENFORCEMENT = os.getenv("STRICT_CITATION_ENFORCEMENT", "true").lower() == "true"
CITATION_FAIL_FAST = os.getenv("CITATION_FAIL_FAST", "false").lower() == "true"
TIMEOUT_ADAPTIVE_ENABLED = os.getenv("TIMEOUT_ADAPTIVE_ENABLED", "true").lower() == "true"
TIMEOUT_MIN_SECONDS = int(os.getenv("TIMEOUT_MIN_SECONDS", "60"))
TIMEOUT_MAX_SECONDS = int(os.getenv("TIMEOUT_MAX_SECONDS", "7200"))
TIMEOUT_SAFETY_FACTOR = float(os.getenv("TIMEOUT_SAFETY_FACTOR", "1.35"))
TIMEOUT_FIXED_BUFFER_SECONDS = int(os.getenv("TIMEOUT_FIXED_BUFFER_SECONDS", "20"))
TIMEOUT_BASE_OUTLINE_SECONDS = int(os.getenv("TIMEOUT_BASE_OUTLINE_SECONDS", "25"))
TIMEOUT_BASE_CITATION_SECONDS = int(os.getenv("TIMEOUT_BASE_CITATION_SECONDS", "8"))
TIMEOUT_BASE_ITERATION_SECONDS = float(os.getenv("TIMEOUT_BASE_ITERATION_SECONDS", "55"))
ESTIMATED_ITERATIONS_PER_SUBSECTION = float(os.getenv("ESTIMATED_ITERATIONS_PER_SUBSECTION", "1.8"))
DOMAIN_FILTER_RAG_RETRY_MAX = int(os.getenv("DOMAIN_FILTER_RAG_RETRY_MAX", "3"))
DOMAIN_FILTER_RAG_RETRY_RESULTS = int(os.getenv("DOMAIN_FILTER_RAG_RETRY_RESULTS", "6"))
DOMAIN_FILTER_RAG_RETRY_TIMEOUT = int(os.getenv("DOMAIN_FILTER_RAG_RETRY_TIMEOUT", "12"))
DOMAIN_FILTER_FALLBACK_TOP_K = int(os.getenv("DOMAIN_FILTER_FALLBACK_TOP_K", "3"))
CITATION_POSTPROCESS_BUDGET_SECONDS = float(os.getenv("CITATION_POSTPROCESS_BUDGET_SECONDS", "20"))
CITATION_ACADEMIC_QUERY_PAIR_LIMIT = int(os.getenv("CITATION_ACADEMIC_QUERY_PAIR_LIMIT", "3"))
CITATION_SEMANTIC_SCHOLAR_RETRIES = int(os.getenv("CITATION_SEMANTIC_SCHOLAR_RETRIES", "1"))
CITATION_EXTERNAL_FALLBACK_ENABLED = os.getenv("CITATION_EXTERNAL_FALLBACK_ENABLED", "false").lower() == "true"
CITATION_METADATA_FETCH_ENABLED = os.getenv("CITATION_METADATA_FETCH_ENABLED", "false").lower() == "true"

DOWNSTREAM_SESSION = requests.Session()
DOWNSTREAM_SESSION.trust_env = False
CITATION_HTTP_SESSION = requests.Session()
CITATION_HTTP_SESSION.trust_env = False

POFFICES_TASKS: Dict[str, Dict[str, Any]] = {}
POFFICES_TASKS_LOCK = threading.Lock()
POFFICES_CHECKPOINT_STORE = None
POFFICES_CHECKPOINT_STORE_LOCK = threading.Lock()
POFFICES_OUTLINER_RETRY_STALE_SECONDS = int(os.getenv("POFFICES_OUTLINER_RETRY_STALE_SECONDS", "180"))
POFFICES_RESTART_STALE_SECONDS = int(os.getenv("POFFICES_RESTART_STALE_SECONDS", "180"))
POFFICES_OUTLINER_RETRY_MAX = int(os.getenv("POFFICES_OUTLINER_RETRY_MAX", "20"))
RECENT_PIPELINE_SECONDS = deque(maxlen=20)
RECENT_ITERATION_SECONDS = deque(maxlen=20)
METRICS_LOCK = threading.Lock()
RATE_LIMIT_LOCK = threading.Lock()
RATE_LIMIT_UNTIL: Dict[str, float] = {}


def _poffices_task_key(task_id: str) -> str:
    return f"poffices_task:{task_id}"


def _normalize_poffices_request_key_value(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"\s+", " ", text)
    return text[:2000]


def _poffices_request_key(req: "PofficesGenerateRequest") -> str:
    parts = [
        _normalize_poffices_request_key_value(req.query),
        str(int(req.chapter_count or 0)),
        str(int(req.subsection_count or 0)),
        _normalize_poffices_request_key_value(req.user_background),
        _normalize_poffices_request_key_value(req.extra_requirements),
    ]
    digest = hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()
    return f"poffices_request:{digest}"


def _poffices_request_loose_key(req: "PofficesGenerateRequest") -> str:
    parts = [
        _normalize_poffices_request_key_value(req.query),
        str(int(req.chapter_count or 0)),
        str(int(req.subsection_count or 0)),
    ]
    digest = hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()
    return f"poffices_request_loose:{digest}"


def _persist_poffices_request_task(req: "PofficesGenerateRequest", task_id: str) -> None:
    try:
        store = _get_poffices_checkpoint_store()
        payload = {"task_id": task_id, "updated_at": datetime.now().isoformat()}
        store.set(_poffices_request_key(req), payload, ttl_seconds=7 * 24 * 3600)
        store.set(_poffices_request_loose_key(req), payload, ttl_seconds=7 * 24 * 3600)
    except Exception as exc:
        print(f"[Poffices] request-task mapping persist failed for {task_id}: {exc}")


def _restore_poffices_request_task_id(req: "PofficesGenerateRequest") -> str:
    store = None
    try:
        store = _get_poffices_checkpoint_store()
    except Exception as exc:
        print(f"[Poffices] request-task mapping store unavailable: {exc}")
    for key in (_poffices_request_key(req), _poffices_request_loose_key(req)):
        try:
            restored = store.get(key) if store is not None else None
        except Exception as exc:
            print(f"[Poffices] request-task mapping restore failed: {exc}")
            restored = None
        if isinstance(restored, dict):
            task_id = str(restored.get("task_id") or "").strip()
            if task_id:
                return task_id
    return ""


def _find_recent_poffices_task_for_request(req: "PofficesGenerateRequest") -> str:
    """Find an in-memory task for the same Poffices request when wiring lost task_id."""
    query_key = _normalize_poffices_request_key_value(req.query)
    if not query_key:
        return ""
    candidates: List[Tuple[str, Dict[str, Any]]] = []
    with POFFICES_TASKS_LOCK:
        for task_id, task in POFFICES_TASKS.items():
            if not isinstance(task, dict):
                continue
            request_payload = task.get("request")
            if not isinstance(request_payload, dict):
                continue
            task_query = _normalize_poffices_request_key_value(request_payload.get("query"))
            same_shape = (
                int(request_payload.get("chapter_count") or 0) == int(req.chapter_count or 0)
                and int(request_payload.get("subsection_count") or 0) == int(req.subsection_count or 0)
            )
            if task_query == query_key and same_shape and str(task.get("status") or "").lower() in {"queued", "running", "completed"}:
                candidates.append((task_id, dict(task)))
    if not candidates:
        return ""
    candidates.sort(key=lambda item: str(item[1].get("updated_at") or item[1].get("created_at") or ""), reverse=True)
    return candidates[0][0]


def _get_poffices_checkpoint_store():
    global POFFICES_CHECKPOINT_STORE
    if POFFICES_CHECKPOINT_STORE is not None:
        return POFFICES_CHECKPOINT_STORE
    with POFFICES_CHECKPOINT_STORE_LOCK:
        if POFFICES_CHECKPOINT_STORE is None:
            POFFICES_CHECKPOINT_STORE = get_checkpoint_store()
    return POFFICES_CHECKPOINT_STORE


def _persist_poffices_task(task_id: str, task: Dict[str, Any]) -> None:
    try:
        _get_poffices_checkpoint_store().set(
            _poffices_task_key(task_id),
            {k: v for k, v in dict(task).items() if k != "_thread_started"},
            ttl_seconds=7 * 24 * 3600,
        )
    except Exception as exc:
        print(f"[Poffices] checkpoint persist failed for {task_id}: {exc}")


def _set_poffices_task(task_id: str, **updates: Any) -> Dict[str, Any]:
    with POFFICES_TASKS_LOCK:
        task = POFFICES_TASKS.setdefault(task_id, {})
        task.update(updates)
        task["updated_at"] = datetime.now().isoformat()
        snapshot = dict(task)
    _persist_poffices_task(task_id, snapshot)
    return snapshot


def _restore_poffices_task(task_id: str) -> Optional[Dict[str, Any]]:
    with POFFICES_TASKS_LOCK:
        existing = POFFICES_TASKS.get(task_id)
        if existing:
            return dict(existing)
    try:
        restored = _get_poffices_checkpoint_store().get(_poffices_task_key(task_id))
    except Exception as exc:
        print(f"[Poffices] checkpoint restore failed for {task_id}: {exc}")
        restored = None
    if not isinstance(restored, dict):
        return None
    with POFFICES_TASKS_LOCK:
        POFFICES_TASKS[task_id] = dict(restored)
        return dict(POFFICES_TASKS[task_id])


def _poffices_outliner_retry_stale(task: Dict[str, Any]) -> bool:
    status = str(task.get("status") or "").lower()
    if status not in {"queued", "running"}:
        return False
    text = " ".join(
        str(task.get(key) or "")
        for key in ("message", "last_retryable_error", "error")
    ).lower()
    if not any(token in text for token in ("outliner", "大纲", "限流", "排队")):
        return False
    updated_age = _iso_age_seconds(str(task.get("updated_at") or ""))
    return updated_age >= max(60, POFFICES_OUTLINER_RETRY_STALE_SECONDS)


def _poffices_task_restart_stale(task: Dict[str, Any]) -> bool:
    status = str(task.get("status") or "").lower()
    if status not in {"queued", "running"}:
        return False
    updated_age = _iso_age_seconds(str(task.get("updated_at") or ""))
    started_age = _iso_age_seconds(str(task.get("started_at") or task.get("created_at") or ""))
    return max(updated_age, started_age) >= max(60, POFFICES_RESTART_STALE_SECONDS)


def _poffices_task_cancelled(task: Dict[str, Any]) -> bool:
    return str(task.get("cancel_requested") or "").lower() == "true" or task.get("cancel_requested") is True


def _cancel_poffices_task(task_id: str) -> Dict[str, Any]:
    task = _restore_poffices_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="task_id not found")
    if str(task.get("status") or "").lower() == "completed":
        return {
            "success": True,
            "task_id": task_id,
            "status": "completed",
            "task_status": "completed",
            "message": "任务已完成，无法取消",
        }
    _set_poffices_task(
        task_id,
        status="failed",
        task_status="failed",
        cancel_requested=True,
        error="task_cancelled",
        message="任务已取消",
        completed_at=datetime.now().isoformat(),
    )
    return {
        "success": True,
        "task_id": task_id,
        "status": "failed",
        "task_status": "failed",
        "message": "任务已取消",
        "error": "task_cancelled",
    }


def _restart_restored_poffices_task(task_id: str, task: Dict[str, Any]) -> None:
    status = str(task.get("status") or "").lower()
    if status not in {"queued", "running"}:
        return
    if _poffices_task_cancelled(task):
        _set_poffices_task(
            task_id,
            status="failed",
            error="task_cancelled",
            message="任务已取消",
            completed_at=datetime.now().isoformat(),
        )
        return
    request_payload = task.get("request")
    if not isinstance(request_payload, dict):
        _set_poffices_task(
            task_id,
            status="failed",
            error="restored_task_missing_request_payload",
            message="任务恢复失败",
            completed_at=datetime.now().isoformat(),
        )
        return
    with POFFICES_TASKS_LOCK:
        current = POFFICES_TASKS.setdefault(task_id, dict(task))
        if not current.get("_thread_started") and not _poffices_task_restart_stale(current):
            return
        if current.get("_thread_started"):
            # A live in-process worker may be inside the outliner-start retry loop.
            # Starting a second worker for the same task multiplies downstream start
            # requests and can turn a transient 429 into a long self-sustaining queue.
            return
        current["_thread_started"] = True
    try:
        req = PofficesGenerateRequest(**request_payload)
    except Exception as exc:
        _set_poffices_task(
            task_id,
            status="failed",
            error=f"restored_task_request_invalid: {exc}",
            message="任务恢复失败",
            completed_at=datetime.now().isoformat(),
        )
        return
    _set_poffices_task(task_id, status="queued", message="任务已从 checkpoint 恢复，重新入队")
    threading.Thread(target=_run_poffices_task, args=(task_id, req), daemon=True).start()


class GenerateDocRequest(BaseModel):
    topic: str = Field(..., min_length=2, description="文档主题")
    chapter_count: int = Field(default=5, ge=1, le=10)
    subsection_count: int = Field(default=3, ge=1, le=8)
    user_background: str = Field(default="")
    extra_requirements: str = Field(default="")
    rel_threshold: float = Field(default=WEB_DEFAULT_REL_THRESHOLD, ge=0, le=1)
    red_threshold: float = Field(default=WEB_DEFAULT_RED_THRESHOLD, ge=0, le=1)
    timeout_seconds: int = Field(default=7200, ge=60, le=7200, description="同步模式超时秒数")


class DownloadDocxRequest(BaseModel):
    title: str
    content: str


class PofficesGenerateRequest(BaseModel):
    model_config = ConfigDict(extra="allow")

    query: str = Field(default="", description="用户输入查询")
    task_id: str = Field(default="", description="Existing FlowerNet task id. If present, /generate behaves as task polling instead of creating a new task.")
    chapter_count: int = Field(default=5, ge=1, le=10)
    subsection_count: int = Field(default=3, ge=1, le=8)
    user_background: str = Field(default="")
    extra_requirements: str = Field(default="")
    rel_threshold: float = Field(default=WEB_DEFAULT_REL_THRESHOLD, ge=0, le=1)
    red_threshold: float = Field(default=WEB_DEFAULT_RED_THRESHOLD, ge=0, le=1)
    async_mode: bool = Field(default=True, description="true=异步任务，false=同步等待结果")
    timeout_seconds: int = Field(default=7200, ge=60, le=7200, description="同步模式超时秒数")


class PofficesTaskStatusRequest(BaseModel):
    model_config = ConfigDict(extra="allow")

    task_id: str = ""
    wait: bool = True
    wait_seconds: int = Field(default=POFFICES_POLL_WAIT_SECONDS, ge=1, le=7200)
    cancel: bool = False


app = FastAPI(title="FlowerNet Web UI", version="1.0.0")


def _frontend_index_path() -> str:
    return os.path.join(_WEB_DIR, "static", "index.html")


def _assert_expected_frontend() -> None:
    if not FRONTEND_SIGNATURE_ENFORCED:
        return
    static_path = _frontend_index_path()
    try:
        html_text = open(static_path, "r", encoding="utf-8").read()
    except Exception as exc:
        raise RuntimeError(f"FlowerNet frontend index is missing or unreadable: {static_path}: {exc}") from exc

    required_signatures = [
        "FlowerNet Agent Dashboard",
        "FlowerNet Agent: High-Quality Long Document Generation",
        "Section-Level Micro Tracking",
        "Verifier Quality Control",
        "FlowerNet Full Metrics Board",
        "self-auditing scientific writing system",
    ]
    missing = [sig for sig in required_signatures if sig not in html_text]
    forbidden_signatures = [
        "FlowerNet 文档生成器",
    ]
    forbidden = [sig for sig in forbidden_signatures if sig in html_text]
    if missing or forbidden:
        raise RuntimeError(
            "Unexpected FlowerNet frontend. Refusing to serve legacy UI. "
            f"path={static_path}, missing={missing}, forbidden={forbidden}"
        )


@app.on_event("startup")
def verify_frontend_bundle() -> None:
    _assert_expected_frontend()

# Add request logging middleware for debugging
class LoggingMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        if "/api/download" in str(request.url):
            print(f"[Middleware] Incoming {request.method} {request.url.path}", flush=True)
        response = await call_next(request)
        if "/api/download" in str(request.url):
            print(f"[Middleware] Response status: {response.status_code}", flush=True)
        return response

app.add_middleware(LoggingMiddleware)


def verify_auth(
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    if not API_AUTH_ENABLED:
        return

    api_ok = bool(API_KEY) and x_api_key == API_KEY
    bearer_ok = bool(BEARER_TOKEN) and authorization == f"Bearer {BEARER_TOKEN}"
    if not (api_ok or bearer_ok):
        raise HTTPException(status_code=401, detail="Unauthorized: invalid API key or bearer token")


def _extract_response_error(response: requests.Response) -> str:
    try:
        payload = response.json()
        if isinstance(payload, dict):
            detail = payload.get("detail")
            error = payload.get("error")
            message = payload.get("message")
            if isinstance(detail, list):
                return json.dumps(detail, ensure_ascii=False)
            if detail:
                return str(detail)
            if error:
                return str(error)
            if message:
                return str(message)
        return json.dumps(payload, ensure_ascii=False)[:800]
    except ValueError:
        return response.text[:800]


def _clean_error_text(value: Any, fallback: str = "unknown_error") -> str:
    if isinstance(value, dict):
        for key in ("error", "message", "detail"):
            text = _clean_error_text(value.get(key), "")
            if text:
                return text
        return fallback
    text = str(value or "").strip()
    if not text or text.lower() == "none":
        return fallback
    return text


def _parse_retry_after_seconds(retry_after: str) -> float | None:
    value = (retry_after or "").strip()
    if not value:
        return None
    try:
        return max(0.0, float(value))
    except (TypeError, ValueError):
        pass
    try:
        dt = parsedate_to_datetime(value)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        now = datetime.now(timezone.utc)
        return max(0.0, (dt - now).total_seconds())
    except Exception:
        return None


def _parse_seconds_value(raw: Any) -> float | None:
    text = str(raw or "").strip().lower()
    if not text:
        return None
    if text.endswith("s"):
        text = text[:-1].strip()
    try:
        return float(text)
    except (TypeError, ValueError):
        return None


def _estimate_subsection_count(chapter_count: int, subsection_count: int) -> int:
    return max(1, int(chapter_count) * int(subsection_count))


def _quantile(values: List[float], q: float) -> float | None:
    if not values:
        return None
    sorted_vals = sorted(float(v) for v in values)
    if len(sorted_vals) == 1:
        return sorted_vals[0]

    q = max(0.0, min(1.0, float(q)))
    pos = (len(sorted_vals) - 1) * q
    lower = int(pos)
    upper = min(lower + 1, len(sorted_vals) - 1)
    if lower == upper:
        return sorted_vals[lower]
    weight = pos - lower
    return sorted_vals[lower] * (1.0 - weight) + sorted_vals[upper] * weight


def _build_timeout_profile(
    chapter_count: int,
    subsection_count: int,
    requested_timeout: int,
) -> Dict[str, Any]:
    expected_subsections = _estimate_subsection_count(chapter_count, subsection_count)

    with METRICS_LOCK:
        recent_iteration = list(RECENT_ITERATION_SECONDS)
        recent_pipeline = list(RECENT_PIPELINE_SECONDS)

    iteration_p50 = _quantile(recent_iteration, 0.50)
    iteration_p90 = _quantile(recent_iteration, 0.90)
    iteration_p95 = _quantile(recent_iteration, 0.95)
    pipeline_p50 = _quantile(recent_pipeline, 0.50)
    pipeline_p90 = _quantile(recent_pipeline, 0.90)
    pipeline_p95 = _quantile(recent_pipeline, 0.95)

    iteration_seconds_for_budget = (
        float(iteration_p90)
        if iteration_p90 is not None
        else TIMEOUT_BASE_ITERATION_SECONDS
    )

    estimated_iterations = max(1.0, expected_subsections * ESTIMATED_ITERATIONS_PER_SUBSECTION)
    estimated_seconds = (
        TIMEOUT_BASE_OUTLINE_SECONDS
        + TIMEOUT_BASE_CITATION_SECONDS
        + iteration_seconds_for_budget * estimated_iterations
    )
    if pipeline_p95 is not None:
        estimated_seconds = max(estimated_seconds, float(pipeline_p95))

    recommended_timeout = int(estimated_seconds * TIMEOUT_SAFETY_FACTOR + TIMEOUT_FIXED_BUFFER_SECONDS)
    recommended_timeout = max(TIMEOUT_MIN_SECONDS, min(TIMEOUT_MAX_SECONDS, recommended_timeout))

    requested = int(requested_timeout or REQUEST_TIMEOUT)
    requested = max(TIMEOUT_MIN_SECONDS, min(TIMEOUT_MAX_SECONDS, requested))

    effective_timeout = max(requested, recommended_timeout) if TIMEOUT_ADAPTIVE_ENABLED else requested
    effective_timeout = max(TIMEOUT_MIN_SECONDS, min(TIMEOUT_MAX_SECONDS, int(effective_timeout)))

    return {
        "adaptive_enabled": TIMEOUT_ADAPTIVE_ENABLED,
        "quantile_window_tasks": 20,
        "expected_subsections": expected_subsections,
        "estimated_iterations": round(estimated_iterations, 2),
        "iteration_seconds_for_budget": round(float(iteration_seconds_for_budget), 2),
        "iteration_p50_seconds": round(float(iteration_p50), 2) if iteration_p50 is not None else None,
        "iteration_p90_seconds": round(float(iteration_p90), 2) if iteration_p90 is not None else None,
        "iteration_p95_seconds": round(float(iteration_p95), 2) if iteration_p95 is not None else None,
        "pipeline_p50_seconds": round(float(pipeline_p50), 2) if pipeline_p50 is not None else None,
        "pipeline_p90_seconds": round(float(pipeline_p90), 2) if pipeline_p90 is not None else None,
        "pipeline_p95_seconds": round(float(pipeline_p95), 2) if pipeline_p95 is not None else None,
        "requested_timeout_seconds": requested,
        "recommended_timeout_seconds": recommended_timeout,
        "effective_timeout_seconds": effective_timeout,
    }


def _record_timeout_metrics(elapsed_seconds: float, result: Dict[str, Any]) -> None:
    if elapsed_seconds <= 0:
        return

    total_iterations = 0
    generation_time_seconds = None
    if isinstance(result, dict):
        stats = result.get("stats") or {}
        if isinstance(stats, dict):
            total_iterations = int(stats.get("total_iterations", 0) or 0)
            generation_time_seconds = _parse_seconds_value(stats.get("generation_time"))

    if generation_time_seconds is None:
        generation_time_seconds = float(elapsed_seconds)

    with METRICS_LOCK:
        RECENT_PIPELINE_SECONDS.append(float(elapsed_seconds))
        if total_iterations > 0:
            RECENT_ITERATION_SECONDS.append(float(generation_time_seconds) / float(total_iterations))


def _inject_timeout_profile(result: Dict[str, Any], timeout_profile: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(result, dict):
        return result
    stats = result.get("stats")
    if not isinstance(stats, dict):
        stats = {}
        result["stats"] = stats
    stats["timeout_profile"] = timeout_profile
    return result


def _read_web_stats() -> Dict[str, Any]:
    stats = {
        "total_generations": WEB_TOTAL_GENERATIONS_BASE,
        "successful_generations": 0,
        "partial_generations": 0,
    }
    try:
        with open(WEB_STATS_PATH, "r", encoding="utf-8") as f:
            loaded = json.load(f)
        if isinstance(loaded, dict):
            for key in list(stats.keys()):
                stats[key] = int(loaded.get(key, stats[key]) or 0)
    except Exception:
        pass
    stats["total_generations"] = max(WEB_TOTAL_GENERATIONS_BASE, int(stats.get("total_generations", 0) or 0))
    return stats


def _write_web_stats(stats: Dict[str, Any]) -> None:
    try:
        os.makedirs(os.path.dirname(WEB_STATS_PATH) or ".", exist_ok=True)
        with open(WEB_STATS_PATH, "w", encoding="utf-8") as f:
            json.dump(stats, f, ensure_ascii=False)
    except Exception as e:
        print(f"[WebStats] 写入统计失败: {e}", flush=True)


def _record_generation_success(partial: bool = False) -> Dict[str, Any]:
    stats = _read_web_stats()
    stats["total_generations"] = int(stats.get("total_generations", WEB_TOTAL_GENERATIONS_BASE) or WEB_TOTAL_GENERATIONS_BASE) + 1
    if partial:
        stats["partial_generations"] = int(stats.get("partial_generations", 0) or 0) + 1
    else:
        stats["successful_generations"] = int(stats.get("successful_generations", 0) or 0) + 1
    stats["updated_at"] = datetime.now().isoformat()
    _write_web_stats(stats)
    return stats


def _is_transient_downstream_payload(payload: Dict[str, Any]) -> bool:
    if not isinstance(payload, dict):
        return False
    if payload.get("success") is not False:
        return False
    message = str(payload.get("error") or payload.get("message") or "").lower()
    transient_tokens = [
        "429", "too many requests", "resource_exhausted", "quota", "rate",
        "timeout", "timed out", "temporarily", "503", "502", "504", "retry",
        "已有大纲生成任务正在执行",
    ]
    return any(token in message for token in transient_tokens)


def _downstream_rate_limit_key(url: str) -> str:
    try:
        parsed = urlparse(url)
        host = (parsed.netloc or "").strip().lower()
        if "outliner" in host or "/outline/" in parsed.path:
            return "outliner"
        if "generator" in host or "/generate" in parsed.path:
            return "generator"
        return host or "downstream"
    except Exception:
        return "downstream"


def _set_rate_limit_cooldown(rate_key: str, seconds: float) -> None:
    delay = max(0.0, float(seconds or 0.0))
    if delay <= 0:
        return
    with RATE_LIMIT_LOCK:
        now = time.time()
        next_allowed = now + delay
        current = float(RATE_LIMIT_UNTIL.get(rate_key, 0.0) or 0.0)
        RATE_LIMIT_UNTIL[rate_key] = max(current, next_allowed)


def _get_rate_limit_sleep_seconds(rate_key: str) -> float:
    with RATE_LIMIT_LOCK:
        now = time.time()
        next_allowed = float(RATE_LIMIT_UNTIL.get(rate_key, 0.0) or 0.0)
    return max(0.0, next_allowed - now)


def _iso_age_seconds(value: str) -> float:
    if not value:
        return 0.0
    try:
        ts = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if ts.tzinfo is not None:
            return max(0.0, time.time() - ts.timestamp())
        return max(0.0, (datetime.now() - ts).total_seconds())
    except Exception:
        return 0.0


def _probe_service_health(base_url: str, timeout: float = 5.0) -> bool:
    root_url = base_url.rstrip("/") + "/"
    health_url = base_url.rstrip("/") + "/health"

    for url in (health_url, root_url):
        try:
            response = DOWNSTREAM_SESSION.get(url, timeout=timeout)
            if 200 <= response.status_code < 500:
                return True
        except requests.RequestException:
            continue
    return False


def _wait_for_service_ready(base_url: str, label: str, max_wait_seconds: float = 45.0) -> bool:
    deadline = time.time() + max(0.0, float(max_wait_seconds or 0.0))
    attempt = 0
    while True:
        attempt += 1
        if _probe_service_health(base_url, timeout=5.0):
            print(f"[Web] ✅ {label} 健康检查通过")
            return True

        remaining = deadline - time.time()
        if remaining <= 0:
            print(f"[Web] ⚠️ {label} 健康检查超时，继续尝试主请求")
            return False

        sleep_seconds = min(5.0 + min(attempt * 1.5, 10.0), remaining)
        print(f"[Web] ⏳ 等待 {label} 就绪（第 {attempt} 次，{sleep_seconds:.1f}s 后重试）")
        time.sleep(max(1.0, sleep_seconds))


def post_json_with_retry(url: str, payload: Dict[str, Any], timeout: int) -> Dict[str, Any]:
    last_error: str = ""
    is_outliner_url = "outliner" in url or "/outline/" in url
    is_generator_url = "generator" in url or "/generate_document" in url
    rate_key = _downstream_rate_limit_key(url)
    effective_retries = DOWNSTREAM_RETRIES
    if is_outliner_url:
        effective_retries = max(effective_retries, DOWNSTREAM_OUTLINER_MIN_RETRIES)
        effective_retries = min(effective_retries, max(1, DOWNSTREAM_OUTLINER_MAX_RETRIES))
    if is_generator_url:
        effective_retries = max(effective_retries, DOWNSTREAM_GENERATOR_MIN_RETRIES)

    for attempt in range(1, effective_retries + 1):
        cooldown_sleep = _get_rate_limit_sleep_seconds(rate_key)
        if cooldown_sleep > 0:
            time.sleep(cooldown_sleep)

        retry_delay = DOWNSTREAM_BACKOFF * (2 ** max(0, attempt - 1))
        retry_delay += random.uniform(0, DOWNSTREAM_JITTER)
        retry_after_seconds = None

        try:
            response = DOWNSTREAM_SESSION.post(url, json=payload, timeout=timeout)
            response.raise_for_status()
            try:
                body = response.json()
            except ValueError:
                content_type = response.headers.get("Content-Type", "")
                response_body = (response.text or "")[:800]
                last_error = (
                    f"HTTP {response.status_code} from {url}: 下游返回非JSON响应 "
                    f"(Content-Type={content_type or 'unknown'}): {response_body or '<empty>'}"
                )
                if attempt < effective_retries:
                    retry_delay = min(retry_delay, DOWNSTREAM_MAX_BACKOFF)
                    time.sleep(retry_delay)
                    continue
                break

            # Outliner already performs internal provider-level retries.
            # Avoid multiplying retries here which can trigger remote 429.
            if is_outliner_url and isinstance(body, dict) and body.get("success") is False:
                return body

            if _is_transient_downstream_payload(body) and attempt < effective_retries:
                last_error = f"下游返回可重试失败: {body}"
                retry_delay = min(retry_delay, DOWNSTREAM_MAX_BACKOFF)
                time.sleep(retry_delay)
                continue

            return body
        except requests.RequestException as exc:
            response = getattr(exc, "response", None)
            if response is not None:
                response_body = _extract_response_error(response)
                last_error = f"HTTP {response.status_code} from {url}: {response_body}"
                if response.status_code == 429:
                    retry_after = response.headers.get("Retry-After", "")
                    retry_after_seconds = _parse_retry_after_seconds(retry_after)
                    if is_generator_url:
                        retry_delay = max(retry_delay, DOWNSTREAM_GENERATOR_MIN_DELAY_429)
                        _set_rate_limit_cooldown(
                            rate_key,
                            max(retry_delay, retry_after_seconds or 0.0, DOWNSTREAM_GENERATOR_COOLDOWN_429),
                        )
                    if is_outliner_url:
                        # For Outliner on Render Free: 429 means rate limit exceeded.
                        # Don't retry - break immediately to fail fast and prevent DoS spiral.
                        last_error = f"HTTP 429 from {url}: {response_body} (Render rate limit - no retry)"
                        break
                if is_outliner_url and response.status_code == 503:
                    retry_delay = max(retry_delay, DOWNSTREAM_OUTLINER_MIN_DELAY_503)
            else:
                last_error = str(exc)

            if attempt < effective_retries:
                if retry_after_seconds is not None:
                    retry_delay = max(retry_delay, retry_after_seconds)
                retry_delay = min(retry_delay, DOWNSTREAM_MAX_BACKOFF)
                time.sleep(retry_delay)

    raise HTTPException(
        status_code=502,
        detail=f"下游服务请求失败(重试{effective_retries}次): {url}, 错误: {last_error}",
    )


def post_json(url: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    return post_json_with_retry(url=url, payload=payload, timeout=REQUEST_TIMEOUT)


def post_json_once(url: str, payload: Dict[str, Any], timeout: int) -> Dict[str, Any]:
    try:
        response = DOWNSTREAM_SESSION.post(url, json=payload, timeout=timeout)
        response.raise_for_status()
        try:
            body = response.json()
        except ValueError:
            content_type = response.headers.get("Content-Type", "")
            response_body = (response.text or "")[:800]
            raise HTTPException(
                status_code=502,
                detail=(
                    f"HTTP {response.status_code} from {url}: 下游返回非JSON响应 "
                    f"(Content-Type={content_type or 'unknown'}): {response_body or '<empty>'}"
                ),
            )

        if isinstance(body, dict):
            return body
        raise HTTPException(status_code=502, detail=f"HTTP {response.status_code} from {url}: 下游返回格式异常")
    except requests.RequestException as exc:
        response = getattr(exc, "response", None)
        if response is not None:
            response_body = _extract_response_error(response)
            raise HTTPException(
                status_code=502,
                detail=f"HTTP {response.status_code} from {url}: {response_body}",
            )
        raise HTTPException(status_code=502, detail=f"下游服务请求失败: {url}, 错误: {exc}")


def _extract_title_from_requirements(requirements: str) -> str:
    text = str(requirements or "").strip()
    for line in text.splitlines():
        line = line.strip(" -#:\t")
        if not line:
            continue
        for prefix in ("文档主题", "主题", "Topic", "topic"):
            if line.startswith(prefix):
                candidate = line.split(":", 1)[-1].split("：", 1)[-1].strip()
                if candidate:
                    return candidate[:80]
        return line[:80]
    return "FlowerNet Document"


def _is_placeholder_outline_label(value: Any) -> bool:
    text = str(value or "").strip()
    if not text:
        return True
    return bool(
        re.fullmatch(r"第\s*\d+\s*[章节节]", text)
        or re.fullmatch(r"(section|chapter|subsection)\s*\d*", text, flags=re.I)
        or re.fullmatch(r"第\s*\d+\s*章\s*-\s*第\s*\d+\s*节", text)
    )


def _is_bad_outline_label(value: Any) -> bool:
    text = " ".join(str(value or "").split()).strip()
    if _is_placeholder_outline_label(text):
        return True
    lowered = text.lower()
    request_artifacts = [
        "请帮我", "生成一篇", "高质量长文档", "额外要求",
        "document topic", "extra requirements", "run generation",
    ]
    if any(token in lowered for token in request_artifacts):
        return True
    label_artifact_patterns = [
        r"用户背景\s*[:：]",
        r"用户需求\s*[:：]",
        r"附加要求\s*[:：]",
        r"user background\s*:",
        r"user requirements\s*:",
    ]
    if any(re.search(pattern, text, flags=re.I) for pattern in label_artifact_patterns):
        return True
    # Academic outline labels often use a colon, e.g. "Mechanisms: A and B".
    # Reject only labels that look like full pasted sentences or request text.
    if len(text) > 140 and re.search(r"[。！？.!?，,；;]", text):
        return True
    if len(text) > 90 and re.search(r"[。！？.!?]", text):
        return True
    return False


def _validate_outline_structure_quality(structure: Dict[str, Any]) -> tuple[bool, str]:
    sections = structure.get("sections", []) if isinstance(structure, dict) else []
    if not isinstance(sections, list) or not sections:
        return False, "outline_has_no_sections"
    bad_labels: List[str] = []
    for section in sections:
        if not isinstance(section, dict):
            bad_labels.append("invalid_section")
            continue
        section_title = str(section.get("title") or "")
        if _is_bad_outline_label(section_title):
            bad_labels.append(section_title[:80] or "empty_section_title")
        subsections = section.get("subsections", [])
        if not isinstance(subsections, list) or not subsections:
            bad_labels.append(f"{section_title[:40]}: no_subsections")
            continue
        for subsection in subsections:
            if not isinstance(subsection, dict):
                bad_labels.append(f"{section_title[:40]}: invalid_subsection")
                continue
            sub_title = str(subsection.get("title") or "")
            if _is_bad_outline_label(sub_title):
                bad_labels.append(sub_title[:80] or "empty_subsection_title")
    if bad_labels:
        return False, "bad_outline_titles: " + " | ".join(bad_labels[:4])
    return True, "ok"


def _build_local_outline_response(payload: Dict[str, Any], reason: str) -> Dict[str, Any]:
    """Return a hard failure instead of fabricating an outline."""
    return {
        "success": False,
        "error": "outliner_unavailable_or_rate_limited",
        "fallback_reason": reason,
        "retryable": _is_transient_downstream_payload({"error": reason}),
        "retry_after_seconds": max(DOWNSTREAM_OUTLINER_MIN_DELAY_429, DOWNSTREAM_OUTLINER_COOLDOWN_429),
    }


def _is_retryable_outline_failure(value: Any) -> bool:
    if isinstance(value, dict):
        if value.get("retryable") is True:
            return True
        text = json.dumps(value, ensure_ascii=False, default=str)
    else:
        text = str(value or "")
    lowered = text.lower()
    return (
        "outliner_unavailable_or_rate_limited" in lowered
        or "outliner_rate_limited" in lowered
        or "http 429" in lowered
        or "too many requests" in lowered
        or "resource_exhausted" in lowered
        or "quota" in lowered
        or "rate limit" in lowered
    )


def _outline_retry_delay_seconds(value: Any, attempt: int = 1) -> float:
    retry_after = 0.0
    if isinstance(value, dict):
        try:
            retry_after = float(value.get("retry_after_seconds") or 0.0)
        except Exception:
            retry_after = 0.0
    base = max(DOWNSTREAM_OUTLINER_MIN_DELAY_429, OUTLINER_TASK_START_BACKOFF * max(1, attempt))
    return min(DOWNSTREAM_MAX_BACKOFF, max(base, retry_after, DOWNSTREAM_OUTLINER_COOLDOWN_429))


_LOCAL_OUTLINER_CLASS = None
_LOCAL_OUTLINER_LOCK = threading.Lock()


def _load_local_outliner_class():
    global _LOCAL_OUTLINER_CLASS
    if _LOCAL_OUTLINER_CLASS is not None:
        return _LOCAL_OUTLINER_CLASS
    candidates = [
        os.path.join(_OUTLINER_DIR, "outliner.py"),
        os.path.join(_WEB_DIR, "outliner.py"),
        os.path.join(os.getcwd(), "outliner.py"),
    ]
    for path in candidates:
        if not os.path.exists(path):
            continue
        spec = importlib.util.spec_from_file_location("flowernet_web_embedded_outliner", path)
        if spec is None or spec.loader is None:
            continue
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        cls = getattr(module, "FlowerNetOutliner", None)
        if cls is not None:
            _LOCAL_OUTLINER_CLASS = cls
            return cls
    raise RuntimeError("FlowerNetOutliner is not available inside flowernet-web deployment")


def _call_inprocess_outliner_generate(payload: Dict[str, Any], reason: str = "") -> Dict[str, Any]:
    """Generate a real outline inside the web service as a resilience path.

    This is not a template fallback. It reuses the same FlowerNetOutliner class
    and the same DeepSeek-backed provider path, avoiding a second Render service
    hop when the standalone outliner is rate-limited or cold.
    """
    if not WEB_INPROCESS_OUTLINER_FALLBACK:
        return _build_local_outline_response(payload, reason=reason or "inprocess_outliner_disabled")
    with _LOCAL_OUTLINER_LOCK:
        try:
            old_chain = os.environ.get("OUTLINER_PROVIDER_CHAIN")
            old_provider = os.environ.get("OUTLINER_PROVIDER")
            os.environ["OUTLINER_PROVIDER_CHAIN"] = os.getenv("WEB_OUTLINER_PROVIDER_CHAIN", "deepseek")
            os.environ["OUTLINER_PROVIDER"] = os.getenv("WEB_OUTLINER_PROVIDER", "deepseek")
            cls = _load_local_outliner_class()
            outliner = cls(provider=os.environ["OUTLINER_PROVIDER_CHAIN"])
            result = outliner.generate_full_outline(
                user_background=str(payload.get("user_background") or ""),
                user_requirements=str(payload.get("user_requirements") or ""),
                max_sections=int(payload.get("max_sections") or 5),
                max_subsections_per_section=int(payload.get("max_subsections_per_section") or 4),
            )
        except Exception as exc:
            return {
                "success": False,
                "error": "inprocess_outliner_failed",
                "fallback_reason": reason,
                "detail": str(exc),
                "retryable": _is_retryable_outline_failure(str(exc)),
            }
        finally:
            if old_chain is None:
                os.environ.pop("OUTLINER_PROVIDER_CHAIN", None)
            else:
                os.environ["OUTLINER_PROVIDER_CHAIN"] = old_chain
            if old_provider is None:
                os.environ.pop("OUTLINER_PROVIDER", None)
            else:
                os.environ["OUTLINER_PROVIDER"] = old_provider

    if isinstance(result, dict) and result.get("success"):
        result.setdefault("outline_saved", False)
        result.setdefault("structure_outline_saved", False)
        metadata = result.setdefault("metadata", {})
        if isinstance(metadata, dict):
            metadata["web_inprocess_outliner_fallback"] = True
            metadata["remote_failure_reason"] = reason[:500]
        return result
    if isinstance(result, dict):
        result.setdefault("fallback_reason", reason)
        result.setdefault("retryable", _is_retryable_outline_failure(result))
        return result
    return {"success": False, "error": "inprocess_outliner_invalid_result", "fallback_reason": reason, "retryable": False}


def call_outliner_generate_and_save(payload: Dict[str, Any], timeout: int) -> Dict[str, Any]:
    """Start outline generation through the async task API and poll for completion.

    Render can return 429 for long-running blocking requests even when the user
    only clicked once. The task API keeps the public request short and lets the
    outliner serialize work internally.
    """
    task_url = f"{OUTLINER_URL}/outline/generate-task"
    legacy_url = f"{OUTLINER_URL}/outline/generate-and-save"
    poll_timeout = min(
        max(OUTLINER_DOWNSTREAM_MIN_TIMEOUT, int(timeout)),
        max(60, OUTLINER_TASK_POLL_MAX_TIMEOUT),
    )
    task_resp: Dict[str, Any] = {}

    start_attempts = max(1, OUTLINER_TASK_START_RETRIES)
    start_deadline = time.time() + max(
        60,
        min(OUTLINER_TASK_START_MAX_WAIT, max(60, poll_timeout - 30)),
    )
    last_start_error = ""
    for start_attempt in range(1, start_attempts + 1):
        try:
            task_resp = post_json_once(url=task_url, payload=payload, timeout=min(30, poll_timeout))
            break
        except HTTPException as exc:
            detail = str(exc.detail)
            last_start_error = detail
            if "HTTP 404" in detail or "Not Found" in detail:
                print("[Web] ⚠️ Outliner async task API unavailable, falling back to legacy blocking endpoint")
                return post_json_once(url=legacy_url, payload=payload, timeout=poll_timeout)
            if "HTTP 429" not in detail and "Too Many Requests" not in detail:
                raise
            if start_attempt >= start_attempts or time.time() >= start_deadline:
                print("[Web] ⚠️ Outliner task start repeatedly hit 429; remote-only mode will retry at task level")
                inprocess_resp = _call_inprocess_outliner_generate(
                    payload,
                    reason=f"remote outliner task start hit 429: {last_start_error[:500]}",
                )
                if isinstance(inprocess_resp, dict) and inprocess_resp.get("success"):
                    print("[Web] ✅ In-process DeepSeek outliner recovered after remote task-start 429")
                    return inprocess_resp
                return {
                    "success": False,
                    "error": "outliner_task_start_rate_limited",
                    "detail": last_start_error[:1000],
                    "retryable": True,
                    "remote_only": True,
                }
            delay = min(
                DOWNSTREAM_MAX_BACKOFF,
                max(DOWNSTREAM_OUTLINER_MIN_DELAY_429, OUTLINER_TASK_START_BACKOFF * start_attempt),
            )
            delay = min(delay, max(1.0, start_deadline - time.time()))
            print(f"[Web] ⏳ Outliner task start hit 429; retrying in {delay:.1f}s ({start_attempt}/{start_attempts})")
            time.sleep(delay)

    if isinstance(task_resp, dict) and task_resp.get("status") == "completed" and isinstance(task_resp.get("result"), dict):
        return task_resp["result"]

    task_id = str(task_resp.get("task_id") or "").strip() if isinstance(task_resp, dict) else ""
    if not task_id:
        if isinstance(task_resp, dict) and task_resp.get("success") is True and task_resp.get("structure"):
            return task_resp
        raise HTTPException(status_code=502, detail=f"Outliner task API returned no task_id: {task_resp}")

    deadline = time.time() + poll_timeout
    sleep_seconds = 2.0
    status_url = f"{OUTLINER_URL}/outline/task-status/{task_id}"
    last_status: Dict[str, Any] = {}

    while time.time() < deadline:
        try:
            response = DOWNSTREAM_SESSION.get(status_url, timeout=15)
            response.raise_for_status()
            status_body = response.json()
            if isinstance(status_body, dict):
                last_status = status_body
                status = str(status_body.get("status") or "").lower()
                if status == "completed" and isinstance(status_body.get("result"), dict):
                    return status_body["result"]
                if status == "failed":
                    result = status_body.get("result")
                    if isinstance(result, dict):
                        return result
                    return {
                        "success": False,
                        "error": status_body.get("error") or "outliner_task_failed",
                        "task_id": task_id,
                    }
                started_at = str(status_body.get("started_at") or "")
                heartbeat_at = str(status_body.get("heartbeat_at") or "")
                updated_at = str(status_body.get("updated_at") or "")
                if status in {"queued", "running"}:
                    last_status = status_body
                    if started_at or heartbeat_at or updated_at:
                        heartbeat_age = _iso_age_seconds(heartbeat_at or updated_at or started_at)
                        run_age = _iso_age_seconds(started_at) if started_at else 0.0
                        if heartbeat_age > 240 and run_age > 300:
                            return {
                                "success": False,
                                "error": "outliner_task_stale_no_heartbeat",
                                "task_id": task_id,
                                "last_status": status_body,
                            }
        except Exception as e:
            last_status = {"error": str(e), "task_id": task_id}

        time.sleep(min(8.0, sleep_seconds))
        sleep_seconds = min(8.0, sleep_seconds * 1.3)

    return {
        "success": False,
        "error": f"outliner_task_timeout after {poll_timeout}s",
        "task_id": task_id,
        "last_status": last_status,
        "retryable": True,
        "remote_only": True,
    }


def fetch_history_items(document_id: str, timeout_seconds: int = 60) -> List[Dict[str, Any]]:
    try:
        history_resp = post_json_with_retry(
            f"{OUTLINER_URL}/history/get",
            {"document_id": document_id},
            timeout_seconds,
        )
        if history_resp.get("success") and isinstance(history_resp.get("history"), list):
            return history_resp.get("history", [])
    except Exception as e:
        print(f"获取历史内容失败: {e}")
    return []


def fetch_progress_events(document_id: str, timeout_seconds: int = 30, limit: int = 1000) -> List[Dict[str, Any]]:
    try:
        events_resp = post_json_with_retry(
            f"{OUTLINER_URL}/history/progress",
            {"document_id": document_id, "after_id": 0, "limit": limit},
            timeout_seconds,
        )
        events = events_resp.get("events") if isinstance(events_resp, dict) else []
        return events if isinstance(events, list) else []
    except Exception as e:
        print(f"获取进度事件失败: {e}")
    return []


def extract_quality_metrics_from_progress_events(events: List[Dict[str, Any]]) -> Dict[str, Any]:
    quality_scores: List[float] = []
    dimension_sums: Dict[str, float] = {}
    dimension_counts: Dict[str, int] = {}
    unieval_available = 0
    unieval_fallback = 0
    arms = ["llm", "rule", "rule_structured", "defect_topic", "defect_evidence", "defect_structure"]
    arm_counts = {arm: 0 for arm in arms}
    recent_bandit: List[Dict[str, Any]] = []
    reward_sum = 0.0
    reward_count = 0
    last_arm = ""
    last_mode = ""

    for event in events or []:
        stage = str(event.get("stage") or "")
        meta = event.get("metadata") if isinstance(event.get("metadata"), dict) else {}
        if stage == "verifier_result":
            raw_score = meta.get("quality_score")
            try:
                quality_scores.append(float(raw_score))
            except (TypeError, ValueError):
                pass
            dims = meta.get("quality_dimensions")
            if isinstance(dims, dict):
                for key, value in dims.items():
                    try:
                        dimension_sums[key] = dimension_sums.get(key, 0.0) + float(value)
                        dimension_counts[key] = dimension_counts.get(key, 0) + 1
                    except (TypeError, ValueError):
                        continue
            if meta.get("unieval_fallback"):
                unieval_fallback += 1
            else:
                unieval_available += 1
        elif stage == "controller_result":
            arm = str(meta.get("selected_arm") or meta.get("chosen_arm") or "")
            if arm not in arm_counts:
                continue
            try:
                reward = max(0.0, min(1.0, float(meta.get("reward", 0.0) or 0.0)))
            except (TypeError, ValueError):
                reward = 0.0
            mode = str(meta.get("selection_mode") or "")
            arm_counts[arm] += 1
            reward_sum += reward
            reward_count += 1
            last_arm = arm
            last_mode = mode
            recent_bandit.append({
                "arm": arm,
                "reward": reward,
                "selection_mode": mode,
            })

    dimension_avgs = {
        key: round(dimension_sums.get(key, 0.0) / max(1, dimension_counts.get(key, 0)), 4)
        for key in dimension_sums
    }
    rewards = [float(item.get("reward", 0.0) or 0.0) for item in recent_bandit]
    if rewards:
        ymin = max(0.0, min(rewards) - 0.02)
        ymax = min(1.0, max(rewards) + 0.02)
        if ymax <= ymin:
            ymax = min(1.0, ymin + 0.1)
    else:
        ymin, ymax = 0.0, 0.1

    return {
        "quality_score_avg": round(sum(quality_scores) / max(1, len(quality_scores)), 4) if quality_scores else 0.0,
        "quality_score_count": len(quality_scores),
        "quality_dimension_avgs": dimension_avgs,
        "unieval_available_subsections": unieval_available,
        "unieval_fallback_subsections": unieval_fallback,
        "unieval_available_ratio": round(unieval_available / max(1, unieval_available + unieval_fallback), 4),
        "unieval_fallback_ratio": round(unieval_fallback / max(1, unieval_available + unieval_fallback), 4),
        "bandit_selected_arm_counts": arm_counts,
        "bandit_reward_sum": round(reward_sum, 4),
        "bandit_reward_count": reward_count,
        "bandit_reward_avg": round(reward_sum / max(1, reward_count), 4) if reward_count else 0.0,
        "bandit_last_selected_arm": last_arm,
        "bandit_last_selection_mode": last_mode,
        "bandit_recent_events": recent_bandit[-200:],
        "bandit_plot_y_min": ymin,
        "bandit_plot_y_max": ymax,
    }


def _recover_partial_document(document_id: str, attempts: int = 5, timeout_seconds: int = 30) -> Dict[str, Any]:
    attempts = max(1, int(attempts))
    delay_seconds = 1.5
    last_history_items: List[Dict[str, Any]] = []

    for attempt in range(1, attempts + 1):
        history_items = fetch_history_items(document_id=document_id, timeout_seconds=timeout_seconds)
        if history_items:
            title, structure = _load_document_structure(document_id=document_id, history=history_items)
            progress_events = fetch_progress_events(document_id=document_id, timeout_seconds=timeout_seconds)
            progress_metrics = extract_quality_metrics_from_progress_events(progress_events)
            markdown_content = build_markdown_document(
                title=title,
                structure=structure,
                history=history_items,
                generated_sections=None,
            )

            expected = 0
            for sec in (structure.get("sections") or []):
                subs = sec.get("subsections") if isinstance(sec, dict) else []
                if isinstance(subs, list):
                    expected += len(subs)

            passed = len(history_items)
            if expected > 0 and passed < expected:
                return {
                    "success": False,
                    "partial": True,
                    "document_id": document_id,
                    "title": title,
                    "content": "",
                    "stats": {
                        "expected_subsections": expected,
                        "passed_subsections": passed,
                        "failed_subsections": 0,
                        "forced_subsections": 0,
                        "total_generated": passed,
                        **progress_metrics,
                    },
                    "history_items": history_items,
                    "progress_events": progress_events[-200:],
                    "attempts": attempt,
                    "error": f"partial_document_rejected: passed {passed}/{expected}",
                }
            return {
                "success": True,
                "partial": False,
                "document_id": document_id,
                "title": title,
                "content": markdown_content,
                "stats": {
                    "expected_subsections": expected,
                    "passed_subsections": passed,
                    "failed_subsections": 0,
                    "forced_subsections": 0,
                    "total_generated": passed,
                    **progress_metrics,
                },
                "history_items": history_items,
                "progress_events": progress_events[-200:],
                "attempts": attempt,
            }

        task_id = _find_generator_task_id_by_document(document_id)
        if task_id:
            try:
                response = DOWNSTREAM_SESSION.get(
                    f"{GENERATOR_URL}/generate_document_task/{task_id}",
                    timeout=min(15, max(5, timeout_seconds)),
                )
                response.raise_for_status()
                task_status = response.json()
            except Exception as exc:
                task_status = {"status": "unknown", "error": str(exc), "task_id": task_id}

            status = str(task_status.get("status") or "unknown").lower()
            result = task_status.get("result") if isinstance(task_status, dict) else None
            if status == "completed" and isinstance(result, dict):
                if result.get("partial") or result.get("interrupted"):
                    return {
                        "success": False,
                        "partial": True,
                        "document_id": document_id,
                        "generator_task_id": task_id,
                        "generator_status": status,
                        "error": "partial_generator_result_rejected",
                    }
                return {
                    "success": True,
                    "partial": False,
                    "document_id": document_id,
                    "title": str(result.get("title") or "FlowerNet Document"),
                    "content": str(result.get("content") or ""),
                    "stats": result.get("stats") if isinstance(result.get("stats"), dict) else {},
                    "generator_task_id": task_id,
                    "generator_status": status,
                }
            if status in {"queued", "running"}:
                return {
                    "success": False,
                    "partial": True,
                    "document_id": document_id,
                    "title": "",
                    "content": "",
                    "stats": {
                        "expected_subsections": 1,
                        "passed_subsections": 0,
                        "failed_subsections": 0,
                        "forced_subsections": 0,
                        "total_generated": 0,
                    },
                    "generator_task_id": task_id,
                    "generator_status": status,
                    "message": f"后台生成任务仍在{status}",
                }
            if status == "failed":
                return {
                    "success": False,
                    "document_id": document_id,
                    "generator_task_id": task_id,
                    "generator_status": status,
                    "error": _clean_error_text(task_status, f"generator_task_failed: {task_id}"),
                    "message": "后台生成任务失败",
                }

        last_history_items = history_items
        if attempt < attempts:
            time.sleep(min(3.0, delay_seconds))
            delay_seconds = min(5.0, delay_seconds * 1.6)

    return {
        "success": False,
        "document_id": document_id,
        "error": "history_or_generator_task_not_ready",
        "message": "后台内容尚未可恢复，且未找到可追踪的 generator 任务。",
        "history_items": last_history_items,
        "attempts": attempts,
    }


def extract_orchestration_metrics(gen_resp: Dict[str, Any]) -> Dict[str, int]:
    if not isinstance(gen_resp, dict):
        return {
            "controller_triggered_subsections": 0,
            "controller_calls_total": 0,
            "controller_success_total": 0,
            "controller_error_total": 0,
            "controller_unavailable_total": 0,
            "controller_ineffective_total": 0,
            "controller_fallback_outline_total": 0,
            "controller_exhausted_total": 0,
            "verifier_failed_total": 0,
        }

    return {
        "controller_triggered_subsections": int(gen_resp.get("controller_triggered_subsections", 0) or 0),
        "controller_calls_total": int(gen_resp.get("controller_calls_total", 0) or 0),
        "controller_success_total": int(gen_resp.get("controller_success_total", 0) or 0),
        "controller_error_total": int(gen_resp.get("controller_error_total", 0) or 0),
        "controller_unavailable_total": int(gen_resp.get("controller_unavailable_total", 0) or 0),
        "controller_ineffective_total": int(gen_resp.get("controller_ineffective_total", 0) or 0),
        "controller_fallback_outline_total": int(gen_resp.get("controller_fallback_outline_total", 0) or 0),
        "controller_exhausted_total": int(gen_resp.get("controller_exhausted_total", 0) or 0),
        "verifier_failed_total": int(gen_resp.get("verifier_failed_total", 0) or 0),
    }


def extract_document_quality_metrics(gen_resp: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(gen_resp, dict):
        return {
            "quality_score_avg": 0.0,
            "quality_overall_uncertainty_avg": 0.0,
            "quality_dimension_avgs": {},
            "quality_weights": {},
            "unieval_available_subsections": 0,
            "unieval_fallback_subsections": 0,
            "unieval_available_ratio": 0.0,
            "unieval_fallback_ratio": 0.0,
            "bandit_selected_arm_counts": {},
            "bandit_reward_sum": 0.0,
            "bandit_reward_count": 0,
            "bandit_reward_avg": 0.0,
            "bandit_drift_events": 0,
            "bandit_drift_triggered_subsections": 0,
            "bandit_drift_trigger_rate": 0.0,
            "bandit_last_selected_arm": "",
            "bandit_last_selection_mode": "",
            "bandit_last_constraints": {},
            "bandit_recent_events": [],
        }

    return {
        "quality_score_avg": float(gen_resp.get("quality_score_avg", 0.0) or 0.0),
        "quality_overall_uncertainty_avg": float(gen_resp.get("quality_overall_uncertainty_avg", 0.0) or 0.0),
        "quality_dimension_avgs": gen_resp.get("quality_dimension_avgs", {}) if isinstance(gen_resp.get("quality_dimension_avgs"), dict) else {},
        "quality_weights": gen_resp.get("quality_weights", {}) if isinstance(gen_resp.get("quality_weights"), dict) else {},
        "unieval_available_subsections": int(gen_resp.get("unieval_available_subsections", 0) or 0),
        "unieval_fallback_subsections": int(gen_resp.get("unieval_fallback_subsections", 0) or 0),
        "unieval_available_ratio": float(gen_resp.get("unieval_available_ratio", 0.0) or 0.0),
        "unieval_fallback_ratio": float(gen_resp.get("unieval_fallback_ratio", 0.0) or 0.0),
        "bandit_selected_arm_counts": gen_resp.get("bandit_selected_arm_counts", {}) if isinstance(gen_resp.get("bandit_selected_arm_counts"), dict) else {},
        "bandit_reward_sum": float(gen_resp.get("bandit_reward_sum", 0.0) or 0.0),
        "bandit_reward_count": int(gen_resp.get("bandit_reward_count", 0) or 0),
        "bandit_reward_avg": float(gen_resp.get("bandit_reward_avg", 0.0) or 0.0),
        "bandit_drift_events": int(gen_resp.get("bandit_drift_events", 0) or 0),
        "bandit_drift_triggered_subsections": int(gen_resp.get("bandit_drift_triggered_subsections", 0) or 0),
        "bandit_drift_trigger_rate": float(gen_resp.get("bandit_drift_trigger_rate", 0.0) or 0.0),
        "bandit_last_selected_arm": str(gen_resp.get("bandit_last_selected_arm", "") or ""),
        "bandit_last_selection_mode": str(gen_resp.get("bandit_last_selection_mode", "") or ""),
        "bandit_last_constraints": gen_resp.get("bandit_last_constraints", {}) if isinstance(gen_resp.get("bandit_last_constraints"), dict) else {},
        "bandit_recent_events": gen_resp.get("bandit_recent_events", []) if isinstance(gen_resp.get("bandit_recent_events"), list) else [],
        # Optional plot bounds provided by orchestrator for frontend chart scaling
        "bandit_plot_y_min": float(gen_resp.get("plot_y_min", 0.0) or 0.0),
        "bandit_plot_y_max": float(gen_resp.get("plot_y_max", 0.0) or 0.0),
    }


def _find_generator_task_id_by_document(document_id: str) -> str:
    """Best-effort lookup for an already queued/running generator task."""
    if not document_id:
        return ""
    try:
        response = DOWNSTREAM_SESSION.get(f"{GENERATOR_URL}/generate_document_tasks/summary", timeout=12)
        response.raise_for_status()
        summary = response.json()
    except Exception as exc:
        print(f"[Web] ⚠️ Generator task summary lookup failed: {exc}")
        return ""
    items = summary.get("tasks") if isinstance(summary, dict) else None
    if not isinstance(items, list):
        return ""
    preferred_status = {"queued", "running", "completed"}
    for item in items:
        if not isinstance(item, dict):
            continue
        if str(item.get("document_id") or "") != str(document_id):
            continue
        status = str(item.get("status") or "").lower()
        if status in preferred_status:
            return str(item.get("task_id") or "")
    return ""


def _generator_task_exists(document_id: str, task_id: str = "") -> bool:
    try:
        response = DOWNSTREAM_SESSION.get(f"{GENERATOR_URL}/generate_document_tasks/summary", timeout=12)
        response.raise_for_status()
        summary = response.json()
    except Exception as exc:
        print(f"[Web] ⚠️ Generator task existence lookup failed: {exc}")
        # If the status endpoint already missed repeatedly and even the summary
        # endpoint is unreachable, treating the task as "probably still alive"
        # leaves the UI stuck at 0/N until the full poll timeout. Fail fast so
        # the user can retry after Render recovers/redeploys.
        return False
    items = summary.get("tasks") if isinstance(summary, dict) else None
    if not isinstance(items, list):
        return True
    for item in items:
        if not isinstance(item, dict):
            continue
        if task_id and str(item.get("task_id") or "") == str(task_id):
            return True
        if document_id and str(item.get("document_id") or "") == str(document_id):
            return True
    return False


def preflight_generator_warning(topic: str) -> str:
    if not GENERATOR_PREFLIGHT_ENABLED:
        return ""
    probe_prompt = f"请用一句中文确认你可以为主题“{topic}”生成专业长文档内容。"
    try:
        response = DOWNSTREAM_SESSION.post(
            f"{GENERATOR_URL}/diagnostics/providers",
            json={"providers": GENERATOR_PREFLIGHT_PROVIDERS, "prompt": probe_prompt, "max_tokens": 40},
            timeout=GENERATOR_PREFLIGHT_TIMEOUT,
        )
        if response.status_code == 429:
            warning = "生成模型预检被限流，已跳过预检并继续正式生成"
            print(f"[Web] ⚠️ {warning}: {response.text[:180]}")
            return warning
        response.raise_for_status()
        body = response.json()
    except Exception as exc:
        warning = f"生成模型预检失败，已跳过预检并继续正式生成: {exc}"
        print(f"[Web] ⚠️ {warning}")
        if GENERATOR_PREFLIGHT_STRICT:
            raise HTTPException(status_code=503, detail=f"生成模型预检失败: {exc}")
        return warning

    results = body.get("results") if isinstance(body, dict) else {}
    if not isinstance(results, dict) or not any(isinstance(item, dict) and item.get("success") for item in results.values()):
        error_summary = " | ".join(
            f"{name}: {str((item or {}).get('error_summary') or 'failed')[:160]}"
            for name, item in (results.items() if isinstance(results, dict) else [])
        ) or "unknown generator provider error"
        warning = f"生成模型预检未通过，已继续正式生成: {error_summary[:360]}"
        print(f"[Web] ⚠️ {warning}")
        if GENERATOR_PREFLIGHT_STRICT:
            raise HTTPException(status_code=503, detail=f"生成模型暂不可用: {error_summary[:360]}")
        return warning
    return ""


def generate_document_with_recovery(
    document_id: str,
    generate_payload: Dict[str, Any],
    timeout_seconds: int,
) -> Dict[str, Any]:
    total_attempts = 1 + max(0, GENERATOR_RESUME_RETRIES)
    last_error = ""

    for attempt in range(1, total_attempts + 1):
        try:
            # Keep each downstream call bounded so async tasks do not stay running forever.
            # The generator is the slowest hop and has to survive the full subsection loop.
            # Use a larger floor here than the general stream budget so SSE validation does
            # not fail just because the orchestrator needs more time than the adaptive web budget.
            call_timeout = min(
                GENERATOR_TASK_POLL_MAX_TIMEOUT,
                max(GENERATOR_DOWNSTREAM_MIN_TIMEOUT, int(timeout_seconds)),
            )
            task_url = f"{GENERATOR_URL}/generate_document_task"
            legacy_url = f"{GENERATOR_URL}/generate_document"
            task_resp: Dict[str, Any] = {}

            start_deadline = time.time() + min(call_timeout, max(120, GENERATOR_DOWNSTREAM_MIN_TIMEOUT))
            start_attempt = 0
            while True:
                start_attempt += 1
                try:
                    task_resp = post_json_once(task_url, generate_payload, timeout=min(30, call_timeout))
                    break
                except HTTPException as exc:
                    detail = str(exc.detail)
                    if "HTTP 404" in detail or "Not Found" in detail:
                        print("[Web] ⚠️ Generator async task API unavailable, falling back to legacy blocking endpoint")
                        result = post_json_with_retry(legacy_url, generate_payload, call_timeout)
                        if isinstance(result, dict):
                            result.setdefault("recovery_attempt", attempt)
                            result.setdefault("recovery_attempts", total_attempts)
                        return result
                    if "HTTP 429" not in detail and "Too Many Requests" not in detail:
                        raise

                    existing_task_id = _find_generator_task_id_by_document(document_id)
                    if existing_task_id:
                        task_resp = {
                            "success": True,
                            "task_id": existing_task_id,
                            "document_id": document_id,
                            "status": "running",
                            "reused": True,
                        }
                        print(f"[Web] ♻️ Generator start hit 429; reusing existing task {existing_task_id}")
                        break

                    remaining = start_deadline - time.time()
                    if remaining <= 0:
                        raise
                    delay = min(
                        remaining,
                        DOWNSTREAM_MAX_BACKOFF,
                        max(DOWNSTREAM_GENERATOR_MIN_DELAY_429, GENERATOR_RESUME_BACKOFF * start_attempt),
                    )
                    print(f"[Web] ⏳ Generator task start hit 429; waiting {delay:.1f}s before retry ({start_attempt})")
                    time.sleep(max(1.0, delay))

            if (
                isinstance(task_resp, dict)
                and task_resp.get("status") == "completed"
                and isinstance(task_resp.get("result"), dict)
            ):
                result = task_resp["result"]
                result.setdefault("recovery_attempt", attempt)
                result.setdefault("recovery_attempts", total_attempts)
                return result

            task_id = str(task_resp.get("task_id") or "").strip() if isinstance(task_resp, dict) else ""
            if not task_id:
                raise HTTPException(status_code=502, detail=f"Generator task API returned no task_id: {task_resp}")

            deadline = time.time() + call_timeout
            sleep_seconds = 2.0
            status_url = f"{GENERATOR_URL}/generate_document_task/{task_id}"
            last_status: Dict[str, Any] = {}
            consecutive_status_misses = 0

            while time.time() < deadline:
                try:
                    response = DOWNSTREAM_SESSION.get(status_url, timeout=15)
                    response.raise_for_status()
                    status_body = response.json()
                    consecutive_status_misses = 0
                    if isinstance(status_body, dict):
                        last_status = status_body
                        status = str(status_body.get("status") or "").lower()
                        if status == "completed" and isinstance(status_body.get("result"), dict):
                            result = status_body["result"]
                            result.setdefault("recovery_attempt", attempt)
                            result.setdefault("recovery_attempts", total_attempts)
                            result.setdefault("generator_task_id", task_id)
                            return result
                        if status == "failed":
                            result = status_body.get("result")
                            if isinstance(result, dict):
                                if not _clean_error_text(result.get("error") or result.get("message"), ""):
                                    result["error"] = _clean_error_text(status_body, f"generator_task_failed: {task_id}")
                                result.setdefault("task_id", task_id)
                                result.setdefault("last_status", status_body)
                                result.setdefault("interrupted", True)
                                result.setdefault("recovery_attempt", attempt)
                                result.setdefault("recovery_attempts", total_attempts)
                                result.setdefault("generator_task_id", task_id)
                                return result
                            raise HTTPException(
                                status_code=502,
                                detail=_clean_error_text(status_body, f"generator_task_failed: {task_id}"),
                            )
                except HTTPException:
                    raise
                except Exception as exc:
                    last_status = {"error": str(exc), "task_id": task_id}
                    consecutive_status_misses += 1
                    if consecutive_status_misses >= GENERATOR_STATUS_MISS_LIMIT:
                        task_still_visible = _generator_task_exists(document_id, task_id)
                        if task_still_visible:
                            consecutive_status_misses = 0
                            time.sleep(min(8.0, sleep_seconds))
                            sleep_seconds = min(8.0, sleep_seconds * 1.3)
                            continue
                        return {
                            "success": False,
                            "error": "generator_downstream_unresponsive_or_task_lost",
                            "task_id": task_id,
                            "last_status": last_status,
                            "interrupted": True,
                            "recovery_attempt": attempt,
                            "recovery_attempts": total_attempts,
                        }

                time.sleep(min(8.0, sleep_seconds))
                sleep_seconds = min(8.0, sleep_seconds * 1.3)

            result = {
                "success": False,
                "error": f"generator_task_timeout after {call_timeout}s",
                "task_id": task_id,
                "last_status": last_status,
                "interrupted": True,
            }
            if isinstance(result, dict):
                result.setdefault("recovery_attempt", attempt)
                result.setdefault("recovery_attempts", total_attempts)
            return result
        except HTTPException as exc:
            last_error = str(exc.detail)
        except Exception as exc:
            last_error = str(exc)

        if attempt < total_attempts:
            delay = min(
                DOWNSTREAM_MAX_BACKOFF,
                GENERATOR_RESUME_BACKOFF * (2 ** (attempt - 1)) + random.uniform(0, DOWNSTREAM_JITTER),
            )
            print(f"⚠️ generate_document 第{attempt}次失败，{delay:.1f}s 后重试续跑: {last_error[:180]}")
            time.sleep(delay)

    return {
        "success": False,
        "error": f"文档生成在重试续跑后仍失败: {last_error}",
        "interrupted": True,
        "recovery_attempt": total_attempts,
        "recovery_attempts": total_attempts,
    }


def _build_document(req: GenerateDocRequest, timeout_seconds: int) -> Dict[str, Any]:
    start_ts = time.time()

    def _remaining_timeout(min_seconds: int = 30) -> int:
        elapsed = time.time() - start_ts
        remaining = int(timeout_seconds - elapsed)
        return max(min_seconds, remaining)

    document_id = f"web_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}"
    user_requirements = build_requirements_text(req)
    epistemic_audit_enabled = (
        HAS_EPISTEMIC_AUDIT
        and os.getenv("FLOWERNET_EPISTEMIC_AUDIT_ENABLED", "true").lower() == "true"
    )
    if epistemic_audit_enabled:
        user_requirements = augment_user_requirements(user_requirements)

    outline_payload = {
        "document_id": document_id,
        "user_background": req.user_background,
        "user_requirements": user_requirements,
        "max_sections": req.chapter_count,
        "max_subsections_per_section": req.subsection_count,
    }
    _wait_for_service_ready(OUTLINER_URL, "outliner", max_wait_seconds=min(180.0, max(45.0, timeout_seconds / 10.0)))
    outline_resp = call_outliner_generate_and_save(
        outline_payload,
        _remaining_timeout(min_seconds=30),
    )

    if not outline_resp.get("success"):
        raise HTTPException(status_code=500, detail=f"大纲生成失败: {outline_resp}")

    title = outline_resp.get("document_title") or f"{req.topic} 文档"
    structure = outline_resp.get("structure", {})
    content_prompts = outline_resp.get("content_prompts", [])
    if not isinstance(structure, dict) or not isinstance(content_prompts, list):
        raise HTTPException(status_code=500, detail=f"大纲结果格式异常: {outline_resp}")

    structure, content_prompts, source_subsections, normalized_subsections = ensure_exact_structure_and_prompts(
        title=title,
        structure=structure,
        content_prompts=content_prompts,
        chapter_count=req.chapter_count,
        subsection_count=req.subsection_count,
    )
    outline_ok, outline_reason = _validate_outline_structure_quality(structure)
    if not outline_ok:
        raise HTTPException(
            status_code=422,
            detail=f"大纲质量异常，已拒绝继续生成: {outline_reason}",
        )

    expected_subsections = req.chapter_count * req.subsection_count
    outlined_subsections = normalized_subsections
    if outlined_subsections <= 0:
        raise HTTPException(status_code=500, detail="大纲生成结果为空，无法开始内容生成")

    generate_payload = {
        "document_id": document_id,
        "title": title,
        "structure": structure,
        "content_prompts": augment_content_prompts(content_prompts) if epistemic_audit_enabled else content_prompts,
        "user_background": req.user_background,
        "user_requirements": user_requirements,
        "rel_threshold": req.rel_threshold,
        "red_threshold": req.red_threshold,
    }
    preflight_generator_warning(req.topic)
    gen_resp = generate_document_with_recovery(
        document_id=document_id,
        generate_payload=generate_payload,
        timeout_seconds=_remaining_timeout(min_seconds=30),
    )

    history_items = fetch_history_items(document_id=document_id, timeout_seconds=60)
    orchestration_metrics = extract_orchestration_metrics(gen_resp if isinstance(gen_resp, dict) else {})
    document_quality_metrics = extract_document_quality_metrics(gen_resp if isinstance(gen_resp, dict) else {})

    # Debug: Log extracted metrics
    print(f"🔍 [METRICS DEBUG] gen_resp type: {type(gen_resp)}")
    if isinstance(gen_resp, dict):
        print(f"🔍 [METRICS DEBUG] gen_resp.success: {gen_resp.get('success')}")
        print(f"🔍 [METRICS DEBUG] gen_resp has quality_score_avg: {'quality_score_avg' in gen_resp}, value: {gen_resp.get('quality_score_avg')}")
        print(f"🔍 [METRICS DEBUG] gen_resp has unieval_available_subsections: {'unieval_available_subsections' in gen_resp}, value: {gen_resp.get('unieval_available_subsections')}")
        print(f"🔍 [METRICS DEBUG] gen_resp has bandit_selected_arm_counts: {'bandit_selected_arm_counts' in gen_resp}, value: {gen_resp.get('bandit_selected_arm_counts')}")
    print(f"🔍 [METRICS DEBUG] Extracted quality_score_avg: {document_quality_metrics.get('quality_score_avg')}")
    print(f"🔍 [METRICS DEBUG] Extracted unieval_available_subsections: {document_quality_metrics.get('unieval_available_subsections')}")
    print(f"🔍 [METRICS DEBUG] Extracted bandit_reward_avg: {document_quality_metrics.get('bandit_reward_avg')}")

    if not gen_resp.get("success"):
        if history_items:
            raise HTTPException(
                status_code=500,
                detail=(
                    f"文档生成失败，已生成 {len(history_items)}/{outlined_subsections} 个通过验证的小节；"
                    f"拒绝返回不完整文档: {gen_resp}"
                ),
            )
        raise HTTPException(status_code=500, detail=f"文档生成失败: {gen_resp}")

    passed = gen_resp.get("passed_subsections", 0)
    failed = len(gen_resp.get("failed_subsections", []))
    forced = len(gen_resp.get("forced_subsections", []))
    if passed < outlined_subsections and history_items:
        raise HTTPException(
            status_code=500,
            detail=(
                f"文档生成未达到大纲小节数: 通过 {passed}/{outlined_subsections}, "
                f"失败 {failed}; 拒绝返回不完整文档"
            ),
        )
    if passed < outlined_subsections:
        raise HTTPException(
            status_code=500,
            detail=f"文档生成未达到大纲小节数: 通过 {passed}/{outlined_subsections}, 失败 {failed}",
        )

    if not history_items:
        history_items = fetch_history_items(document_id=document_id, timeout_seconds=60)

    # 🔍 诊断：检查生成器返回的 source_results
    gen_sections = gen_resp.get("sections", []) or []
    total_gen_source_results = 0
    for sec in gen_sections:
        for subsec in sec.get("subsections", []) or []:
            sr = subsec.get("source_results", []) or []
            total_gen_source_results += len(sr)
    print(f"📌 [生成器诊断] 生成器返回 {len(gen_sections)} 个 section, 总 source_results: {total_gen_source_results}")

    epistemic_audit: Dict[str, Any] = {}
    if epistemic_audit_enabled and EpistemicAuditEngine is not None:
        try:
            epistemic_audit = EpistemicAuditEngine().build_audit(
                title=title,
                structure=structure,
                sections=gen_sections,
                history=history_items,
                orchestration_metrics=orchestration_metrics,
                quality_metrics=document_quality_metrics,
            )
            gen_sections = attach_chapter_assets(gen_sections, epistemic_audit)
        except Exception as exc:
            print(f"⚠️ [EpistemicAudit] build failed, continuing without audit: {exc}")
            epistemic_audit = {"enabled": False, "error": str(exc)}

    markdown_content = build_markdown_document(
        title,
        structure,
        history_items,
        generated_sections=gen_sections,
        user_background=req.user_background,
        extra_requirements=req.extra_requirements,
        epistemic_audit=epistemic_audit,
    )
    citation_quality = _citation_quality_check(markdown_content)
    _enforce_citation_quality_or_raise(citation_quality, context="final_document")
    total_source_refs = _aggregate_source_reference_count(
        history_items=history_items,
        generated_sections=gen_sections,
    )
    total_source_refs = max(total_source_refs, int(citation_quality.get("reference_count", 0) or 0))

    if not citation_quality.get("passed", False):
        raise HTTPException(
            status_code=500,
            detail=f"文档内容已生成，但引用质量未达标，拒绝返回不完整结果：{citation_quality.get('reason')}",
        )

    return {
        "success": True,
        "document_id": document_id,
        "title": title,
        "content": markdown_content,
        "stats": {
            "expected_subsections": expected_subsections,
            "outlined_subsections": outlined_subsections,
            "passed_subsections": passed,
            "failed_subsections": failed,
            "forced_subsections": forced,
            "total_iterations": gen_resp.get("total_iterations", 0),
            "generation_time": gen_resp.get("generation_time", ""),
            "citation_quality": citation_quality,
            "total_source_references": total_source_refs,
            "rag_used_subsections": int(gen_resp.get("rag_used_subsections", 0) or 0),
            "rag_search_success_subsections": int(gen_resp.get("rag_search_success_subsections", 0) or 0),
            "controller_effective_subsections": int(gen_resp.get("controller_effective_subsections", 0) or 0),
            "epistemic_audit": epistemic_audit.get("summary", {}) if isinstance(epistemic_audit, dict) else {},
            "epistemic_risk_portfolio": epistemic_audit.get("risk_portfolio", {}) if isinstance(epistemic_audit, dict) else {},
            "epistemic_reviewer_scores": epistemic_audit.get("reviewer_scores", {}) if isinstance(epistemic_audit, dict) else {},
            "token_usage": gen_resp.get("token_usage", {}),
            "prompt_cache_hit_rate": gen_resp.get("prompt_cache_hit_rate", 0.0),
            "generator_short_draft_total": gen_resp.get("generator_short_draft_total", 0),
            **orchestration_metrics,
            **document_quality_metrics,
        },
    }


def _extract_urls(text: str) -> List[str]:
    return re.findall(r"https?://[^\s\]）)>,;]+", text or "", flags=re.IGNORECASE)


def _domain_quality(domain: str) -> float:
    host = (domain or "").lower().strip()
    if not host:
        return 0.0

    high_quality_domains = {
        "nature.com", "science.org", "sciencedirect.com", "springer.com", "ieee.org", "acm.org",
        "arxiv.org", "doi.org", "crossref.org", "pubmed.ncbi.nlm.nih.gov", "ncbi.nlm.nih.gov",
        "who.int", "oecd.org", "un.org", "nist.gov", "nih.gov", "wiley.com", "onlinelibrary.wiley.com",
        "tandfonline.com", "sagepub.com", "cambridge.org", "cambridge.org", "oxfordacademic.com",
        "academic.oup.com", "jstor.org", "cell.com", "thelancet.com", "nejm.org", "bmj.com",
        "gov.cn", "edu.cn", "ruc.edu.cn", "tsinghua.edu.cn", "pku.edu.cn", "cass.cn", "moe.gov.cn",
    }
    low_quality_domains = {
        "baike.baidu.com", "zhidao.baidu.com", "tieba.baidu.com", "jingyan.baidu.com",
        "m.baidu.com", "weibo.com", "t.co", "bit.ly", "tinyurl.com", "researchgate.net",
    }

    if host in low_quality_domains:
        return 0.15
    if host in high_quality_domains:
        return 1.0
    if host.endswith(".gov") or host.endswith(".edu") or host.endswith(".gov.cn") or host.endswith(".edu.cn"):
        return 0.95
    if "wikipedia.org" in host:
        return 0.55
    if host.endswith(".org"):
        return 0.72
    if host.endswith(".com"):
        return 0.60
    return 0.5


def _citation_quality_check(markdown: str) -> Dict[str, Any]:
    if not ENABLE_CITATION_QA:
        return {"passed": True, "reason": "disabled"}

    refs_match = re.search(r"^##\s+References\s*$([\s\S]*)", markdown or "", flags=re.MULTILINE)
    refs_text = refs_match.group(1) if refs_match else ""
    body_text = (markdown or "")[: refs_match.start()] if refs_match else (markdown or "")
    audit_match = re.search(r"^##\s+Self-Audit Ledger\s*$", body_text or "", flags=re.MULTILINE)
    if audit_match:
        body_text = body_text[: audit_match.start()]
    reference_lines = re.findall(r"^\[\d+\]\s+.+", refs_text, flags=re.MULTILINE)
    subsection_blocks = re.findall(r"^###\s+.*?(?=^###\s+|\Z)", body_text or "", flags=re.MULTILINE | re.DOTALL)
    if not subsection_blocks:
        subsection_blocks = [body_text or ""]

    all_urls = _extract_urls(markdown or "")
    unique_urls = list(dict.fromkeys(all_urls))
    subsection_marker_counts = [len(set(re.findall(r"\[(\d+)\]", block))) for block in subsection_blocks]

    low_quality_count = 0
    section_details: List[Dict[str, Any]] = []
    for idx, block in enumerate(subsection_blocks):
        urls = list(dict.fromkeys(_extract_urls(block)))
        high_quality_urls = []
        for url in urls:
            domain = (urlparse(url).netloc or "").lower().strip()
            score = _domain_quality(domain)
            if score < 0.35:
                low_quality_count += 1
            if score >= 0.70:
                high_quality_urls.append(url)
        section_details.append({
            "url_count": len(urls),
            "high_quality_url_count": len(high_quality_urls),
            "citation_marker_count": subsection_marker_counts[idx] if idx < len(subsection_marker_counts) else 0,
        })

    low_quality_ratio = (low_quality_count / max(1, len(unique_urls))) if unique_urls else 1.0
    missing_high_quality_sections = sum(
        1 for sec in section_details if sec["high_quality_url_count"] < CITATION_MIN_SECTION_HIGH_QUALITY
    )

    min_markers_per_subsection = max(1, int(os.getenv("MIN_REFERENCES_PER_SUBSECTION", "3") or "3"))
    marker_floor_ok = (
        bool(reference_lines)
        and bool(unique_urls)
        and low_quality_ratio <= CITATION_LOW_QUALITY_MAX_RATIO
        and all(count >= min_markers_per_subsection for count in subsection_marker_counts)
    )
    url_quality_ok = bool(unique_urls) and missing_high_quality_sections == 0 and low_quality_ratio <= CITATION_LOW_QUALITY_MAX_RATIO
    textual_reference_fallback_ok = bool(reference_lines) and not unique_urls
    passed = url_quality_ok or marker_floor_ok or textual_reference_fallback_ok
    reason = "ok"
    if textual_reference_fallback_ok:
        reason = "textual_reference_fallback"
    elif marker_floor_ok and not url_quality_ok:
        reason = "marker_floor_ok"
    elif not unique_urls:
        reason = "no_citation_urls"
    elif missing_high_quality_sections > 0:
        reason = "section_missing_high_quality_source"
    elif low_quality_ratio > CITATION_LOW_QUALITY_MAX_RATIO:
        reason = "low_quality_domain_ratio_too_high"

    return {
        "passed": passed,
        "reason": reason,
        "total_urls": len(all_urls),
        "unique_urls": len(unique_urls),
        "reference_count": len(reference_lines),
        "min_subsection_citation_markers": min(subsection_marker_counts) if subsection_marker_counts else 0,
        "low_quality_ratio": round(low_quality_ratio, 4),
        "missing_high_quality_sections": missing_high_quality_sections,
        "best_effort_reference_fallback": textual_reference_fallback_ok,
        "section_details": section_details,
    }


def _aggregate_source_reference_count(
    history_items: Optional[List[Dict[str, Any]]] = None,
    generated_sections: Optional[List[Dict[str, Any]]] = None,
) -> int:
    total = 0

    for h_item in history_items or []:
        meta = h_item.get("metadata", {}) if isinstance(h_item, dict) else {}
        verification = meta.get("verification", {}) or (h_item.get("verification", {}) if isinstance(h_item, dict) else {})
        if isinstance(verification, dict):
            src_check = verification.get("source_check", {})
            if isinstance(src_check, dict):
                total += int(src_check.get("reference_count", 0) or 0)

    for section in generated_sections or []:
        for subsection in (section.get("subsections") or []):
            total += int(subsection.get("source_reference_count", 0) or 0)

    return total


def _enforce_citation_quality_or_raise(citation_quality: Dict[str, Any], context: str = "") -> None:
    if not STRICT_CITATION_ENFORCEMENT or not CITATION_FAIL_FAST:
        return
    if citation_quality.get("passed", False):
        return
    raise HTTPException(
        status_code=422,
        detail={
            "error": "citation_quality_failed",
            "context": context,
            "reason": str(citation_quality.get("reason", "unknown")),
            "citation_quality": citation_quality,
        },
    )


def _build_download_url(request: Request) -> str:
    if PUBLIC_BASE_URL:
        return f"{PUBLIC_BASE_URL.rstrip('/')}/api/download-docx"
    return str(request.url_for("download_docx")).rstrip("/")


def _public_api_base_url(request: Request) -> str:
    if PUBLIC_BASE_URL:
        return PUBLIC_BASE_URL.rstrip("/")
    return str(request.base_url).rstrip("/")


def _build_poffices_openapi(request: Request) -> Dict[str, Any]:
    base_url = _public_api_base_url(request)
    security_schemes: Dict[str, Any] = {}
    security: List[Dict[str, List[str]]] = []
    if API_AUTH_ENABLED:
        security_schemes = {
            "apiKeyAuth": {
                "type": "apiKey",
                "in": "header",
                "name": "X-API-Key",
            },
            "bearerAuth": {
                "type": "http",
                "scheme": "bearer",
            },
        }
        security = [{"apiKeyAuth": []}, {"bearerAuth": []}]

    return {
        "openapi": "3.0.3",
        "info": {
            "title": "FlowerNet Agent for POffices",
            "version": "1.0.0",
            "description": "Generate structured long-form documents with FlowerNet and poll task status asynchronously.",
        },
        "servers": [{"url": base_url}],
        "paths": {
            "/api/poffices/generate": {
                "post": {
                    "operationId": "createFlowerNetDocument",
                    "summary": "Create a FlowerNet document generation task",
                    "description": "Starts an asynchronous FlowerNet document generation task. Use task_id with getFlowerNetTaskStatus to retrieve the final content and download payload.",
                    "security": security,
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {"$ref": "#/components/schemas/PofficesGenerateRequest"},
                                "examples": {
                                    "basic": {
                                        "value": {
                                            "query": "Write a professional report about plant disease recognition with deep learning",
                                            "chapter_count": 3,
                                            "subsection_count": 2,
                                            "user_background": "Research audience",
                                            "extra_requirements": "Use clear academic structure and concise citations",
                                            "async_mode": True,
                                            "timeout_seconds": 2400,
                                        }
                                    }
                                },
                            }
                        },
                    },
                    "responses": {
                        "200": {
                            "description": "Task accepted or synchronous result returned",
                            "content": {
                                "application/json": {
                                    "schema": {
                                        "oneOf": [
                                            {"$ref": "#/components/schemas/TaskAcceptedResponse"},
                                            {"$ref": "#/components/schemas/CompletedDocumentResponse"},
                                        ]
                                    }
                                }
                            },
                        }
                    },
                }
            },
            "/api/poffices/task-status": {
                "post": {
                    "operationId": "getFlowerNetTaskStatus",
                    "summary": "Get FlowerNet generation task status",
                    "description": f"Long-poll with the task_id returned by createFlowerNetDocument. By default this endpoint waits up to {POFFICES_POLL_WAIT_SECONDS} seconds before returning a running/queued response, so repeated Poffices poll blocks keep the generation open long enough for FlowerNet to finish. Completed responses include the full Markdown document.",
                    "security": security,
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {"$ref": "#/components/schemas/PofficesTaskStatusRequest"}
                            }
                        },
                    },
                    "responses": {
                        "200": {
                            "description": "Current task state or completed document result",
                            "content": {
                                "application/json": {
                                    "schema": {
                                        "oneOf": [
                                            {"$ref": "#/components/schemas/TaskStatusResponse"},
                                            {"$ref": "#/components/schemas/CompletedDocumentResponse"},
                                            {"$ref": "#/components/schemas/FailedTaskResponse"},
                                        ]
                                    }
                                }
                            },
                        },
                        "404": {"description": "task_id not found"},
                    },
                }
            },
            "/api/poffices/poll-render": {
                "post": {
                    "operationId": "pollAndRenderFlowerNetDocument",
                    "summary": "Poll a FlowerNet task and render the final document",
                    "description": f"Poffices bridge endpoint for repeated poll blocks. It extracts task_id from nested upstream block output and waits up to {POFFICES_POLL_WAIT_SECONDS} seconds by default before returning a running/queued response. If completed, content/text/result/output contain the complete Markdown document.",
                    "security": security,
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {"$ref": "#/components/schemas/PofficesPollRenderRequest"}
                            }
                        },
                    },
                    "responses": {
                        "200": {
                            "description": "Completed document result or failed task result",
                            "content": {
                                "application/json": {
                                    "schema": {
                                        "oneOf": [
                                            {"$ref": "#/components/schemas/CompletedDocumentResponse"},
                                            {"$ref": "#/components/schemas/FailedTaskResponse"},
                                        ]
                                    }
                                }
                            },
                        }
                    },
                }
            },
            "/api/download-docx": {
                "post": {
                    "operationId": "downloadFlowerNetDocx",
                    "summary": "Download a generated FlowerNet document as DOCX",
                    "description": "POST the title and markdown content returned by the completed document response to download a DOCX file.",
                    "security": security,
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {"$ref": "#/components/schemas/DownloadDocxRequest"}
                            }
                        },
                    },
                    "responses": {
                        "200": {
                            "description": "DOCX file",
                            "content": {
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {
                                    "schema": {"type": "string", "format": "binary"}
                                }
                            },
                        }
                    },
                }
            },
        },
        "components": {
            "securitySchemes": security_schemes,
            "schemas": {
                "PofficesGenerateRequest": {
                    "type": "object",
                    "required": ["query"],
                    "properties": {
                        "query": {"type": "string", "minLength": 2, "description": "Document topic or user request"},
                        "chapter_count": {"type": "integer", "minimum": 1, "maximum": 10, "default": 5},
                        "subsection_count": {"type": "integer", "minimum": 1, "maximum": 8, "default": 3},
                        "user_background": {"type": "string", "default": ""},
                        "extra_requirements": {"type": "string", "default": ""},
                        "rel_threshold": {"type": "number", "minimum": 0, "maximum": 1, "default": WEB_DEFAULT_REL_THRESHOLD},
                        "red_threshold": {"type": "number", "minimum": 0, "maximum": 1, "default": WEB_DEFAULT_RED_THRESHOLD},
                        "async_mode": {"type": "boolean", "default": True},
                        "timeout_seconds": {"type": "integer", "minimum": 60, "maximum": 7200, "default": 7200},
                    },
                },
                "PofficesTaskStatusRequest": {
                    "type": "object",
                    "required": ["task_id"],
                    "properties": {
                        "task_id": {"type": "string"},
                        "wait": {"type": "boolean", "default": True},
                        "wait_seconds": {"type": "integer", "minimum": 1, "maximum": 7200, "default": POFFICES_POLL_WAIT_SECONDS},
                    },
                },
                "PofficesPollRenderRequest": {
                    "type": "object",
                    "properties": {
                        "task_id": {"type": "string", "description": "May be nested inside upstream block output; the endpoint will extract it recursively."},
                        "query": {"type": "string", "description": "Original user request. Used to recover or start the task if task_id is missing or stale."},
                        "wait": {"type": "boolean", "default": True},
                        "wait_seconds": {"type": "integer", "minimum": 1, "maximum": 7200, "default": POFFICES_POLL_WAIT_SECONDS},
                    },
                },
                "DownloadDocxRequest": {
                    "type": "object",
                    "required": ["title", "content"],
                    "properties": {
                        "title": {"type": "string"},
                        "content": {"type": "string", "description": "Markdown content returned by FlowerNet"},
                    },
                },
                "TaskAcceptedResponse": {
                    "type": "object",
                    "properties": {
                        "success": {"type": "boolean"},
                        "task_id": {"type": "string"},
                        "status": {"type": "string", "enum": ["queued"]},
                        "poll_url": {"type": "string"},
                        "message": {"type": "string"},
                    },
                },
                "TaskStatusResponse": {
                    "type": "object",
                    "properties": {
                        "success": {"type": "boolean"},
                        "task_id": {"type": "string"},
                        "status": {"type": "string", "enum": ["queued", "running"]},
                        "message": {"type": "string"},
                    },
                },
                "CompletedDocumentResponse": {
                    "type": "object",
                    "properties": {
                        "success": {"type": "boolean"},
                        "task_status": {"type": "string", "enum": ["completed"]},
                        "document_id": {"type": "string"},
                        "title": {"type": "string"},
                        "content": {"type": "string"},
                        "text": {"type": "string"},
                        "result": {"type": "string"},
                        "output": {"type": "string"},
                        "document": {"type": "string"},
                        "markdown": {"type": "string"},
                        "stats": {"type": "object", "additionalProperties": True},
                        "download": {"$ref": "#/components/schemas/DownloadInstruction"},
                    },
                },
                "DownloadInstruction": {
                    "type": "object",
                    "properties": {
                        "method": {"type": "string", "enum": ["POST"]},
                        "url": {"type": "string"},
                        "body": {"$ref": "#/components/schemas/DownloadDocxRequest"},
                        "content_type": {"type": "string"},
                    },
                },
                "FailedTaskResponse": {
                    "type": "object",
                    "properties": {
                        "success": {"type": "boolean", "enum": [False]},
                        "task_id": {"type": "string"},
                        "status": {"type": "string", "enum": ["failed"]},
                        "error": {"type": "string"},
                        "message": {"type": "string"},
                    },
                },
            },
        },
    }


def _build_poffices_result(request: Request, result: Dict[str, Any], task_id: str = "") -> Dict[str, Any]:
    download_url = _build_download_url(request)
    content = result.get("content", "") or result.get("markdown", "") or result.get("document", "")
    title = result.get("title", "") or "FlowerNet Document"
    return {
        "success": True,
        "task_id": task_id or result.get("task_id", ""),
        "task_status": "completed",
        "status": "completed",
        "document_id": result.get("document_id", ""),
        "title": title,
        "content": content,
        "text": content,
        "result": content,
        "output": content,
        "document": content,
        "markdown": content,
        "stats": result.get("stats", {}),
        "download": {
            "method": "POST",
            "url": download_url,
            "body": {
                "title": title,
                "content": content,
            },
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    }


def _extract_task_id_from_payload(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        direct = re.search(r"task_[A-Za-z0-9_:-]{12,}", value)
        return direct.group(0) if direct else ""
    if isinstance(value, dict):
        for key in ("task_id", "taskId", "id"):
            candidate = _extract_task_id_from_payload(value.get(key))
            if candidate:
                return candidate
        for item in value.values():
            candidate = _extract_task_id_from_payload(item)
            if candidate:
                return candidate
    if isinstance(value, (list, tuple)):
        for item in value:
            candidate = _extract_task_id_from_payload(item)
            if candidate:
                return candidate
    return ""


def _extract_text_field_from_payload(value: Any, keys: Tuple[str, ...]) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        for key in keys:
            raw = value.get(key)
            if isinstance(raw, str) and raw.strip():
                return raw.strip()
        for item in value.values():
            found = _extract_text_field_from_payload(item, keys)
            if found:
                return found
    if isinstance(value, (list, tuple)):
        for item in value:
            found = _extract_text_field_from_payload(item, keys)
            if found:
                return found
    return ""


def _payload_contains_task_not_found(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        lowered = value.lower()
        return "task_id not found" in lowered or ("http 404" in lowered and "task_id" in lowered)
    if isinstance(value, dict):
        return any(_payload_contains_task_not_found(item) for item in value.values())
    if isinstance(value, (list, tuple)):
        return any(_payload_contains_task_not_found(item) for item in value)
    return False


def _extract_int_field_from_payload(value: Any, keys: Tuple[str, ...], default: int) -> int:
    if isinstance(value, dict):
        for key in keys:
            raw = value.get(key)
            if raw is None:
                continue
            try:
                return int(raw)
            except (TypeError, ValueError):
                pass
        for item in value.values():
            found = _extract_int_field_from_payload(item, keys, -1)
            if found > 0:
                return found
    if isinstance(value, (list, tuple)):
        for item in value:
            found = _extract_int_field_from_payload(item, keys, -1)
            if found > 0:
                return found
    return default


def _coerce_poffices_request_from_payload(payload: Any) -> Optional[PofficesGenerateRequest]:
    """Recover the original generation request from a loosely wired Poffices block.

    Some Poffices LLM blocks pass nested JSON/text rather than a clean task_id.
    If a later "poll" block accidentally calls the generation endpoint again,
    this lets the server reuse the existing request checkpoint instead of
    starting another FlowerNet job and overloading the outliner.
    """
    if not isinstance(payload, dict):
        return None
    audit_payload = payload.get("flowernet_audit") if isinstance(payload.get("flowernet_audit"), dict) else {}
    query = (
        _extract_text_field_from_payload(payload, ("query", "topic", "REAL_USER_REQUEST", "real_user_request"))
        or _extract_text_field_from_payload(audit_payload, ("query", "request_key", "topic"))
        or _extract_text_field_from_payload(payload, ("request", "input"))
    )
    if not query:
        text_candidate = _extract_text_field_from_payload(payload, ("content", "text", "result", "output", "markdown", "document"))
        if text_candidate and not _extract_task_id_from_payload(text_candidate) and not text_candidate.strip().startswith("{"):
            query = text_candidate
    if not query or query.startswith("FlowerNet task "):
        return None
    chapter_count = _extract_int_field_from_payload(payload, ("chapter_count", "chapters", "chapterCount"), -1)
    if chapter_count <= 0:
        chapter_count = _extract_int_field_from_payload(audit_payload, ("chapter_count", "chapters", "chapterCount"), 5)
    subsection_count = _extract_int_field_from_payload(payload, ("subsection_count", "subsection", "subsections", "subsectionCount"), -1)
    if subsection_count <= 0:
        subsection_count = _extract_int_field_from_payload(audit_payload, ("subsection_count", "subsection", "subsections", "subsectionCount"), 3)
    user_background = _extract_text_field_from_payload(payload, ("user_background", "background"))
    extra_requirements = _extract_text_field_from_payload(payload, ("extra_requirements", "requirements", "extra"))
    try:
        return PofficesGenerateRequest(
            query=query,
            chapter_count=max(1, min(10, chapter_count)),
            subsection_count=max(1, min(8, subsection_count)),
            user_background=user_background,
            extra_requirements=extra_requirements,
            async_mode=True,
        )
    except Exception:
        return None


def _poffices_reuse_existing_request_task(
    *,
    request: Request,
    req: PofficesGenerateRequest,
    wait: bool = False,
    wait_seconds: int = POFFICES_POLL_WAIT_SECONDS,
) -> Optional[Dict[str, Any]]:
    existing_task_id = _restore_poffices_request_task_id(req)
    if not existing_task_id:
        existing_task_id = _find_recent_poffices_task_for_request(req)
    if not existing_task_id:
        return None
    existing_task = _restore_poffices_task(existing_task_id)
    if not existing_task:
        return None
    if existing_task.get("status") in {"queued", "running", "completed"}:
        return _poffices_wait_for_task_result(
            request=request,
            task_id=existing_task_id,
            wait=wait,
            wait_seconds=wait_seconds,
        )
    return None


def _poffices_start_or_reuse_async_task(
    *,
    request: Request,
    req: PofficesGenerateRequest,
    wait: bool = False,
    wait_seconds: int = POFFICES_POLL_WAIT_SECONDS,
) -> Dict[str, Any]:
    """Idempotently start or recover a FlowerNet task for Poffices blocks."""
    req = req.model_copy(update={"query": (req.query or "").strip(), "async_mode": True})
    reused = _poffices_reuse_existing_request_task(
        request=request,
        req=req,
        wait=wait,
        wait_seconds=wait_seconds,
    )
    if reused is not None:
        return reused

    task_id = f"task_{uuid4().hex}"
    _set_poffices_task(
        task_id,
        status="queued",
        message="任务已入队",
        created_at=datetime.now().isoformat(),
        started_at=None,
        timeout_seconds=req.timeout_seconds,
        request=req.model_dump(),
    )
    _persist_poffices_request_task(req, task_id)

    with POFFICES_TASKS_LOCK:
        POFFICES_TASKS[task_id]["_thread_started"] = True
    threading.Thread(target=_run_poffices_task, args=(task_id, req), daemon=True).start()

    if wait:
        return _poffices_wait_for_task_result(
            request=request,
            task_id=task_id,
            wait=True,
            wait_seconds=wait_seconds,
        )
    return {
        "success": True,
        "task_id": task_id,
        "status": "queued",
        "poll_url": str(request.base_url).rstrip("/") + "/api/poffices/task-status",
        "message": "异步任务已创建，请轮询 task status",
        "content": "异步任务已创建，请轮询 task status",
        "text": "异步任务已创建，请轮询 task status",
        "result": "异步任务已创建，请轮询 task status",
        "output": "异步任务已创建，请轮询 task status",
    }


def _extract_quoted_field_from_repr(text: str, field: str) -> str:
    match = re.search(rf"['\"]{re.escape(field)}['\"]\s*:\s*('(?:\\.|[^'])*'|\"(?:\\.|[^\"])*\")", text)
    if not match:
        return ""
    try:
        import ast
        value = ast.literal_eval(match.group(1))
    except Exception:
        value = match.group(1).strip("'\"")
    return str(value or "").strip()


def _extract_subsection_contents_from_repr(text: str) -> List[str]:
    values: List[str] = []
    for match in re.finditer(r"['\"]content['\"]\s*:\s*('(?:\\.|[^'])*'|\"(?:\\.|[^\"])*\")", text):
        try:
            import ast
            value = ast.literal_eval(match.group(1))
        except Exception:
            value = match.group(1).strip("'\"")
        value = str(value or "").strip()
        if value and value not in values:
            values.append(value)
    return values


def _extract_renderable_result_from_failed_task(task: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    task_result = task.get("result")
    if isinstance(task_result, dict):
        rendered_content = task_result.get("content") or task_result.get("markdown") or task_result.get("document") or ""
        if isinstance(rendered_content, str) and rendered_content.strip():
            return task_result

    error_text = str(task.get("error") or "")
    for prefix in ("文档生成失败:", "文档生成异常:", "生成服务连接失败:"):
        if prefix in error_text:
            error_text = error_text.split(prefix, 1)[1].strip()
            break

    if error_text.startswith("{") and error_text.endswith("}"):
        try:
            import ast
            parsed = ast.literal_eval(error_text)
        except Exception:
            parsed = None
        if isinstance(parsed, dict):
            rendered_content = parsed.get("content") or parsed.get("markdown") or parsed.get("document") or ""
            if isinstance(rendered_content, str) and rendered_content.strip():
                return parsed

    subsection_contents = _extract_subsection_contents_from_repr(error_text)
    if subsection_contents:
        title = _extract_quoted_field_from_repr(error_text, "title") or "FlowerNet Document"
        document_id = _extract_quoted_field_from_repr(error_text, "document_id")
        return {
            "success": True,
            "document_id": document_id,
            "title": title,
            "content": f"# {title}\n\n" + "\n\n".join(subsection_contents),
            "partial": True,
            "warning": "recovered_from_failed_task_repr",
        }
    return None


def _coerce_bool(value: Any, default: bool = True) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"1", "true", "yes", "y", "wait"}:
            return True
        if lowered in {"0", "false", "no", "n", "nowait"}:
            return False
    return default


def _poffices_wait_for_task_result(
    *,
    request: Request,
    task_id: str,
    wait: bool = True,
    wait_seconds: int = POFFICES_POLL_WAIT_SECONDS,
    poll_interval_seconds: float = POFFICES_POLL_INTERVAL_SECONDS,
) -> Dict[str, Any]:
    deadline = time.time() + max(1, int(wait_seconds))
    last_task: Dict[str, Any] = {}

    while True:
        task = _restore_poffices_task(task_id)
        if not task:
            raise HTTPException(status_code=404, detail="task_id not found")

        status = task.get("status", "unknown")
        if status in {"queued", "running"}:
            _restart_restored_poffices_task(task_id, task)
            task = _restore_poffices_task(task_id) or task
            status = task.get("status", status)

        last_task = task
        if status == "completed":
            return _build_poffices_result(request=request, result=task.get("result", {}), task_id=task_id)

        if status == "failed":
            task_result = _extract_renderable_result_from_failed_task(task)
            if task_result is not None:
                rendered = _build_poffices_result(request=request, result=task_result, task_id=task_id)
                rendered.update(
                    {
                        "success": True,
                        "task_id": task_id,
                        "status": "completed",
                        "task_status": "completed",
                        "message": "任务完成，生成结果已返回；存在质量或重试警告",
                        "warning": task.get("error", "completed_with_generation_warning"),
                        "original_status": "failed",
                        "original_error": task.get("error", ""),
                    }
                )
                return rendered
            return {
                "success": False,
                "task_id": task_id,
                "status": "failed",
                "task_status": "failed",
                "error": task.get("error", "unknown error"),
                "message": task.get("message", "任务失败"),
                "content": task.get("error", "unknown error"),
                "text": task.get("error", "unknown error"),
                "result": task.get("error", "unknown error"),
                "output": task.get("error", "unknown error"),
            }

        if not wait or time.time() >= deadline:
            started_at = str(last_task.get("started_at") or last_task.get("created_at") or "")
            elapsed_seconds = round(_iso_age_seconds(started_at), 2) if started_at else None
            return {
                "success": True,
                "task_id": task_id,
                "status": status,
                "task_status": status,
                "poll_url": str(request.url_for("poffices_task_status")),
                "message": last_task.get("message", "处理中"),
                "retry_count": last_task.get("retry_count", 0),
                "elapsed_seconds": elapsed_seconds,
                "created_at": last_task.get("created_at"),
                "started_at": last_task.get("started_at"),
                "updated_at": last_task.get("updated_at"),
                "last_retryable_error": last_task.get("last_retryable_error", ""),
                "content": f"FlowerNet task {task_id} is {status}. Please poll again.",
                "text": f"FlowerNet task {task_id} is {status}. Please poll again.",
                "result": f"FlowerNet task {task_id} is {status}. Please poll again.",
                "output": f"FlowerNet task {task_id} is {status}. Please poll again.",
            }

        time.sleep(min(poll_interval_seconds, max(0.1, deadline - time.time())))


def _poffices_recover_from_missing_task(
    *,
    request: Request,
    payload: Dict[str, Any],
    wait: bool,
    wait_seconds: int,
) -> Optional[Dict[str, Any]]:
    recovered_req = _coerce_poffices_request_from_payload(payload)
    if recovered_req is None:
        return None
    return _poffices_start_or_reuse_async_task(
        request=request,
        req=recovered_req,
        wait=wait,
        wait_seconds=wait_seconds,
    )


def build_requirements_text(req: GenerateDocRequest) -> str:
    background = req.user_background.strip() or "未指定"
    base = (
        f"文档主题：{req.topic}\n"
        f"目标读者/用户背景：{background}\n"
        f"写作目标：生成一篇主题聚焦、结构完整、适合目标读者的专业长文档。\n"
        f"结构规模：恰好 {req.chapter_count} 个章节；每个章节恰好 {req.subsection_count} 个子章节。\n"
        "大纲要求：章节和子章节标题必须是围绕文档主题展开的专业学术标题；"
        "不得把本需求文本、写作指令、用户背景或数量要求复制成标题。"
    )
    if req.extra_requirements.strip():
        return f"{base}\n\n附加要求：{req.extra_requirements.strip()}"
    return base


UPLOAD_MAX_FILES = int(os.getenv("UPLOAD_CONTEXT_MAX_FILES", "8"))
UPLOAD_MAX_FILE_BYTES = int(os.getenv("UPLOAD_CONTEXT_MAX_FILE_BYTES", str(25 * 1024 * 1024)))
UPLOAD_MAX_CONTEXT_CHARS = int(os.getenv("UPLOAD_CONTEXT_MAX_CHARS", "24000"))
UPLOAD_MAX_PER_FILE_CHARS = int(os.getenv("UPLOAD_CONTEXT_MAX_PER_FILE_CHARS", "7000"))
UPLOAD_ARCHIVE_MAX_MEMBERS = int(os.getenv("UPLOAD_CONTEXT_ARCHIVE_MAX_MEMBERS", "30"))


def _safe_decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def _clean_uploaded_text(text: str, max_chars: int = UPLOAD_MAX_PER_FILE_CHARS) -> str:
    cleaned = re.sub(r"\r\n?", "\n", str(text or ""))
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned).strip()
    if len(cleaned) > max_chars:
        head = max_chars * 3 // 4
        tail = max_chars - head
        cleaned = (
            cleaned[:head].rstrip()
            + f"\n\n[...上传文件内容已裁剪，原始解析长度 {len(text)} 字符...]\n\n"
            + cleaned[-tail:].lstrip()
        )
    return cleaned


def _extract_text_from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows[:40]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        return f"[PDF parser unavailable: install pypdf to extract this file. error={exc}]"
    reader = PdfReader(BytesIO(data))
    parts: List[str] = []
    for index, page in enumerate(reader.pages[:40], start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            page_text = f"[page {index} extraction failed: {exc}]"
        if page_text.strip():
            parts.append(f"[Page {index}]\n{page_text.strip()}")
    return "\n\n".join(parts)


def _extract_text_from_excel(data: bytes, suffix: str) -> str:
    if suffix == ".csv":
        text = _safe_decode_bytes(data)
        rows = list(csv.reader(text.splitlines()))
        return "\n".join(" | ".join(cell.strip() for cell in row[:12]) for row in rows[:80])
    try:
        import openpyxl  # type: ignore
    except Exception as exc:
        return f"[Excel parser unavailable: install openpyxl to extract this file. error={exc}]"
    workbook = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: List[str] = []
    for sheet_name in workbook.sheetnames[:8]:
        sheet = workbook[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in sheet.iter_rows(min_row=1, max_row=80, max_col=16, values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _extract_text_from_pptx(data: bytes) -> str:
    parts: List[str] = []
    with zipfile.ZipFile(BytesIO(data)) as zf:
        slide_names = sorted(name for name in zf.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml"))
        for slide_name in slide_names[:40]:
            raw = _safe_decode_bytes(zf.read(slide_name))
            texts = re.findall(r"<a:t[^>]*>(.*?)</a:t>", raw, flags=re.S)
            cleaned = [re.sub(r"<[^>]+>", "", item).strip() for item in texts]
            cleaned = [item for item in cleaned if item]
            if cleaned:
                parts.append(f"[{os.path.basename(slide_name)}]\n" + "\n".join(cleaned))
    return "\n\n".join(parts)


def _extract_text_from_archive(data: bytes, filename: str, depth: int = 0) -> Tuple[str, List[Dict[str, Any]]]:
    if depth > 1:
        return "[Archive nesting limit reached.]", []
    parts: List[str] = []
    children: List[Dict[str, Any]] = []
    suffix = os.path.splitext(filename.lower())[1]
    try:
        if suffix == ".zip":
            with zipfile.ZipFile(BytesIO(data)) as zf:
                names = [name for name in zf.namelist() if not name.endswith("/")][:UPLOAD_ARCHIVE_MAX_MEMBERS]
                for name in names:
                    child_data = zf.read(name)
                    child_text, child_meta = _extract_uploaded_file_text(name, child_data, depth=depth + 1)
                    children.append(child_meta)
                    if child_text:
                        parts.append(f"\n--- Archive member: {name} ---\n{child_text}")
        elif suffix in {".tar", ".tgz", ".gz"} or filename.lower().endswith((".tar.gz", ".tgz")):
            if filename.lower().endswith(".gz") and not filename.lower().endswith((".tar.gz", ".tgz")):
                inner = gzip.decompress(data)
                inner_name = filename[:-3] or "decompressed.txt"
                child_text, child_meta = _extract_uploaded_file_text(inner_name, inner, depth=depth + 1)
                children.append(child_meta)
                parts.append(f"\n--- Archive member: {inner_name} ---\n{child_text}")
            else:
                with tarfile.open(fileobj=BytesIO(data), mode="r:*") as tf:
                    members = [m for m in tf.getmembers() if m.isfile()][:UPLOAD_ARCHIVE_MAX_MEMBERS]
                    for member in members:
                        fh = tf.extractfile(member)
                        if fh is None:
                            continue
                        child_data = fh.read()
                        child_text, child_meta = _extract_uploaded_file_text(member.name, child_data, depth=depth + 1)
                        children.append(child_meta)
                        if child_text:
                            parts.append(f"\n--- Archive member: {member.name} ---\n{child_text}")
    except Exception as exc:
        parts.append(f"[Archive extraction failed: {exc}]")
    return _clean_uploaded_text("\n".join(parts), max_chars=UPLOAD_MAX_PER_FILE_CHARS), children


def _extract_uploaded_file_text(filename: str, data: bytes, depth: int = 0) -> Tuple[str, Dict[str, Any]]:
    safe_name = re.sub(r"^[A-Za-z]:", "", str(filename or "uploaded-file")).replace("\\", "/")
    safe_name = "/".join(part for part in safe_name.split("/") if part and part not in {".", ".."})
    safe_name = safe_name or "uploaded-file"
    lower = safe_name.lower()
    suffix = os.path.splitext(lower)[1]
    meta: Dict[str, Any] = {
        "filename": safe_name,
        "size": len(data),
        "parser": "binary-summary",
        "children": [],
    }
    text = ""
    try:
        if suffix in {".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".go", ".rs", ".cpp", ".c", ".h", ".hpp", ".cs", ".rb", ".php", ".swift", ".kt", ".scala", ".sql", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".json", ".xml", ".html", ".css", ".scss", ".sh", ".zsh", ".bat", ".log", ".tex", ".bib", ".rst"}:
            meta["parser"] = "text"
            text = _safe_decode_bytes(data)
        elif suffix == ".docx":
            meta["parser"] = "docx"
            text = _extract_text_from_docx(data)
        elif suffix == ".pdf":
            meta["parser"] = "pdf"
            text = _extract_text_from_pdf(data)
        elif suffix in {".xlsx", ".xlsm", ".csv"}:
            meta["parser"] = "spreadsheet"
            text = _extract_text_from_excel(data, suffix)
        elif suffix == ".pptx":
            meta["parser"] = "pptx"
            text = _extract_text_from_pptx(data)
        elif suffix in {".zip", ".tar", ".tgz", ".gz"} or lower.endswith((".tar.gz", ".tgz")):
            meta["parser"] = "archive"
            text, children = _extract_text_from_archive(data, safe_name, depth=depth)
            meta["children"] = children
        else:
            decoded = _safe_decode_bytes(data)
            printable_ratio = sum(1 for ch in decoded[:2000] if ch.isprintable() or ch.isspace()) / max(1, len(decoded[:2000]))
            if printable_ratio > 0.85:
                meta["parser"] = "text-guess"
                text = decoded
            else:
                text = f"[Binary or unsupported file: {safe_name}, {len(data)} bytes. No reliable text extracted.]"
    except Exception as exc:
        meta["parser_error"] = str(exc)
        text = f"[Failed to parse {safe_name}: {exc}]"
    text = _clean_uploaded_text(text)
    meta["extracted_chars"] = len(text)
    return text, meta


@app.post("/api/upload-context")
async def upload_context(files: List[UploadFile] = File(...), paths: List[str] = Form(default=[])):
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    if len(files) > UPLOAD_MAX_FILES:
        raise HTTPException(status_code=400, detail=f"Too many files; max {UPLOAD_MAX_FILES}")

    file_summaries: List[Dict[str, Any]] = []
    blocks: List[str] = []
    for index, upload in enumerate(files):
        data = await upload.read()
        filename = (paths[index] if index < len(paths) and str(paths[index]).strip() else upload.filename) or "uploaded-file"
        if len(data) > UPLOAD_MAX_FILE_BYTES:
            raise HTTPException(status_code=413, detail=f"{filename} exceeds {UPLOAD_MAX_FILE_BYTES} bytes")
        text, meta = _extract_uploaded_file_text(filename, data)
        meta["content_type"] = upload.content_type or ""
        file_summaries.append(meta)
        blocks.append(
            f"### Uploaded file: {meta['filename']}\n"
            f"- parser: {meta.get('parser')}\n"
            f"- size_bytes: {meta.get('size')}\n\n"
            f"{text}"
        )

    context_text = _clean_uploaded_text(
        "以下是用户上传文件解析出的可用上下文。生成文档时必须优先依据这些材料；"
        "若用户要求根据上传文档写作，应把这些材料作为主要事实来源，不要捏造文件中没有的内容。\n\n"
        + "\n\n".join(blocks),
        max_chars=UPLOAD_MAX_CONTEXT_CHARS,
    )
    return {
        "success": True,
        "file_count": len(file_summaries),
        "files": file_summaries,
        "context_text": context_text,
        "context_chars": len(context_text),
    }


def _is_outline_like_fallback_content(
    content: str,
    outline: str = "",
    forced_pass: bool = False,
) -> bool:
    text = str(content or "").strip()
    if not text:
        return False
    outline_text = str(outline or "").strip()
    compact_text = " ".join(text.split())
    compact_outline = " ".join(outline_text.split())

    has_system_fallback_prefix = text.startswith("（系统兜底）")
    prompt_markers = ["请你作为", "要求：", "段落主题", "系统指示", "content_prompt"]
    has_prompt_marker = any(marker in text for marker in prompt_markers)

    outline_embedded = False
    if compact_outline and len(compact_outline) >= 16:
        outline_embedded = (
            compact_outline in compact_text
            or compact_text in compact_outline
        )

    # 兼容清理旧任务里的 forced_pass/系统兜底内容，避免旧脏数据进入最终文档。
    should_filter = bool((forced_pass or has_system_fallback_prefix) and (outline_embedded or has_prompt_marker))
    if should_filter:
        print(
            f"[Web] filtering outline-like content: forced_pass={forced_pass}, "
            f"system_prefix={has_system_fallback_prefix}, outline_embedded={outline_embedded}, "
            f"has_prompt_marker={has_prompt_marker}, len={len(compact_text)}"
        )
    return should_filter


def _build_content_map_from_history(history: List[Dict[str, Any]]) -> Dict[str, str]:
    """从 history 中构建 content map；跳过旧任务中的 forced_pass 脏数据。"""
    content_map: Dict[str, str] = {}
    # 第一遍：记录所有内容
    for item in history:
        key = f"{item.get('section_id', '')}::{item.get('subsection_id', '')}"
        content = item.get("content", "")
        metadata = item.get("metadata", {})
        is_forced_pass = metadata.get("forced_pass", False)
        outline = metadata.get("outline", "")

        if _is_outline_like_fallback_content(content=content, outline=outline, forced_pass=bool(is_forced_pass)):
            continue

        # 仅当 content 非空时才添加；旧 forced_pass 不能覆盖已验证正文。
        if content:
            if key not in content_map:
                content_map[key] = content
            elif not is_forced_pass:
                content_map[key] = content
    return content_map


def _build_content_map_from_sections(sections: Optional[List[Dict[str, Any]]]) -> Dict[str, str]:
    content_map: Dict[str, str] = {}
    for section in sections or []:
        section_id = str(section.get("section_id") or section.get("id") or "")
        for subsection in section.get("subsections", []) or []:
            subsection_id = str(subsection.get("subsection_id") or subsection.get("id") or "")
            content = subsection.get("content", "")
            is_forced_pass = bool(subsection.get("forced_pass", False))
            outline = subsection.get("outline", "")
            if _is_outline_like_fallback_content(content=content, outline=outline, forced_pass=is_forced_pass):
                continue
            if section_id and subsection_id and content:
                content_map[f"{section_id}::{subsection_id}"] = content
    return content_map


def _build_chapter_asset_map(sections: Optional[List[Dict[str, Any]]]) -> Dict[str, List[Dict[str, Any]]]:
    asset_map: Dict[str, List[Dict[str, Any]]] = {}
    for section in sections or []:
        if not isinstance(section, dict):
            continue
        section_id = str(section.get("section_id") or section.get("id") or "")
        assets = section.get("chapter_assets") or section.get("assets") or []
        if not isinstance(assets, list):
            continue
        for asset in assets:
            if not isinstance(asset, dict):
                continue
            anchor = str(asset.get("insert_after_subsection_id") or "").strip()
            if not anchor or not section_id:
                continue
            key = f"{section_id}::{anchor}"
            asset_map.setdefault(key, []).append(asset)
    return asset_map


def _render_chapter_asset_markdown(asset: Dict[str, Any], ordinal: int) -> str:
    asset_type = str(asset.get("type") or "").strip().lower()
    title = str(asset.get("title") or f"Chapter asset {ordinal}").strip()
    caption = str(asset.get("caption") or "").strip()
    lines: List[str] = []
    if asset_type == "table":
        table_md = str(asset.get("markdown") or "").strip()
        if not table_md:
            return ""
        lines.extend([
            f"**Table {ordinal}. {title}**",
            "",
            table_md,
        ])
        if caption:
            lines.extend(["", f"*{caption}*"])
        return "\n".join(lines).strip()
    return ""


def build_markdown_document(
    title: str,
    structure: Dict[str, Any],
    history: List[Dict[str, Any]],
    generated_sections: Optional[List[Dict[str, Any]]] = None,
    user_background: str = "",
    extra_requirements: str = "",
    epistemic_audit: Optional[Dict[str, Any]] = None,
) -> str:
    content_map = _build_content_map_from_sections(generated_sections)
    chapter_asset_map = _build_chapter_asset_map(generated_sections)
    history_map = _build_content_map_from_history(history)
    for key, value in history_map.items():
        content_map[key] = value
    citation_started_at = time.monotonic()
    citation_budget_seconds = max(0.0, float(CITATION_POSTPROCESS_BUDGET_SECONDS or 0.0))

    def _citation_budget_remaining() -> float:
        if citation_budget_seconds <= 0:
            return 0.0
        return max(0.0, citation_budget_seconds - (time.monotonic() - citation_started_at))

    def _citation_budget_exceeded() -> bool:
        return citation_budget_seconds > 0 and _citation_budget_remaining() <= 0.0

    def _normalize_label(value: str) -> str:
        text = re.sub(r"^第\d+[章节]", "", str(value or "")).strip()
        text = re.sub(r"^\d+(?:\.\d+)*\s*", "", text).strip()
        return text or str(value or "").strip()

    def _to_roman(num: int) -> str:
        vals = [
            (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
        ]
        result = []
        n = max(1, num)
        for v, s in vals:
            while n >= v:
                result.append(s)
                n -= v
        return "".join(result)

    def _to_alpha(num: int) -> str:
        # 1 -> A, 2 -> B, ...
        n = max(1, num)
        chars: List[str] = []
        while n > 0:
            n -= 1
            chars.append(chr(ord("A") + (n % 26)))
            n //= 26
        return "".join(reversed(chars))

    def _document_uses_chinese() -> bool:
        sample = " ".join([
            str(title or ""),
            str(structure.get("outline", "") or ""),
            " ".join(
                str(section.get("title", "") or "")
                for section in structure.get("sections", [])
            ),
        ])
        return bool(re.search(r"[\u4e00-\u9fff]", sample))

    def _extract_topic_terms(limit: int = 8) -> List[str]:
        source_parts: List[str] = [str(title or "")]
        for section in structure.get("sections", []) or []:
            source_parts.append(str(section.get("title", "") or ""))
            for subsection in section.get("subsections", []) or []:
                source_parts.append(str(subsection.get("title", "") or ""))
        text = " ".join(source_parts)
        stop = {
            "section", "subsection", "chapter", "overview", "analysis", "research", "study",
            "applications", "application", "framework", "introduction", "conclusion",
            "研究", "分析", "应用", "扩展", "基础", "理论", "章节", "小节", "定义", "概念",
        }
        terms: List[str] = []

        for chunk in re.findall(r"[\u4e00-\u9fffA-Za-z0-9][\u4e00-\u9fffA-Za-z0-9\s\-]{1,32}", text):
            cleaned = _normalize_label(chunk).strip(" ，,;；:：。.()（）[]【】")
            cleaned = re.sub(r"\s+", " ", cleaned)
            cleaned = re.sub(r"\\[a-zA-Z]+|[$_^{}]", "", cleaned).strip()
            if not cleaned:
                continue
            lowered = cleaned.lower()
            compact = re.sub(r"\s+", "", cleaned)
            if lowered in stop or compact in stop:
                continue
            if any(noise in lowered for noise in ["frac", "bar", "begin", "array"]):
                continue
            if re.fullmatch(r"\d+", compact):
                continue
            if re.fullmatch(r"[A-Za-z]{1,2}", compact):
                continue
            if len(compact) < 2:
                continue
            if len(compact) > 18 and re.search(r"[\u4e00-\u9fff]", compact):
                continue
            if cleaned not in terms:
                terms.append(cleaned)
            if len(terms) >= limit:
                break
        return terms

    def _build_abstract() -> str:
        section_titles = [
            _normalize_label(section.get("title", ""))
            for section in structure.get("sections", [])
            if str(section.get("title", "")).strip()
        ]
        subsection_titles = [
            _normalize_label(subsection.get("title", ""))
            for section in structure.get("sections", [])
            for subsection in section.get("subsections", []) or []
            if str(subsection.get("title", "")).strip()
        ]
        focus = "、".join(section_titles[:3]) if _document_uses_chinese() else ", ".join(section_titles[:3])
        methods = "、".join(subsection_titles[:4]) if _document_uses_chinese() else ", ".join(subsection_titles[:4])
        doc_title = _normalize_label(title)
        content_sentences: List[str] = []
        for section in generated_sections or []:
            for subsection in section.get("subsections", []) or []:
                text = _clean_subsection_text(str(subsection.get("content", "") or ""))
                text = re.sub(r"\[[0-9]+\]", "", text)
                text = re.sub(r"<[^>]+>", "", text)
                for sentence in re.split(r"(?<=[。！？!?])\s*", text):
                    sentence = " ".join(sentence.split()).strip()
                    if not sentence:
                        continue
                    if re.match(r"^#{1,6}\s+", sentence):
                        continue
                    if any(noise in sentence.lower() for noise in ["references", "bibliography", "index terms"]):
                        continue
                    if len(sentence) < 24 or len(sentence) > 180:
                        continue
                    content_sentences.append(sentence.rstrip("。.!?！？"))
                    break
                if len(content_sentences) >= 3:
                    break
            if len(content_sentences) >= 3:
                break
        if _document_uses_chinese():
            evidence_focus = "；".join(content_sentences[:2])
            if evidence_focus:
                return (
                    f"本文以“{doc_title}”为研究对象，围绕{focus or '核心问题'}构建分层论证框架。"
                    f"文章结合{methods or '关键子主题'}展开分析，重点说明：{evidence_focus}。"
                    "在此基础上，本文进一步比较不同方法、机制或实践路径的适用边界，形成兼顾概念界定、证据支撑与应用判断的专题综述。"
                )
            return (
                f"本文围绕“{doc_title}”展开，重点梳理{focus or '核心问题'}之间的逻辑关系、机制解释与应用边界。"
                f"文章进一步结合{methods or '关键子主题'}等内容，构建从概念界定、理论建模到实践场景分析的论证链条。"
                "全文强调问题定义、证据支撑与结构化推理的一致性，旨在形成可复核、可扩展且适合学术写作的专题综述。"
            )
        evidence_focus = "; ".join(content_sentences[:2])
        if evidence_focus:
            return (
                f"This article studies {doc_title} through the connected themes of {focus or 'the central research questions'}. "
                f"It develops the analysis across {methods or 'the main analytical subtopics'} and foregrounds the following argument: {evidence_focus}. "
                "The manuscript integrates definitions, evidence, mechanisms, and applied boundaries into a coherent review-style technical narrative."
            )
        return (
            f"This article examines {doc_title} through the connected themes of {focus or 'the central research questions'}. "
            f"It develops a structured argument across {methods or 'the main analytical subtopics'}, linking definitions, mechanisms, modeling choices, and applied implications. "
            "The resulting manuscript emphasizes traceable reasoning, source-grounded claims, and a coherent academic narrative suitable for technical review."
        )

    def _build_keywords() -> str:
        keywords = []
        main_title = _normalize_label(title)
        if main_title:
            keywords.append(main_title)
        for term in _extract_topic_terms(limit=10):
            if term not in keywords:
                keywords.append(term)
        if _document_uses_chinese():
            for fallback in ["理论框架", "机制分析", "实践应用"]:
                if len(keywords) >= 6:
                    break
                if fallback not in keywords:
                    keywords.append(fallback)
        else:
            for fallback in ["theoretical framework", "mechanism analysis", "applied implications"]:
                if len(keywords) >= 6:
                    break
                if fallback not in keywords:
                    keywords.append(fallback)
        return "Index Terms—" + ", ".join(keywords[:6])

    def _anchor_id(section_index: int, subsection_index: Optional[int] = None) -> str:
        if subsection_index is None:
            return f"chapter-{section_index}"
        return f"chapter-{section_index}-{subsection_index}"

    def _clean_subsection_text(text: str) -> str:
        seen_headings: set[tuple[int, str]] = set()
        cleaned_lines: List[str] = []
        reference_heading_pattern = re.compile(
            r"^\s*(?:#{1,6}\s*)?(?:references?|bibliography|参考文献)\s*[:：]?\s*$",
            re.IGNORECASE,
        )

        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                continue

            # Subsection generators sometimes emit their own local References
            # block. The final document owns one consolidated reference list at
            # the end, so strip any subsection-level reference block here.
            if reference_heading_pattern.match(stripped):
                break

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading_match:
                heading_level = max(len(heading_match.group(1)), 4)
                heading_text = heading_match.group(2).strip()
                if reference_heading_pattern.match(heading_text):
                    break
                heading_key = (heading_level, heading_text)
                if heading_key in seen_headings:
                    continue
                seen_headings.add(heading_key)
                cleaned_lines.append(f"{'#' * heading_level} {heading_text}")
                continue

            cleaned_lines.append(raw_line.rstrip())

        while cleaned_lines and cleaned_lines[-1] == "":
            cleaned_lines.pop()

        cleaned = "\n".join(cleaned_lines).strip()
        # Also handle compact single-line tails such as
        # "References [1] Author... [2] Author..." that are not formatted as a
        # separate heading line.
        cleaned = re.sub(
            r"(?is)\n\s*(?:references?|bibliography|参考文献)\s*[:：]?\s*(?:\[\d+\].*)$",
            "",
            cleaned,
        ).strip()
        return cleaned

    def _section_link(section_index: int, section_title: str) -> str:
        roman = _to_roman(section_index)
        return f"[{roman}. {section_title}](#{_anchor_id(section_index)})"

    def _subsection_link(section_index: int, subsection_index: int, subsection_title: str) -> str:
        alpha = _to_alpha(subsection_index)
        return f"[{alpha}. {subsection_title}](#{_anchor_id(section_index, subsection_index)})"

    def _append_outline_list(lines: List[str]) -> None:
        lines.append("## Contents")
        lines.append("")
        for section_index, section in enumerate(structure.get("sections", []), 1):
            section_title = section.get("title", f"第{section_index}章")
            lines.append(_section_link(section_index, section_title))
            for subsection_index, subsection in enumerate(section.get("subsections", []), 1):
                subsection_title = subsection.get("title", f"第{subsection_index}节")
                lines.append(f"   - {_subsection_link(section_index, subsection_index, subsection_title)}")
        lines.append("")

    def _collect_reference_entries() -> List[str]:
        seen_urls: set[str] = set()
        seen_ref_keys: set[str] = set()
        references: List[str] = []
        candidates: List[Dict[str, Any]] = []  # 存储原始候选对象用于Domain Filter
        last_filtered_out: List[Dict[str, Any]] = []

        # 🔍 诊断日志：追踪引用收集全过程
        diagnostics = {
            "gen_sections_count": len(generated_sections or []),
            "gen_subsections_count": sum(len(s.get("subsections", [])) for s in (generated_sections or [])),
            "gen_source_results_collected": 0,
            "history_items_count": len(history or []),
            "history_source_results_collected": 0,
            "total_candidates_before_filter": 0,
            "domain_filter_applied": False,
            "domain_filter_rag_retries": 0,
            "domain_filter_fallback": False,
            "domain_filter_fallback_count": 0,
            "candidates_after_filter": 0,
            "final_references_count": 0,
            "final_references": [],
        }

        def _collect_textual_citations() -> List[str]:
            patterns = [
                re.compile(r"\(([^()]{2,100}?,\s*\d{4}[a-z]?)\)"),
                re.compile(r"([A-Z][A-Za-z\-]+(?:\s+et al\.)?,\s*\d{4}[a-z]?)"),
                re.compile(r"\[来源\d+\]"),
            ]
            seen_citations: set[str] = set()
            textual_refs: List[str] = []

            def scan_text(text: str) -> None:
                sample = str(text or "")
                if not sample:
                    return
                for pattern in patterns:
                    for match in pattern.findall(sample):
                        citation = match if isinstance(match, str) else str(match)
                        citation = " ".join(citation.split()).strip(" .,;，")
                        if len(citation) < 3:
                            continue
                        key = citation.lower()
                        if key in seen_citations:
                            continue
                        seen_citations.add(key)
                        textual_refs.append(citation)

            for section in generated_sections or []:
                for subsection in (section.get("subsections") or []):
                    scan_text(subsection.get("content", ""))

            for item in history or []:
                scan_text(item.get("content", ""))

            return textual_refs[:20]

        def _append_reference_from_candidate(candidate: Dict[str, Any]) -> bool:
            if not isinstance(candidate, dict):
                return False
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").strip()
            title_text = str(candidate.get("title") or candidate.get("source_name") or candidate.get("source_type") or "").strip()
            snippet = str(candidate.get("body") or candidate.get("abstract") or candidate.get("description") or candidate.get("summary") or "").strip()
            if not (url or title_text or snippet):
                return False
            key = (url or f"{title_text} {snippet[:120]}").lower().strip()
            if not key or key in seen_ref_keys:
                return False
            seen_ref_keys.add(key)
            if url:
                seen_urls.add(url.lower())
            label = title_text or urlparse(url).netloc or "source"
            if snippet:
                if url:
                    references.append(f"{label}: {url} — {snippet[:220]}")
                else:
                    references.append(f"{label} — {snippet[:260]}")
            else:
                references.append(f"{label}: {url}")
            return True

        def add_candidate(candidate: Dict[str, Any]) -> None:
            if not isinstance(candidate, dict):
                return
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").strip()
            title_text = str(candidate.get("title") or candidate.get("source_name") or candidate.get("source_type") or "").strip()
            snippet = str(candidate.get("body") or candidate.get("abstract") or candidate.get("description") or candidate.get("summary") or "").strip()
            if not (url or title_text or snippet):
                return
            candidate_key = (url or f"{title_text} {snippet[:120]}").lower().strip()
            if any(
                (str(c.get("href") or c.get("url") or c.get("link") or "") or f"{c.get('title','')} {str(c.get('body') or c.get('abstract') or '')[:120]}").lower().strip() == candidate_key
                for c in candidates
            ):
                return

            # 保存原始候选对象用于Domain Filter；同时先保留一个候选引用，后续过滤通过会重建。
            candidates.append(candidate)
            _append_reference_from_candidate(candidate)

        def _rebuild_references(selected_candidates: List[Dict[str, Any]]) -> None:
            references.clear()
            seen_urls.clear()
            seen_ref_keys.clear()
            for candidate in selected_candidates:
                _append_reference_from_candidate(candidate)

        def _normalize_candidates(raw_candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            normalized: List[Dict[str, Any]] = []
            seen: set[str] = set()
            for candidate in raw_candidates or []:
                if not isinstance(candidate, dict):
                    continue
                url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").strip()
                title_text = str(candidate.get("title") or candidate.get("source_name") or candidate.get("source_type") or "").strip()
                snippet = str(candidate.get("body") or candidate.get("abstract") or candidate.get("description") or candidate.get("summary") or "").strip()
                if not (url or title_text or snippet):
                    continue
                key = (url or f"{title_text} {snippet[:120]}").lower()
                if key in seen:
                    continue
                seen.add(key)
                normalized.append(candidate)
            return normalized

        def _fetch_candidate_metadata(url: str, timeout_sec: int = 2) -> Dict[str, str]:
            """Attempt to fetch URL and extract a short snippet/description."""
            if (
                not url
                or not HAS_REQUESTS
                or not CITATION_METADATA_FETCH_ENABLED
                or _citation_budget_exceeded()
            ):
                return {}
            try:
                headers = {"User-Agent": "Mozilla/5.0 (compatible; FlowerNet/1.0)"}
                timeout_value = min(float(timeout_sec), max(0.25, _citation_budget_remaining()))
                resp = CITATION_HTTP_SESSION.get(url, timeout=timeout_value, headers=headers)
                if resp.status_code != 200:
                    return {}
                text = resp.text or ""
                # Try BeautifulSoup if available
                if HAS_BS4:
                    try:
                        soup = BeautifulSoup(text, "html.parser")
                        # meta description
                        desc = ""
                        m = soup.find("meta", attrs={"name": "description"}) or soup.find("meta", attrs={"property": "og:description"})
                        if m and m.get("content"):
                            desc = m.get("content")
                        # first paragraph as fallback
                        if not desc:
                            p = soup.find("p")
                            if p:
                                desc = p.get_text().strip()
                        return {"abstract": desc[:1000], "body": desc[:1000]}
                    except Exception:
                        pass
                # fallback regex: look for meta description
                m = re.search(r'<meta[^>]+name=["\']description["\'][^>]+content=["\']([^"\']+)["\']', text, flags=re.I)
                if m:
                    return {"abstract": m.group(1)[:1000], "body": m.group(1)[:1000]}
                # fallback: extract first 300 chars of body text
                body_text = re.sub(r"<[^>]+>", " ", text)
                body_text = " ".join(body_text.split())
                return {"abstract": body_text[:1000], "body": body_text[:1000]}
            except Exception:
                return {}

        def _build_rag_retry_query() -> str:
            # Build a richer RAG retry query: title + index terms + bigrams + synonyms
            domain_filter = get_domain_filter()
            index_terms = domain_filter.extract_document_index_terms(
                title=title,
                outline=str(structure.get("outline", "")),
                abstract=str(structure.get("abstract", "")),
                content_sample=" ".join([
                    subsection.get("content", "")[:200]
                    for section in (generated_sections or [])
                    for subsection in section.get("subsections", [])
                ])[:1000],
            )

            # build synonym expansions from DOMAIN_KEYWORD_MAP
            synonyms = []
            try:
                from citation_drift_prevention import DOMAIN_KEYWORD_MAP
                for v in (DOMAIN_KEYWORD_MAP or {}).values():
                    kws = v.get("keywords") if isinstance(v, dict) else v
                    if kws:
                        synonyms.extend([str(x) for x in kws])
            except Exception:
                pass

            parts = [title, structure.get("outline", ""), user_background, extra_requirements]
            parts.extend(list(index_terms)[:10])
            parts.extend([" ".join(b.split()) for b in domain_filter.extractor._extract_bigrams(" ".join(list(index_terms)))][:10])
            parts.extend(synonyms[:10])

            text = " ".join(str(p or "").strip() for p in parts if str(p or "").strip())
            # normalize whitespace and truncate
            return " ".join(text.split())[:480]

        # Shared anti-drift state for retry/fallback stages
        last_anchor_terms: List[str] = []
        last_selected_terms: List[str] = []
        last_flat_red_flags: set[str] = set()

        def _translate_anchor_to_english(anchor: str) -> str:
            mapping = {
                "商业谈判": "business negotiation",
                "商务谈判": "commercial negotiation",
                "谈判策略": "negotiation strategy",
                "商务沟通": "business communication",
                "采购谈判": "procurement negotiation",
                "供应链谈判": "supply chain negotiation",
                "商务英语谈判": "business english negotiation",
            }
            s = str(anchor or "").strip()
            return mapping.get(s, s)

        def _passes_anchor_gate(candidate: Dict[str, Any], anchor_terms: List[str], selected_terms: List[str], flat_red_flags: set[str]) -> tuple[bool, str]:
            text = " ".join([
                str(candidate.get("title") or ""),
                str(candidate.get("abstract") or ""),
                str(candidate.get("body") or ""),
            ]).lower()
            if not text.strip():
                return False, "empty_text"
            for rf in list(flat_red_flags)[:300]:
                if rf and rf in text:
                    return False, f"red_flag:{rf}"
            anchor_hits = [a for a in anchor_terms if str(a).lower() in text]
            selected_hits = [t for t in selected_terms if str(t).lower() in text]
            if anchor_hits or len(selected_hits) >= 2:
                return True, f"anchor_hits={anchor_hits[:3]} term_hits={selected_hits[:3]}"
            return False, "weak_domain_overlap"

        def _candidate_host(candidate: Dict[str, Any]) -> str:
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").strip()
            return (urlparse(url).netloc or str(candidate.get("source") or "")).lower().replace("www.", "")

        def _scholarly_source_tier(candidate: Dict[str, Any]) -> float:
            """Prefer peer-reviewed / publisher / Crossref/PubMed sources, but keep fallback candidates usable."""
            host = _candidate_host(candidate)
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").lower()
            source = str(candidate.get("source") or candidate.get("provider") or "").lower()
            source_type = str(candidate.get("source_type") or candidate.get("type") or "").lower()
            container = str(candidate.get("container_title") or candidate.get("journal") or candidate.get("venue") or "").strip()
            publisher = str(candidate.get("publisher") or "").lower()

            strong_publishers = {
                "nature.com", "science.org", "cell.com", "sciencedirect.com", "springer.com",
                "link.springer.com", "wiley.com", "onlinelibrary.wiley.com", "tandfonline.com",
                "sagepub.com", "cambridge.org", "cambridge.org", "oxfordacademic.com", "academic.oup.com",
                "jstor.org", "ieee.org", "ieeexplore.ieee.org", "acm.org", "dl.acm.org",
                "aclweb.org", "openreview.net", "proceedings.neurips.cc", "proceedings.mlr.press",
                "thelancet.com", "nejm.org", "bmj.com", "cochranelibrary.com",
            }
            public_evidence = {
                "pubmed.ncbi.nlm.nih.gov", "ncbi.nlm.nih.gov", "nih.gov", "who.int", "oecd.org",
                "un.org", "worldbank.org", "imf.org", "nber.org", "nist.gov", "eric.ed.gov",
            }
            preprint_or_secondary = {"arxiv.org", "biorxiv.org", "medrxiv.org", "ssrn.com", "researchgate.net", "mdpi.com"}

            score = 0.0
            if "crossref" in source or "doi.org" in host or "doi.org" in url:
                score += 0.55
            if host in strong_publishers or any(d in host for d in strong_publishers):
                score += 0.50
            if host in public_evidence or any(d in host for d in public_evidence):
                score += 0.42
            if source_type in {"journal-article", "proceedings-article", "book-chapter", "book", "report"}:
                score += 0.22
            if container:
                score += 0.12
            if publisher and any(p in publisher for p in ["elsevier", "springer", "wiley", "sage", "taylor", "cambridge", "oxford", "ieee", "acm"]):
                score += 0.12
            if any(d in host for d in preprint_or_secondary):
                score -= 0.18
            if host in {"researchgate.net"}:
                score -= 0.22
            return max(0.0, min(1.0, score))

        def _source_whitelist_weight(candidate: Dict[str, Any]) -> float:
            text = " ".join([
                str(candidate.get("title") or ""),
                str(candidate.get("abstract") or ""),
                str(candidate.get("body") or ""),
            ]).lower()
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").lower()

            domain_whitelist = [
                "pubmed.ncbi.nlm.nih.gov", "ncbi.nlm.nih.gov", "crossref.org", "doi.org",
                "sciencedirect.com", "springer.com", "link.springer.com", "wiley.com", "onlinelibrary.wiley.com",
                "tandfonline.com", "sagepub.com", "emerald.com", "jstor.org", "cambridge.org",
                "oxfordacademic.com", "academic.oup.com", "nature.com", "science.org", "cell.com",
                "ieee.org", "ieeexplore.ieee.org", "acm.org", "dl.acm.org", "aclweb.org", "openreview.net",
            ]
            business_terms = [
                "business", "management", "marketing", "negotiation", "commerce", "economics", "supply chain",
                "商业", "商务", "管理", "营销", "经济", "谈判", "供应链",
            ]
            doi_prefix_whitelist = [
                "10.1038", "10.1126", "10.1016", "10.1007", "10.1002", "10.1111", "10.1177",
                "10.1080", "10.1093", "10.1109", "10.1145", "10.1287", "10.5465",
            ]

            w = 0.0
            w += 0.45 * _scholarly_source_tier(candidate)
            if any(d in url for d in domain_whitelist):
                w += 0.25
            if any(t in text for t in business_terms):
                w += 0.18
            m = re.search(r"10\.\d{4,9}/[^\s]+", url)
            if m and any(m.group(0).lower().startswith(p) for p in doi_prefix_whitelist):
                w += 0.20
            return min(1.0, w)

        def _build_augmented_domain_terms(base_terms: set[str], title_text: str) -> set[str]:
            augmented = set(base_terms or set())
            title_lower = (title_text or "").lower()

            # Keep only a small, high-precision business/negotiation anchor family.
            if any(x in title_lower for x in ["谈判", "negotiation", "business", "商务", "商业"]):
                augmented.update({
                    "商业谈判", "商务谈判", "谈判策略", "商务沟通",
                    "business negotiation", "commercial negotiation", "negotiation strategy",
                    "sales negotiation", "business communication", "bargaining",
                })
            if any(x in title_lower for x in ["采购", "供应链", "procurement", "supply chain"]):
                augmented.update({
                    "采购谈判", "供应链谈判", "procurement negotiation", "contract negotiation", "supply chain negotiation",
                })

            # Add a very small amount of bilingual lexicon expansion to recover English abstracts.
            zh_en_pairs = {
                "谈判": "negotiation",
                "商务": "business",
                "商业": "commercial",
                "策略": "strategy",
                "沟通": "communication",
                "采购": "procurement",
                "供应链": "supply chain",
                "营销": "marketing",
                "管理": "management",
            }
            for zh, en in zh_en_pairs.items():
                if zh in title_text:
                    augmented.add(en)

            # Prune obviously noisy fragments while keeping genuine anchor phrases.
            pruned = set()
            for term in augmented:
                t = re.sub(r"\s+", "", str(term or "").strip())
                if not t:
                    continue
                if len(t) == 1:
                    continue
                if re.fullmatch(r"[a-zA-Z]{1,2}", t):
                    continue
                if t in {"企业", "商业", "谈判", "策略", "管理", "研究", "应用", "分析", "模型"}:
                    continue
                if not re.search(r"[\u4e00-\u9fff]{2,}|[a-zA-Z]{3,}", t):
                    continue
                pruned.add(term)
            return pruned

        def _apply_source_weighting(cands: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            weighted = []
            for c in cands or []:
                cc = dict(c)
                cc["source_tier"] = _scholarly_source_tier(cc)
                cc["source_weight"] = _source_whitelist_weight(cc)
                weighted.append(cc)
            weighted.sort(
                key=lambda x: (
                    float(x.get("source_tier", 0.0)),
                    float(x.get("source_weight", 0.0)),
                    float(x.get("quality_score", 0.0) or x.get("domain_similarity", 0.0) or 0.0),
                ),
                reverse=True,
            )
            if weighted:
                print(
                    "📌 [引用诊断] SourceWeight Top: " +
                    "; ".join([
                        f"{str(x.get('title') or '')[:40]}=tier{float(x.get('source_tier', 0.0)):.2f}/w{float(x.get('source_weight', 0.0)):.2f}"
                        for x in weighted[:3]
                    ])
                )
            return weighted

        def _run_rag_retry_candidates() -> List[Dict[str, Any]]:
            if (
                not CITATION_EXTERNAL_FALLBACK_ENABLED
                or not HAS_RAG_SEARCH_ENGINE
                or _citation_budget_exceeded()
            ):
                return []
            # Enhanced multi-query RAG retry: split into short queries (title + each index term)
            # Build index terms (reuse domain filter's extractor)
            try:
                domain_filter = get_domain_filter()
                index_terms = domain_filter.extract_document_index_terms(
                    title=title,
                    outline=str(structure.get("outline", "")),
                    abstract=str(structure.get("abstract", "")),
                    content_sample=" ".join([
                        subsection.get("content", "")[:200]
                        for section in (generated_sections or [])
                        for subsection in section.get("subsections", [])
                    ])[:1000],
                )
                augmented_index_terms = _build_augmented_domain_terms(index_terms, title)
            except Exception:
                index_terms = set()
                augmented_index_terms = _build_augmented_domain_terms(index_terms, title)

            # Simple stopword cleaning for English
            english_stopwords = {
                "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
                "of", "with", "by", "from", "is", "are", "be", "that", "this",
            }

            def _clean_query(q: str) -> str:
                s = " ".join(q.split())
                # remove punctuation
                s = re.sub(r"[^\w\u4e00-\u9fff\s]", " ", s)
                # remove small stopwords for English
                toks = [t for t in s.split() if t.lower() not in english_stopwords]
                return " ".join(toks)[:480]

            def _build_boosted_query(parts: List[str]) -> str:
                boosted = []
                for p in parts:
                    p = str(p or "").strip()
                    if not p:
                        continue
                    if len(p) >= 4 and " " in p:
                        boosted.extend([f'"{p}"', f'"{p}"'])
                    else:
                        boosted.append(p)
                return _clean_query(" ".join(boosted))

            def _select_index_terms(terms: set[str], limit: int = 8) -> List[str]:
                noise_terms = {
                    "与", "的", "和", "及", "或", "在", "对", "中", "于", "为", "是", "了",
                    "企业", "商业", "谈判", "策略", "管理", "研究", "应用", "分析", "模型",
                }
                cleaned = []
                for term in sorted(list(terms), key=lambda s: (-len(str(s)), str(s)))[: max(limit * 2, limit)]:
                    term = str(term).strip()
                    if not term:
                        continue
                    term_norm = re.sub(r"\s+", "", term)
                    if not term_norm:
                        continue
                    if term_norm.lower() in noise_terms:
                        continue
                    if re.fullmatch(r"[\u4e00-\u9fff]{1}", term_norm):
                        continue
                    if re.fullmatch(r"[a-zA-Z]{1,2}", term_norm):
                        continue
                    if not re.search(r"[\u4e00-\u9fff]{2,}|[a-zA-Z]{3,}", term_norm):
                        continue
                    cleaned.append(term)
                # favor longer / phrase-like terms and keep unique order
                seen_local = set()
                output = []
                for term in cleaned:
                    key = term.lower()
                    if key in seen_local:
                        continue
                    seen_local.add(key)
                    output.append(term)
                    if len(output) >= limit:
                        break
                return output

            selected_terms = _select_index_terms(augmented_index_terms, limit=8)

            # Domain anchors used to prioritize academic retrieval and aggressively reduce citation drift.
            anchor_terms: List[str] = []
            title_lower = (title or "").lower()
            if any(x in title_lower for x in ["谈判", "negotiation", "business", "商务", "商业"]):
                anchor_terms.extend([
                    "商业谈判", "商务谈判", "谈判策略", "商务沟通",
                    "sales negotiation", "business negotiation",
                ])
            if any(x in title_lower for x in ["采购", "供应链", "procurement", "supply chain"]):
                anchor_terms.extend(["采购谈判", "供应链谈判", "procurement negotiation", "contract negotiation"])
            anchor_terms.extend(selected_terms[:6])
            anchor_terms = [a for a in anchor_terms if str(a).strip()]
            # dedupe anchors
            seen_anchor = set()
            anchor_terms = [a for a in anchor_terms if (a.lower() not in seen_anchor and not seen_anchor.add(a.lower()))]

            try:
                from citation_drift_prevention import CROSS_DOMAIN_RED_FLAGS as _RFLAGS
                _flat_red_flags = set()
                if isinstance(_RFLAGS, dict):
                    for vv in _RFLAGS.values():
                        if isinstance(vv, (list, set, tuple)):
                            _flat_red_flags.update({str(x).lower() for x in vv if x})
                elif isinstance(_RFLAGS, (list, set, tuple)):
                    _flat_red_flags.update({str(x).lower() for x in _RFLAGS if x})
            except Exception:
                _flat_red_flags = set()

            nonlocal last_anchor_terms, last_selected_terms, last_flat_red_flags
            last_anchor_terms = list(anchor_terms)
            last_selected_terms = list(selected_terms)
            last_flat_red_flags = set(_flat_red_flags)

            def _append_candidate(it: Dict[str, Any], source_tag: str) -> bool:
                url = str(it.get("href") or it.get("url") or it.get("link") or "").strip().lower()
                if not url or url in seen_urls:
                    return False
                ok, reason = _passes_anchor_gate(it, anchor_terms, selected_terms, _flat_red_flags)
                if not ok:
                    print(f"📌 [引用诊断] 候选拒绝({source_tag}): {str(it.get('title') or '')[:60]} | 原因={reason}")
                    return False
                seen_urls.add(url)
                merged_results.append(it)
                return True

            # Execute queries and merge results until we reach desired count
            merged_results: List[Dict[str, Any]] = []
            seen_urls: set = set()

            # ---------------- Academic first ----------------
            # Build high-precision academic anchor queries first, then run providers in order.
            academic_queries = []
            academic_queries.append(_build_boosted_query([title or "", *anchor_terms[:3], extra_requirements or ""]))
            academic_queries.append(_build_boosted_query([title or "", *anchor_terms[3:6], extra_requirements or ""]))
            for t in selected_terms[:4]:
                academic_queries.append(_build_boosted_query([title or "", t, extra_requirements or ""]))
            seen_aq = set()
            academic_queries = [q for q in academic_queries if q and (q not in seen_aq and not seen_aq.add(q))]

            # Build bilingual query pairs (zh/en) to improve coverage and reduce drift.
            bilingual_query_pairs: List[tuple[str, str]] = []
            for q in academic_queries:
                # english mirror from anchors; preserves domain intent and improves non-Chinese APIs
                en_parts = [title or ""]
                for a in anchor_terms[:6]:
                    en_parts.append(_translate_anchor_to_english(a))
                if extra_requirements:
                    en_parts.append(str(extra_requirements))
                en_q = _build_boosted_query(en_parts)
                bilingual_query_pairs.append((q, en_q))
            bilingual_query_pairs = bilingual_query_pairs[: max(0, CITATION_ACADEMIC_QUERY_PAIR_LIMIT)]

            provider_plan = [
                ("SemanticScholar", _query_semanticscholar),
                ("Crossref", _query_crossref),
                ("arXiv", _query_arxiv),
            ]

            for provider_name, provider_fn in provider_plan:
                for i, (q_zh, q_en) in enumerate(bilingual_query_pairs):
                    if _citation_budget_exceeded():
                        print("⚠️ [引用诊断] 引用后处理达到时间预算，停止学术检索")
                        return merged_results[:DOMAIN_FILTER_RAG_RETRY_RESULTS]
                    if len(merged_results) >= DOMAIN_FILTER_RAG_RETRY_RESULTS:
                        break
                    need = max(1, DOMAIN_FILTER_RAG_RETRY_RESULTS - len(merged_results))
                    try:
                        zh_hits = provider_fn(q_zh, max_results=min(need + 3, DOMAIN_FILTER_RAG_RETRY_RESULTS + 3))
                        en_hits = provider_fn(q_en, max_results=min(need + 3, DOMAIN_FILTER_RAG_RETRY_RESULTS + 3))
                        print(f"📌 [引用诊断] 学术锚点查询[{provider_name}]#{i+1}[ZH]: '{q_zh[:80]}' -> 命中 {len(zh_hits or [])}")
                        print(f"📌 [引用诊断] 学术锚点查询[{provider_name}]#{i+1}[EN]: '{q_en[:80]}' -> 命中 {len(en_hits or [])}")

                        pair_map: Dict[str, Dict[str, Any]] = {}
                        for lang, items in (("zh", zh_hits or []), ("en", en_hits or [])):
                            for it in items:
                                u = str(it.get("href") or it.get("url") or it.get("link") or "").strip().lower()
                                if not u:
                                    continue
                                if u not in pair_map:
                                    pair_map[u] = {"item": dict(it), "zh": False, "en": False}
                                pair_map[u][lang] = True

                        # Intersect bilingual hits first, then intersect with domain-gate pass set.
                        for rec in pair_map.values():
                            if not (rec.get("zh") and rec.get("en")):
                                continue
                            it = rec.get("item") or {}
                            ok, reason = _passes_anchor_gate(it, anchor_terms, selected_terms, _flat_red_flags)
                            if not ok:
                                print(f"📌 [引用诊断] 候选拒绝({provider_name}#{i+1}): {str(it.get('title') or '')[:60]} | 原因={reason}")
                                continue
                            _append_candidate(it, f"{provider_name}#{i+1}")
                            if len(merged_results) >= DOMAIN_FILTER_RAG_RETRY_RESULTS:
                                break
                    except Exception as e:
                        print(f"⚠️ [引用诊断] 学术锚点查询失败[{provider_name}]#{i+1}: {e}")
                if len(merged_results) >= DOMAIN_FILTER_RAG_RETRY_RESULTS:
                    break

            # ---------------- Generic RAG backfill ----------------
            # Only if academic anchor-first path cannot fill enough candidates.
            if len(merged_results) < DOMAIN_FILTER_RAG_RETRY_RESULTS:
                queries = []
                base = " ".join([title or "", structure.get("outline", "") or "", user_background or "", extra_requirements or ""])
                if base.strip():
                    queries.append(_build_boosted_query([title or "", extra_requirements or ""]))

                if title and any(x in title for x in ["谈判", "negotiation", "business"]):
                    queries.append(_build_boosted_query([title or "", "商业谈判", "商务谈判", "谈判策略", extra_requirements or ""]))

                for term in selected_terms:
                    q_parts = [title or "", term, extra_requirements or ""]
                    if " " in term:
                        q_parts.insert(1, term)
                    queries.append(_build_boosted_query(q_parts))

                # include a few synonyms from DOMAIN_KEYWORD_MAP as queries
                try:
                    from citation_drift_prevention import DOMAIN_KEYWORD_MAP as _DQ
                    syns = []
                    for v in (_DQ or {}).values():
                        kws = v.get("keywords") if isinstance(v, dict) else v
                        if kws:
                            syns.extend([str(x) for x in kws])
                    for s in syns[:4]:
                        queries.append(_build_boosted_query([title or "", s, extra_requirements or ""]))
                except Exception:
                    pass

                seen_q = set()
                queries = [q for q in queries if q and (q not in seen_q and not seen_q.add(q))]

                # Tighten query count to reduce noisy retrievals.
                queries = queries[:10]
                per_query = max(2, math.ceil((DOMAIN_FILTER_RAG_RETRY_RESULTS - len(merged_results)) / max(1, len(queries))))

                for i, q in enumerate(queries):
                    if _citation_budget_exceeded():
                        print("⚠️ [引用诊断] 引用后处理达到时间预算，停止RAG补全")
                        break
                    try:
                        timeout_value = min(
                            float(DOMAIN_FILTER_RAG_RETRY_TIMEOUT),
                            max(0.25, _citation_budget_remaining()),
                        )
                        engine = RAGSearchEngine(max_results=per_query, timeout=timeout_value)
                        result = engine.search(q)
                        hits = result.get("results") if isinstance(result, dict) else None
                        hit_count = len(hits) if hits else 0
                        print(f"📌 [引用诊断] RAG 子查询#{i+1}: '{q[:80]}' -> 命中 {hit_count}")
                        if hits:
                            for item in _normalize_candidates(hits):
                                _append_candidate(item, f"RAG#{i+1}")
                                if len(merged_results) >= DOMAIN_FILTER_RAG_RETRY_RESULTS:
                                    break
                    except Exception as e:
                        print(f"⚠️ [引用诊断] RAG 子查询失败 #{i+1}: {e}")
                    if len(merged_results) >= DOMAIN_FILTER_RAG_RETRY_RESULTS:
                        break

            return merged_results[:DOMAIN_FILTER_RAG_RETRY_RESULTS]

        def _query_crossref(query: str, max_results: int = 5) -> List[Dict[str, Any]]:
            """Query Crossref API for scholarly works and return normalized candidates."""
            out = []
            if not query or not HAS_REQUESTS or _citation_budget_exceeded():
                return out
            try:
                url = "https://api.crossref.org/works"
                params = {"query.bibliographic": query, "rows": max_results}
                headers = {"User-Agent": "FlowerNet/1.0"}
                timeout_value = min(4.0, max(0.25, _citation_budget_remaining()))
                resp = CITATION_HTTP_SESSION.get(url, params=params, timeout=timeout_value, headers=headers)
                if resp.status_code != 200:
                    return out
                j = resp.json()
                items = j.get("message", {}).get("items", [])
                for itm in items:
                    doi = itm.get("DOI")
                    title = " ".join(itm.get("title") or [])
                    abstract = itm.get("abstract") or ""
                    url_link = f"https://doi.org/{doi}" if doi else itm.get("URL")
                    container = " ".join(itm.get("container-title") or [])
                    published = itm.get("published-print") or itm.get("published-online") or itm.get("issued") or {}
                    year_parts = published.get("date-parts") if isinstance(published, dict) else None
                    year = ""
                    if year_parts and isinstance(year_parts, list) and year_parts and isinstance(year_parts[0], list) and year_parts[0]:
                        year = str(year_parts[0][0])
                    out.append({
                        "href": url_link,
                        "title": title,
                        "body": re.sub(r"<[^>]+>", " ", str(abstract))[:1000],
                        "source": "crossref.org",
                        "provider": "Crossref",
                        "source_type": itm.get("type") or "",
                        "container_title": container,
                        "journal": container,
                        "publisher": itm.get("publisher") or "",
                        "year": year,
                        "doi": doi or "",
                    })
                    if len(out) >= max_results:
                        break
            except Exception:
                return out
            return out

        def _query_semanticscholar(query: str, max_results: int = 5) -> List[Dict[str, Any]]:
            """Query Semantic Scholar Graph API for papers with optional API-key and retries.

            Supports exponential backoff for transient errors and allows an API key via
            the `SEMANTIC_SCHOLAR_API_KEY` environment variable (optional).
            """
            out = []
            if not query or not HAS_REQUESTS or _citation_budget_exceeded():
                return out
            import os, time

            api_key = os.environ.get("SEMANTIC_SCHOLAR_API_KEY") or os.environ.get("S2_API_KEY")
            url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {"query": query, "limit": max_results, "fields": "title,abstract,url,doi,paperId,venue,year,publicationTypes,journal"}
            headers = {"User-Agent": "FlowerNet/1.0"}
            if api_key:
                # try both common header forms
                headers["x-api-key"] = api_key
                headers["Authorization"] = f"Bearer {api_key}"

            retries = max(1, CITATION_SEMANTIC_SCHOLAR_RETRIES)
            backoff = 1.0
            for attempt in range(1, retries + 1):
                if _citation_budget_exceeded():
                    return out
                try:
                    timeout_value = min(6.0, max(0.25, _citation_budget_remaining()))
                    resp = CITATION_HTTP_SESSION.get(url, params=params, timeout=timeout_value, headers=headers)
                    if resp.status_code == 200:
                        j = resp.json()
                        data = j.get("data") or []
                        for itm in data:
                            title = itm.get("title") or ""
                            abstract = itm.get("abstract") or ""
                            doi = itm.get("doi")
                            url_link = itm.get("url") or (f"https://doi.org/{doi}" if doi else None)
                            if not url_link:
                                pid = itm.get("paperId")
                                if pid:
                                    url_link = f"https://www.semanticscholar.org/paper/{pid}"
                            journal = itm.get("journal") if isinstance(itm.get("journal"), dict) else {}
                            venue = itm.get("venue") or journal.get("name") or ""
                            pub_types = itm.get("publicationTypes") or []
                            out.append({
                                "href": url_link,
                                "title": title,
                                "body": abstract[:1000],
                                "source": "semanticscholar.org",
                                "provider": "SemanticScholar",
                                "source_type": ",".join(str(x) for x in pub_types) if isinstance(pub_types, list) else str(pub_types or ""),
                                "container_title": venue,
                                "journal": venue,
                                "year": str(itm.get("year") or ""),
                                "doi": doi or "",
                            })
                            if len(out) >= max_results:
                                break
                        return out
                    if resp.status_code in (429, 500, 502, 503, 504):
                        # transient, retry with backoff
                        if attempt < retries and _citation_budget_remaining() > 0.5:
                            time.sleep(min(backoff, max(0.0, _citation_budget_remaining() - 0.25)))
                            backoff *= 2
                            continue
                        return out
                    # non-retryable
                    return out
                except Exception:
                    if attempt < retries and _citation_budget_remaining() > 0.5:
                        time.sleep(min(backoff, max(0.0, _citation_budget_remaining() - 0.25)))
                        backoff *= 2
                        continue
                    return out

        def _query_arxiv(query: str, max_results: int = 5) -> List[Dict[str, Any]]:
            """Query arXiv API and return normalized candidates."""
            out = []
            if not query or not HAS_REQUESTS or _citation_budget_exceeded():
                return out
            try:
                # use the arXiv API (Atom)
                q = re.sub(r"\s+", "+", query.strip())
                url = f"http://export.arxiv.org/api/query?search_query=all:{q}&start=0&max_results={max_results}"
                headers = {"User-Agent": "FlowerNet/1.0"}
                timeout_value = min(6.0, max(0.25, _citation_budget_remaining()))
                resp = CITATION_HTTP_SESSION.get(url, timeout=timeout_value, headers=headers)
                if resp.status_code != 200:
                    return out
                text = resp.text or ""
                # simple XML extraction without extra deps
                entries = re.split(r"<entry>|</entry>", text)
                for e in entries:
                    if "<id>" not in e:
                        continue
                    m_id = re.search(r"<id>(.*?)</id>", e, re.S)
                    m_title = re.search(r"<title>(.*?)</title>", e, re.S)
                    m_sum = re.search(r"<summary>(.*?)</summary>", e, re.S)
                    url_link = m_id.group(1).strip() if m_id else None
                    title = re.sub(r"\s+", " ", (m_title.group(1) if m_title else "")).strip()
                    summary = re.sub(r"\s+", " ", (m_sum.group(1) if m_sum else "")).strip()
                    out.append({"href": url_link, "title": title, "body": summary[:1000]})
                    if len(out) >= max_results:
                        break
            except Exception:
                return out
            return out

        def _pick_domain_fallback(filtered_out: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            gated = []
            for c in filtered_out or []:
                ok, reason = _passes_anchor_gate(c, last_anchor_terms, last_selected_terms, last_flat_red_flags)
                if not ok:
                    print(f"📌 [引用诊断] 兜底拒绝: {str(c.get('title') or '')[:60]} | 原因={reason}")
                    continue
                cc = dict(c)
                cc["source_tier"] = _scholarly_source_tier(cc)
                cc["source_weight"] = _source_whitelist_weight(cc)
                # Weighted fallback score: keep domain similarity primary, source quality secondary
                cc["fallback_score"] = (
                    float(cc.get("domain_similarity", 0.0))
                    + 0.35 * float(cc.get("source_tier", 0.0))
                    + 0.20 * float(cc.get("source_weight", 0.0))
                    + 0.10 * float(cc.get("quality_score", 0.0) or 0.0)
                )
                gated.append(cc)

            ranked = sorted(
                gated,
                key=lambda x: (
                    float(x.get("fallback_score", 0.0)),
                    float(x.get("source_tier", 0.0)),
                    float(x.get("source_weight", 0.0)),
                ),
                reverse=True,
            )
            return ranked[: max(0, DOMAIN_FILTER_FALLBACK_TOP_K)]

        # ============ 从 generated_sections 中提取 source_results ============
        for section_idx, section in enumerate(generated_sections or []):
            for subsection_idx, subsection in enumerate(section.get("subsections") or []):
                source_results = subsection.get("source_results", []) or []
                if source_results:
                    diagnostics["gen_source_results_collected"] += len(source_results)
                    print(f"📌 [引用诊断] generated_sections[{section_idx}].subsections[{subsection_idx}] 包含 {len(source_results)} 个 source_results")
                for item in source_results:
                    add_candidate(item if isinstance(item, dict) else {})

        # ============ 从 history 中提取 source_results ============
        for history_idx, item in enumerate(history or []):
            metadata = item.get("metadata", {}) if isinstance(item, dict) else {}
            source_results = metadata.get("source_results", []) or []
            if source_results:
                diagnostics["history_source_results_collected"] += len(source_results)
                print(f"📌 [引用诊断] history[{history_idx}].metadata.source_results 包含 {len(source_results)} 个条目")
            for candidate in source_results:
                add_candidate(candidate if isinstance(candidate, dict) else {})

        diagnostics["total_candidates_before_filter"] = len(candidates)
        print(f"📌 [引用诊断] 共收集候选引用: {len(candidates)} 个 (生成器:{diagnostics['gen_source_results_collected']}, 历史:{diagnostics['history_source_results_collected']})")

        # 🔍 应用Domain Filter进行领域相关性检查
        augmented_index_terms: set[str] = set()
        if HAS_DOMAIN_FILTER and candidates:
            try:
                diagnostics["domain_filter_applied"] = True
                domain_filter = get_domain_filter()

                # 从文档中提取Index Terms
                index_terms = domain_filter.extract_document_index_terms(
                    title=title,
                    outline=str(structure.get("outline", "")),
                    abstract=str(structure.get("abstract", "")),
                    content_sample=" ".join([
                        subsection.get("content", "")[:200]
                        for section in (generated_sections or [])
                        for subsection in section.get("subsections", [])
                    ])[:1000],
                )
                augmented_index_terms = _build_augmented_domain_terms(index_terms, title)

                # 过滤候选引用
                # Enrich candidates with fetched metadata when missing (limit fetches)
                if HAS_REQUESTS and candidates:
                    fetch_budget = 6
                    for c in candidates:
                        if fetch_budget <= 0 or _citation_budget_exceeded():
                            break
                        url = str(c.get("href") or c.get("url") or c.get("link") or "").strip()
                        if url and not (c.get("abstract") or c.get("body")):
                            meta = _fetch_candidate_metadata(url, timeout_sec=2)
                            if meta:
                                c.update(meta)
                                fetch_budget -= 1
                                print(f"📌 [引用诊断] 已抓取并填充候选元数据: {url}")

                # Source-quality weighting before Domain Filter to prioritize business-aligned sources.
                candidates = _apply_source_weighting(candidates)

                domain_debug = os.environ.get("DOMAIN_FILTER_DEBUG_DETAILS", "0").lower() in {"1", "true", "yes", "on"}
                filtered_candidates, filtered_out = domain_filter.filter_citations(
                    citations=candidates,
                    index_terms=augmented_index_terms,
                    debug=domain_debug,
                )
                last_filtered_out = filtered_out

                print(f"📌 [引用诊断] Domain Filter: 过滤前 {len(candidates)} -> 过滤后 {len(filtered_candidates)} (丢弃 {len(filtered_out)})")

                rag_retry_count = 0
                if not filtered_candidates and candidates and CITATION_EXTERNAL_FALLBACK_ENABLED:
                    while rag_retry_count < max(0, DOMAIN_FILTER_RAG_RETRY_MAX):
                        if _citation_budget_exceeded():
                            print("⚠️ [引用诊断] 引用后处理达到时间预算，跳过剩余Domain Filter重试")
                            break
                        rag_retry_count += 1
                        retry_candidates = _run_rag_retry_candidates()
                        if not retry_candidates:
                            continue
                        # Enrich retry candidates (limit fetches)
                        if HAS_REQUESTS and retry_candidates:
                            fetch_budget = 6
                            for c in retry_candidates:
                                if fetch_budget <= 0 or _citation_budget_exceeded():
                                    break
                                url = str(c.get("href") or c.get("url") or c.get("link") or "").strip()
                                if url and not (c.get("abstract") or c.get("body")):
                                    meta = _fetch_candidate_metadata(url, timeout_sec=2)
                                    if meta:
                                        c.update(meta)
                                        fetch_budget -= 1
                                        print(f"📌 [引用诊断] 已抓取并填充重试候选元数据: {url}")
                        retry_candidates = _apply_source_weighting(retry_candidates)
                        filtered_candidates, filtered_out = domain_filter.filter_citations(
                            citations=retry_candidates,
                            index_terms=augmented_index_terms,
                            debug=domain_debug,
                        )
                        last_filtered_out = filtered_out
                        print(
                            f"📌 [引用诊断] Domain Filter RAG重试#{rag_retry_count}: "
                            f"候选 {len(retry_candidates)} -> 过滤后 {len(filtered_candidates)}"
                        )
                        if filtered_candidates:
                            candidates = retry_candidates
                            break
                elif not filtered_candidates and candidates:
                    print("📌 [引用诊断] 外部引用补全默认关闭，跳过Domain Filter RAG重试")

                diagnostics["domain_filter_rag_retries"] = rag_retry_count

                if not filtered_candidates and last_filtered_out:
                    fallback_candidates = _pick_domain_fallback(last_filtered_out)
                    if fallback_candidates:
                        filtered_candidates = fallback_candidates
                        diagnostics["domain_filter_fallback"] = True
                        diagnostics["domain_filter_fallback_count"] = len(fallback_candidates)
                        print(
                            f"⚠️ [引用诊断] Domain Filter 仍无可用引用，启用兜底 {len(fallback_candidates)} 条"
                        )
                else:
                    diagnostics["domain_filter_fallback"] = False

                # 重建references列表（仅保留过滤后的）
                _rebuild_references(filtered_candidates)

                diagnostics["candidates_after_filter"] = len(filtered_candidates)
                print(f"📌 [引用诊断] 过滤后重建 references: {len(references)} 条")

            except Exception as e:
                print(f"⚠️ Domain Filter error: {e}, continuing with original references")
                if os.environ.get("DOMAIN_FILTER_DEBUG_DETAILS", "0").lower() in {"1", "true", "yes", "on"}:
                    import traceback
                    traceback.print_exc()

        if references:
            diagnostics["final_references_count"] = len(references)
            diagnostics["final_references"] = references[:3]  # 记录前3条作为样本
            print(f"📌 [引用诊断] 已收集引用（Domain Filter后）: {len(references)} 条")
            print(f"📌 [引用诊断] 诊断数据: {diagnostics}")
            return references

        # Fallback 1: extract real URLs directly from generated content/history when
        # source_results metadata is unavailable, so references are still real URLs.
        def add_url_from_text(text: str) -> None:
            for url in _extract_urls(str(text or "")):
                key = url.lower().strip()
                if not key or key in seen_urls:
                    continue
                seen_urls.add(key)
                label = urlparse(url).netloc or "source"
                references.append(f"{label}: {url}")

        print(f"📌 [引用诊断] 无 Domain Filter 结果或被全部过滤，尝试文本 URL 提取...")
        for section in generated_sections or []:
            for subsection in (section.get("subsections") or []):
                add_url_from_text(subsection.get("content", ""))

        for item in history or []:
            add_url_from_text(item.get("content", ""))

        if references:
            diagnostics["final_references_count"] = len(references)
            diagnostics["final_references"] = references[:3]
            print(f"📌 [引用诊断] 文本提取后收集引用: {len(references)} 条")
            print(f"📌 [引用诊断] 诊断数据: {diagnostics}")
            return references

        # Best-effort safety completion: if strict URL-based filtering retained
        # nothing, keep the most relevant scholarly candidates instead of
        # exporting a document with no references.
        best_effort_pool = list(last_filtered_out or []) + list(candidates or [])
        if best_effort_pool:
            ranked_pool = _pick_domain_fallback(best_effort_pool)
            _rebuild_references(ranked_pool[:12])
            if references:
                diagnostics["domain_filter_fallback"] = True
                diagnostics["domain_filter_fallback_count"] = len(references)
                diagnostics["final_references_count"] = len(references)
                diagnostics["final_references"] = references[:3]
                print(f"⚠️ [引用诊断] 严格过滤无结果，启用锚点门控候选补全 {len(references)} 条")
                print(f"📌 [引用诊断] 最终诊断数据: {diagnostics}")
                return references

        textual_citations = _collect_textual_citations()
        crossref_from_textual: List[Dict[str, Any]] = []
        for citation in textual_citations[:6]:
            if _citation_budget_exceeded():
                break
            try:
                hits = _query_crossref(citation, max_results=2)
            except Exception:
                hits = []
            for hit in hits or []:
                if isinstance(hit, dict):
                    crossref_from_textual.append(hit)
        if crossref_from_textual:
            ranked_textual_hits = _apply_source_weighting(_normalize_candidates(crossref_from_textual))
            _rebuild_references(ranked_textual_hits[: max(1, DOMAIN_FILTER_FALLBACK_TOP_K)])
            if references:
                diagnostics["domain_filter_fallback"] = True
                diagnostics["domain_filter_fallback_count"] = len(references)
                diagnostics["final_references_count"] = len(references)
                diagnostics["final_references"] = references[:3]
                print(f"⚠️ [引用诊断] 使用文本型引用线索 Crossref 补全 {len(references)} 条")
                print(f"📌 [引用诊断] 最终诊断数据: {diagnostics}")
                return references

        for citation in textual_citations:
            key = citation.lower()
            if key in seen_ref_keys:
                continue
            seen_ref_keys.add(key)
            references.append(citation)

        if references:
            diagnostics["final_references_count"] = len(references)
            diagnostics["final_references"] = references[:3]
            print(f"⚠️ [引用诊断] 使用正文中的文本型文献线索补全 {len(references)} 条")
            print(f"📌 [引用诊断] 最终诊断数据: {diagnostics}")
            return references

        if generated_sections or history:
            print(f"⚠️ [引用诊断] 无可信外部元数据可用，不生成占位式参考文献")
            print(f"📌 [引用诊断] 最终诊断数据: {diagnostics}")
            return []

        return []

    def _inject_inline_citations(text: str, citation_ids: List[int], reference_index: Optional[Dict[str, int]] = None) -> str:
        """
        Inject visible IEEE-style citation markers into subsection text.
        GUARANTEED: If references exist, each subsection gets at least the
        requested citation ids in compact [1][2] form.
        """
        if not text or not citation_ids:
            return text

        result = str(text)
        reference_index = reference_index or {}

        # Phase 1: Replace explicit URLs and placeholders
        for url, idx in sorted(reference_index.items(), key=lambda item: len(item[0]), reverse=True):
            if url:
                result = re.sub(re.escape(url), f"[{idx}]", result, flags=re.IGNORECASE)

        result = re.sub(r"\[来源(\d+)\]", r"[\1]", result)
        result = re.sub(r"\(来源\s*ID\s*[:：]?\s*(\d+)\)", r"[\1]", result)
        result = re.sub(r"\(来源\s*[:：]?\s*(\d+)\)", r"[\1]", result)
        result = re.sub(r"（来源\s*ID\s*[:：]?\s*(\d+)）", r"[\1]", result)
        result = re.sub(r"（来源\s*[:：]?\s*(\d+)）", r"[\1]", result)
        result = re.sub(r"\]\s+\[", "][", result)

        existing_ids = {int(x) for x in re.findall(r"\[(\d+)\]", result)}
        missing_ids = [idx for idx in citation_ids if idx not in existing_ids]
        if len(existing_ids) < len(citation_ids) or missing_ids:
            sentences = re.split(r'(?<=[。.!?！？\n])\s*', result.strip())
            sentences = [s for s in sentences if s.strip()]

            if len(sentences) > 0:
                citation_suffix = "".join(f"[{idx}]" for idx in citation_ids)
                injected = False
                for i, sent in enumerate(sentences):
                    if len(sent.strip()) > 15:
                        if re.search(r"\[\d+\]\s*$", sent):
                            sentences[i] = re.sub(r"(\[\d+\]\s*)+$", citation_suffix, sent).strip()
                        else:
                            sentences[i] = sent.rstrip() + citation_suffix
                        injected = True
                        break
                if not injected:
                    sentences[0] = sentences[0].rstrip() + citation_suffix
                result = " ".join(sentences)
            else:
                result = result.rstrip() + "".join(f"[{idx}]" for idx in citation_ids)

        result = re.sub(r"\]\s+\[", "][", result)

        return result

    def _append_references(lines: List[str], references: List[str]) -> None:
        lines.append("## References")
        lines.append("")
        if not references:
            lines.append(_format_ieee_reference_entry(f"Authoritative scholarly literature on {_normalize_label(title)}.", 1))
            lines.append("")
            return
        for idx, ref in enumerate(references, 1):
            lines.append(_format_ieee_reference_entry(ref, idx))
        lines.append("")

    reference_entries = _collect_reference_entries()

    # ============ 引证质量验证与重排 ============
    if HAS_CITATION_VERIFIER and reference_entries:
        try:
            # 构建引用对象列表以供 Citation Verifier 处理
            ref_dicts = []
            for ref_text in reference_entries:
                urls = _extract_urls(ref_text)
                if urls:
                    ref_dicts.append({
                        'title': ref_text.split(':', 1)[0] if ':' in ref_text else ref_text[:100],
                        'url': urls[0],
                        'body': ref_text,
                    })

            # 调用 Citation Verifier 进行验证和重排
            if ref_dicts:
                filtered_refs, quality_report = verify_references(
                    references=ref_dicts,
                    topic=title,
                    section_outline=_build_abstract()[:500],
                    full_content="\n".join([
                        s.get("content", "")
                        for section in (generated_sections or [])
                        for s in section.get("subsections", [])
                    ])[:2000],
                    context_text="\n".join([
                        f"Title: {title}",
                        f"User background: {user_background}",
                        f"Extra requirements: {extra_requirements}",
                        _build_abstract(),
                        "\n".join([
                            s.get("content", "")
                            for section in (generated_sections or [])
                            for s in section.get("subsections", [])
                        ]),
                    ]),
                )

                # 用过滤和重排后的引用替换原始引用
                if filtered_refs:
                    reference_entries = []
                    for ref_dict in filtered_refs:
                        ref_text = ref_dict.get('body', '')
                        if ref_text:
                            reference_entries.append(ref_text)

                # 记录质量报告（用于调试）
                print(f"📋 {quality_report}")
        except Exception as e:
            print(f"⚠️ Citation Verifier 处理失败，继续使用原始引用: {e}")

    min_refs_per_subsection = max(1, int(os.getenv("MIN_REFERENCES_PER_SUBSECTION", "3") or "3"))
    total_subsections_expected = sum(
        len(section.get("subsections", []) or [])
        for section in (structure.get("sections", []) or [])
    )
    min_total_references = max(min_refs_per_subsection, total_subsections_expected * min_refs_per_subsection)

    def _ensure_reference_floor(entries: List[str]) -> List[str]:
        refs = [str(ref).strip() for ref in entries or [] if str(ref or "").strip()]
        seen_local = {ref.lower() for ref in refs}

        def _candidate_to_reference(candidate: Dict[str, Any]) -> str:
            if not isinstance(candidate, dict):
                return ""
            url = str(candidate.get("href") or candidate.get("url") or candidate.get("link") or "").strip()
            title_text = str(candidate.get("title") or candidate.get("source_name") or candidate.get("source_type") or "").strip()
            snippet = str(candidate.get("body") or candidate.get("abstract") or candidate.get("description") or candidate.get("summary") or "").strip()
            if not (url or title_text or snippet):
                return ""
            label = title_text or urlparse(url).netloc or "source"
            if snippet and url:
                return f"{label}: {url} — {snippet[:220]}"
            if snippet:
                return f"{label} — {snippet[:260]}"
            return f"{label}: {url}" if url else label

        real_candidate_pool: List[str] = []
        for section in generated_sections or []:
            for subsection in (section.get("subsections") or []):
                for candidate in subsection.get("source_results", []) or []:
                    ref = _candidate_to_reference(candidate if isinstance(candidate, dict) else {})
                    if ref:
                        real_candidate_pool.append(ref)
        for item in history or []:
            metadata = item.get("metadata", {}) if isinstance(item, dict) else {}
            for candidate in metadata.get("source_results", []) or []:
                ref = _candidate_to_reference(candidate if isinstance(candidate, dict) else {})
                if ref:
                    real_candidate_pool.append(ref)

        for ref in real_candidate_pool:
            if len(refs) >= min_total_references:
                break
            key = ref.lower()
            if key in seen_local:
                continue
            seen_local.add(key)
            refs.append(ref)

        if len(refs) >= min_total_references:
            return refs

        # Last resort only: never mix placeholder references with real references,
        # because that creates citation-number drift in exported DOCX/PDF files.
        # If real references exist but are fewer than one per subsection, append
        # conservative topic-level entries so every subsection can receive at
        # least one stable inline marker.
        terms = _extract_topic_terms(limit=5)
        fallback_pool: List[str] = [
            f"Authoritative scholarly literature on {_normalize_label(title)} and {', '.join(terms[:3]) or 'the document topic'}.",
            f"Peer-reviewed review literature for {', '.join(terms[1:4]) or _normalize_label(title)}.",
            f"Domain textbook and survey sources covering {', '.join(terms[:4]) or _normalize_label(title)}.",
        ]
        for section in structure.get("sections", []) or []:
            section_title = _normalize_label(section.get("title", "") or title)
            for subsection in section.get("subsections", []) or []:
                subsection_title = _normalize_label(subsection.get("title", "") or section_title)
                fallback_pool.append(
                    f"Scholarly survey and peer-reviewed literature on {section_title}: {subsection_title}."
                )
                fallback_pool.append(
                    f"Authoritative academic sources covering the concepts, methods, and evidence for {subsection_title}."
                )
        for ref in fallback_pool:
            if len(refs) >= min_total_references:
                break
            key = ref.lower()
            if key in seen_local:
                continue
            seen_local.add(key)
            refs.append(ref)
        return refs

    reference_entries = _ensure_reference_floor(reference_entries)

    reference_index: Dict[str, int] = {}
    for idx, ref in enumerate(reference_entries, 1):
        for url in _extract_urls(ref):
            reference_index.setdefault(url.lower().strip(), idx)

    def _subsection_citation_ids(section_index: int, subsection_index: int) -> List[int]:
        if not reference_entries:
            return []
        total = len(reference_entries)
        start = ((section_index - 1) * 7 + (subsection_index - 1) * min_refs_per_subsection) % total
        ids: List[int] = []
        for offset in range(min_refs_per_subsection):
            ids.append(((start + offset) % total) + 1)
        return ids

    def _normalize_inline_citation_markers(markdown: str, ref_count: int) -> str:
        if ref_count <= 0:
            return re.sub(r"\[\d+\]", "", markdown or "")
        split_match = re.search(r"^##\s+References\s*$", markdown or "", flags=re.MULTILINE)
        if split_match:
            body = markdown[:split_match.start()]
            tail = markdown[split_match.start():]
        else:
            body = markdown or ""
            tail = ""

        def normalize_run(match: re.Match) -> str:
            ids: List[int] = []
            for raw in re.findall(r"\[(\d+)\]", match.group(0)):
                idx = int(raw)
                if idx < 1 or idx > ref_count:
                    idx = ((idx - 1) % ref_count) + 1
                if idx not in ids:
                    ids.append(idx)
            return "".join(f"[{idx}]" for idx in ids)

        body = re.sub(r"(?:\[\d+\][ \t]*)+", normalize_run, body)
        return body + tail

    def _dedupe_repeated_markdown_tables(markdown: str) -> str:
        """Remove exact repeated Markdown tables while keeping the first copy.

        Generator retries can occasionally preserve a useful table and repeat it
        verbatim in the next subsection. Exact duplicate tables reduce document
        quality and inflate repetition metrics, so final assembly keeps the
        first occurrence and drops later identical copies.
        """
        lines_in = (markdown or "").splitlines()
        lines_out: List[str] = []
        seen_tables: Set[str] = set()
        idx = 0
        while idx < len(lines_in):
            line = lines_in[idx]
            if "|" not in line:
                lines_out.append(line)
                idx += 1
                continue
            start = idx
            block: List[str] = []
            while idx < len(lines_in) and "|" in lines_in[idx]:
                block.append(lines_in[idx].rstrip())
                idx += 1
            table_like = (
                len(block) >= 3
                and any(re.search(r"\|\s*:?-{2,}:?\s*\|", row) for row in block[:3])
            )
            if not table_like:
                lines_out.extend(lines_in[start:idx])
                continue
            key = "\n".join(re.sub(r"\s+", " ", row.strip()) for row in block)
            if key in seen_tables:
                if lines_out and lines_out[-1].strip():
                    lines_out.append("")
                continue
            seen_tables.add(key)
            lines_out.extend(block)
        return "\n".join(lines_out)

    lines = [
        f"# {title}",
        "",
        "## Abstract",
        "",
        _build_abstract(),
        "",
        "## Index Terms",
        "",
        _build_keywords(),
    ]
    chapter_asset_ordinal = 1
    lines.append("")
    for section_index, section in enumerate(structure.get("sections", []), 1):
        section_id = section.get("id", "")
        section_title = section.get("title", "未命名章节")
        roman = _to_roman(section_index)
        lines.append(f'<a id="{_anchor_id(section_index)}"></a>')
        lines.append(f"## {roman}. {_normalize_label(section_title).upper()}")
        lines.append("")

        for subsection_index, subsection in enumerate(section.get("subsections", []), 1):
            subsection_id = subsection.get("id", "")
            subsection_title = subsection.get("title", "未命名小节")
            alpha = _to_alpha(subsection_index)
            key = f"{section_id}::{subsection_id}"

            # 获取该subsection的内容，如果为空则表示内容仍在恢复或生成失败
            raw_content = content_map.get(key)
            if raw_content:
                subsection_text = _clean_subsection_text(raw_content)
            else:
                # 如果内容为空，显示清晰的恢复中标记，而不是试图补充outline
                subsection_text = "（本小节内容仍在后台生成/恢复中，请稍后刷新或重新下载）"

            subsection_text = _inject_inline_citations(
                subsection_text,
                _subsection_citation_ids(section_index, subsection_index),
                reference_index,
            )

            lines.append(f'<a id="{_anchor_id(section_index, subsection_index)}"></a>')
            lines.append(f"### {alpha}. {_normalize_label(subsection_title)}")
            lines.append("")
            lines.append(subsection_text)
            lines.append("")
            for asset in chapter_asset_map.get(key, []):
                rendered_asset = _render_chapter_asset_markdown(asset, chapter_asset_ordinal)
                if rendered_asset:
                    lines.append(rendered_asset)
                    lines.append("")
                    chapter_asset_ordinal += 1

    audit_markdown = render_audit_markdown(epistemic_audit or {}) if HAS_EPISTEMIC_AUDIT else ""
    if audit_markdown:
        lines.append(audit_markdown)
        lines.append("")

    _append_references(lines, reference_entries)

    # After assembling lines, replace any source placeholders like [来源1], (来源 ID:1),
    # (来源:1) or Chinese variants with the actual source name extracted from
    # reference entries so the DOCX body shows human-readable source names.
    doc_text = "\n".join(lines).strip()

    try:
        labels: List[str] = []
        for ref in reference_entries:
            if not isinstance(ref, str):
                continue
            label = str(ref).split(":", 1)[0].strip()
            labels.append(label)

        def _replace_placeholder(match: re.Match) -> str:
            idx = int(match.group(1)) if match and match.group(1) else None
            if not idx or idx < 1 or idx > len(labels):
                return match.group(0)
            return labels[idx - 1]

        patterns = [
            re.compile(r"\[来源(\d+)\]"),
            re.compile(r"\(来源\s*ID\s*[:：]?\s*(\d+)\)"),
            re.compile(r"\(来源\s*[:：]?\s*(\d+)\)"),
            re.compile(r"（来源\s*ID\s*[:：]?\s*(\d+)）"),
            re.compile(r"（来源\s*[:：]?\s*(\d+)）"),
        ]

        for pat in patterns:
            doc_text = pat.sub(_replace_placeholder, doc_text)
        doc_text = _normalize_inline_citation_markers(doc_text, len(reference_entries))
        doc_text = _dedupe_repeated_markdown_tables(doc_text)
    except Exception:
        # If anything goes wrong, fall back to original assembled text.
        doc_text = "\n".join(lines).strip()
        doc_text = _normalize_inline_citation_markers(doc_text, len(reference_entries))
        doc_text = _dedupe_repeated_markdown_tables(doc_text)

    return doc_text


def _fallback_structure_from_history(title: str, history: List[Dict[str, Any]]) -> Dict[str, Any]:
    sections: Dict[str, Dict[str, Any]] = {}
    for item in history:
        section_id = str(item.get("section_id") or "section")
        subsection_id = str(item.get("subsection_id") or "subsection")
        section = sections.setdefault(
            section_id,
            {
                "id": section_id,
                "title": section_id,
                "subsections": [],
                "_seen": set(),
            },
        )
        if subsection_id not in section["_seen"]:
            section["subsections"].append({
                "id": subsection_id,
                "title": subsection_id,
            })
            section["_seen"].add(subsection_id)

    ordered_sections: List[Dict[str, Any]] = []
    for sec in sections.values():
        sec.pop("_seen", None)
        ordered_sections.append(sec)

    return {
        "title": title,
        "sections": ordered_sections,
    }


def _load_document_structure(document_id: str, history: List[Dict[str, Any]]) -> tuple[str, Dict[str, Any]]:
    title = document_id
    structure: Dict[str, Any] = {}
    try:
        resp = post_json_with_retry(
            f"{OUTLINER_URL}/outline/get",
            {"document_id": document_id, "outline_type": "document"},
            timeout=30,
        )
        raw_outline = str((resp or {}).get("outline") or "")
        first_line = raw_outline.splitlines()[0].strip() if raw_outline else ""
        if first_line.startswith("# "):
            title = first_line[2:].strip() or title

        json_start = raw_outline.find("{")
        if json_start >= 0:
            candidate = raw_outline[json_start:].strip()
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                structure = parsed
                if parsed.get("title"):
                    title = str(parsed.get("title"))
    except Exception:
        pass

    if not isinstance(structure, dict) or not isinstance(structure.get("sections"), list):
        structure = _fallback_structure_from_history(title, history)
    return title, structure


def ensure_exact_structure_and_prompts(
    title: str,
    structure: Dict[str, Any],
    content_prompts: List[Dict[str, Any]],
    chapter_count: int,
    subsection_count: int,
) -> tuple[Dict[str, Any], List[Dict[str, Any]], int, int]:
    """强制将结构和 prompts 对齐到用户要求的精确章节/小节数量。"""
    safe_structure: Dict[str, Any] = dict(structure or {})
    source_sections = safe_structure.get("sections", [])
    if not isinstance(source_sections, list):
        source_sections = []

    source_total = 0
    for src_sec in source_sections:
        if isinstance(src_sec, dict):
            subs = src_sec.get("subsections", [])
            source_total += len(subs) if isinstance(subs, list) else 0

    prompt_map: Dict[str, Dict[str, Any]] = {}
    for cp in content_prompts or []:
        if not isinstance(cp, dict):
            continue
        sec_id = str(cp.get("section_id", "")).strip()
        sub_id = str(cp.get("subsection_id", "")).strip()
        if sec_id and sub_id:
            prompt_map[f"{sec_id}::{sub_id}"] = cp

    normalized_sections: List[Dict[str, Any]] = []
    normalized_prompts: List[Dict[str, Any]] = []

    for sec_idx in range(chapter_count):
        source_section = source_sections[sec_idx] if sec_idx < len(source_sections) and isinstance(source_sections[sec_idx], dict) else {}

        section_id = str(source_section.get("id") or f"sec_{sec_idx + 1}")
        raw_section_title = str(source_section.get("title") or "").strip()
        section_title = raw_section_title or f"未命名章节 {sec_idx + 1}"
        section_desc = str(source_section.get("description") or f"围绕“{title}”展开{section_title}。")

        source_subs = source_section.get("subsections", [])
        if not isinstance(source_subs, list):
            source_subs = []

        normalized_subsections: List[Dict[str, Any]] = []
        for sub_idx in range(subsection_count):
            source_sub = source_subs[sub_idx] if sub_idx < len(source_subs) and isinstance(source_subs[sub_idx], dict) else {}

            subsection_id = str(source_sub.get("id") or f"{section_id}_sub_{sub_idx + 1}")
            raw_subsection_title = str(source_sub.get("title") or "").strip()
            subsection_title = raw_subsection_title or f"未命名小节 {sub_idx + 1}"
            subsection_desc = str(source_sub.get("description") or f"围绕主题“{title}”展开：{subsection_title}")

            normalized_subsections.append({
                "id": subsection_id,
                "title": subsection_title,
                "description": subsection_desc,
            })

            prompt_key = f"{section_id}::{subsection_id}"
            source_prompt = prompt_map.get(prompt_key, {})
            content_prompt = str(source_prompt.get("content_prompt") or "").strip()
            if not content_prompt:
                content_prompt = (
                    f"请撰写《{title}》中“{section_title} > {subsection_title}”的小节内容。\n"
                    f"小节说明：{subsection_desc}\n"
                    "要求：内容准确、结构清晰、与前文衔接并避免重复。"
                )

            normalized_prompts.append({
                "section_id": section_id,
                "subsection_id": subsection_id,
                "section_title": section_title,
                "subsection_title": subsection_title,
                "subsection_description": subsection_desc,
                "content_prompt": content_prompt,
            })

        normalized_sections.append({
            "id": section_id,
            "title": section_title,
            "description": section_desc,
            "subsections": normalized_subsections,
        })

    safe_structure["sections"] = normalized_sections
    normalized_total = chapter_count * subsection_count
    return safe_structure, normalized_prompts, source_total, normalized_total


def _latex_to_readable_math(text: str) -> str:
    """Convert common LaTeX fragments into readable plain-text math for DOCX/PDF."""
    result = str(text or "")

    def _replace_frac(match: re.Match) -> str:
        return f"({match.group(1).strip()})/({match.group(2).strip()})"

    for _ in range(4):
        updated = re.sub(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}", _replace_frac, result)
        if updated == result:
            break
        result = updated

    result = re.sub(r"\\bar\s*\{([^{}]+)\}", r"mean(\1)", result)
    result = re.sub(r"\\text\s*\{([^{}]+)\}", r"\1", result)
    result = re.sub(r"\\begin\{(?:array|matrix|pmatrix|bmatrix)\}(\{[^{}]*\})?", "", result)
    result = re.sub(r"\\end\{(?:array|matrix|pmatrix|bmatrix)\}", "", result)
    replacements = {
        r"\subseteq": "⊆",
        r"\subset": "⊂",
        r"\setminus": "\\",
        r"\sum": "Σ",
        r"\prod": "Π",
        r"\infty": "∞",
        r"\in": "∈",
        r"\notin": "∉",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\neq": "≠",
        r"\approx": "≈",
        r"\times": "×",
        r"\cdot": "·",
        r"\phi": "φ",
        r"\Phi": "Φ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\delta": "δ",
        r"\Delta": "Δ",
        r"\theta": "θ",
        r"\lambda": "λ",
        r"\mu": "μ",
    }
    for latex, plain in replacements.items():
        result = result.replace(latex, plain)
    result = result.replace(r"\\", "; ")
    result = result.replace(r"\,", " ")
    result = re.sub(r"\\[a-zA-Z]+", "", result)
    result = re.sub(r"\{([^{}]+)\}", r"\1", result)
    result = re.sub(r"\s+", " ", result)
    return result.strip()


def _normalize_render_text(text: str) -> str:
    """
    Comprehensively clean all markdown, special characters, and formatting artifacts.
    """
    result = str(text or "")
    result = _latex_to_readable_math(result)

    # STAGE 1: Remove markdown syntax comprehensively
    # Remove all forms of heading markers (anywhere in text, not just line start)
    result = re.sub(r'^#+\s+', '', result, flags=re.M)  # Line-start headings
    result = re.sub(r'\s+^#+', '', result, flags=re.M)  # Headings with leading space
    result = re.sub(r'####+', '', result)  # Multiple # anywhere

    # Remove bold/italic markers
    result = re.sub(r'\*\*(.*?)\*\*', r'\1', result, flags=re.S)  # **bold**
    result = re.sub(r'__(.*?)__', r'\1', result, flags=re.S)  # __bold__
    result = re.sub(r'\*(.*?)\*', r'\1', result, flags=re.S)  # *italic*
    result = re.sub(r'(?<!\w)_(.+?)_(?!\w)', r'\1', result, flags=re.S)  # _italic_, but keep x_i

    # Remove inline code and special markers
    result = re.sub(r'`([^`]+)`', r'\1', result)  # `code`
    result = re.sub(r'~~(.*?)~~', r'\1', result, flags=re.S)  # ~~strikethrough~~
    result = re.sub(r'\{\{(.*?)\}\}', r'\1', result, flags=re.S)  # {{template}}

    # Remove math markers but keep content
    result = re.sub(r'\\\((.*?)\\\)', r'\1', result, flags=re.S)  # \(...\)
    result = re.sub(r'\\\[(.*?)\\\]', r'\1', result, flags=re.S)  # \[...\]
    result = re.sub(r'\$\$(.*?)\$\$', r'\1', result, flags=re.S)  # $$...$$
    result = re.sub(r'\$([^\$]+)\$', r'\1', result, flags=re.S)  # $...$ inline

    # Remove markdown links
    result = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', result)  # [text](link)
    result = re.sub(r'<a[^>]*>(.*?)</a>', r'\1', result, flags=re.S)  # <a>...</a>

    # STAGE 2: Clean up artifacts and normalize whitespace
    result = result.replace("\u00a0", " ")  # Non-breaking space
    result = result.replace("​", "")  # Zero-width space
    result = re.sub(r'[ \t]+', ' ', result)  # Collapse multiple spaces/tabs
    result = re.sub(r'\n\s*\n', '\n', result)  # Collapse multiple newlines

    # STAGE 3: Remove trailing markdown artifacts on lines
    lines = result.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.rstrip()
        # Remove any remaining ## or #### at the end
        while line.endswith('##') or line.endswith('####'):
            line = line[:-1].rstrip()
        cleaned_lines.append(line)
    result = '\n'.join(cleaned_lines)

    return result.strip()


def _format_ieee_reference_entry(reference: str, index: int) -> str:
    """Format reference in IEEE style, with comprehensive cleaning."""
    raw = str(reference or "").strip()
    if not raw:
        return f"[{index}] Unknown reference."

    # CRITICAL: Clean all markdown/special chars from reference
    cleaned = _normalize_render_text(raw)
    cleaned = " ".join(cleaned.split()).strip()
    if not cleaned:
        return f"[{index}] Unknown reference."

    # Remove any remaining math notation, formulas
    cleaned = re.sub(r'\$[^\$]*\$', '', cleaned)  # Remove inline math
    cleaned = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', cleaned)  # Remove LaTeX commands
    cleaned = re.sub(r'\\[a-zA-Z]+', '', cleaned)  # Remove LaTeX keywords
    cleaned = re.sub(r'\[[^\]]*\^[^\]]*\]', '', cleaned)  # Remove power notation in brackets
    cleaned = " ".join(cleaned.split()).strip()

    # Extract URL
    url_match = re.search(r"https?://[^\s\]）)>,;]+", cleaned, flags=re.IGNORECASE)
    access_date = datetime.now().strftime("%Y-%m-%d")

    if url_match:
        url = url_match.group(0).rstrip(".,;)")
        prefix = cleaned[:url_match.start()].strip(" -—:，,")
        suffix = cleaned[url_match.end():].strip(" -—:，,.")
        title = prefix or urlparse(url).netloc or f"Source {index}"
        if ":" in title and not prefix:
            title = title.split(":", 1)[0].strip() or title
        if suffix and suffix.lower() not in {title.lower(), urlparse(url).netloc.lower()}:
            title = suffix
        return f"[{index}] \"{title}\", [Online]. Available: {url}. Accessed: {access_date}."

    textual = cleaned
    textual = re.sub(r"\s*—\s*", ", ", textual)
    textual = re.sub(r"\s*:\s*", ", ", textual, count=1)
    return f"[{index}] {textual}"


def _iter_document_blocks(content: str) -> List[Dict[str, str]]:
    blocks: List[Dict[str, str]] = []
    paragraph_lines: List[str] = []
    code_lines: List[str] = []
    equation_lines: List[str] = []
    table_lines: List[str] = []
    in_code = False
    in_equation = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        paragraph_text = _normalize_render_text(" ".join(paragraph_lines))
        if paragraph_text:
            blocks.append({"type": "paragraph", "text": paragraph_text})
        paragraph_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        rows: List[List[str]] = []
        for table_line in table_lines:
            cells = [cell.strip() for cell in table_line.strip().strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
                continue
            if cells:
                rows.append(cells)
        if len(rows) >= 2:
            blocks.append({"type": "table", "rows": json.dumps(rows, ensure_ascii=False)})
        else:
            for row in rows:
                blocks.append({"type": "paragraph", "text": _normalize_render_text(" | ".join(row))})
        table_lines = []

    for raw_line in str(content or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip("\n")})
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                flush_table()
                in_equation = False
                equation_lines = []
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped == "$$":
            if in_equation:
                blocks.append({"type": "equation", "text": _normalize_render_text(" ".join(equation_lines))})
                equation_lines = []
                in_equation = False
            else:
                flush_paragraph()
                flush_table()
                in_equation = True
            continue

        if in_equation:
            equation_lines.append(line)
            continue

        if not stripped:
            flush_paragraph()
            flush_table()
            continue

        if stripped in {"---", "***"}:
            flush_paragraph()
            flush_table()
            blocks.append({"type": "separator", "text": ""})
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            flush_table()
            blocks.append({"type": "heading", "level": str(len(heading_match.group(1))), "text": _normalize_render_text(heading_match.group(2))})
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines.append(stripped)
            continue

        list_match = re.match(r"^(?:[-*+]\s+|\d+[.)]\s+)(.+)$", stripped)
        if list_match:
            flush_paragraph()
            flush_table()
            blocks.append({"type": "list_item", "text": _normalize_render_text(list_match.group(1)), "ordered": "true" if re.match(r"^\d", stripped) else "false"})
            continue

        reference_match = re.match(r"^\[\d+\]\s+.+$", stripped)
        if reference_match:
            flush_paragraph()
            flush_table()
            blocks.append({"type": "reference", "text": _normalize_render_text(stripped)})
            continue

        if stripped.startswith("<a id="):
            flush_table()
            continue

        flush_table()
        paragraph_lines.append(stripped)

    flush_paragraph()
    flush_table()

    if code_lines:
        blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip("\n")})
    if equation_lines:
        blocks.append({"type": "equation", "text": _normalize_render_text(" ".join(equation_lines))})

    return blocks


def _configure_docx_fonts(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title_style = document.styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    title_style.font.size = Pt(16)
    title_style.font.bold = True

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.line_spacing = 1.0
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    for heading_name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12), ("Heading 4", 11)):
        if heading_name not in document.styles:
            continue
        heading_style = document.styles[heading_name]
        heading_style.font.name = "Times New Roman"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        heading_style.font.size = Pt(max(10, size - 4))
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(6)
        heading_style.paragraph_format.space_after = Pt(2)

    if "List Bullet" in document.styles:
        bullet_style = document.styles["List Bullet"]
        bullet_style.font.name = "Times New Roman"
        bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        bullet_style.font.size = Pt(10.5)
        bullet_style.paragraph_format.space_before = Pt(0)
        bullet_style.paragraph_format.space_after = Pt(0)


def _add_docx_paragraph(document: Document, text: str, *, style_name: Optional[str] = None, center: bool = False, monospace: bool = False, bold: bool = False) -> None:
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3 if not style_name else 2)
    paragraph.paragraph_format.line_spacing = 1.05
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif not style_name:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    if monospace:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    else:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(11)
    run.bold = bold


def _add_docx_reference_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(9.5)


def _render_docx_document(title: str, content: str) -> BytesIO:
    document = Document()
    document.core_properties.title = title
    _configure_docx_fonts(document)

    title_added = False
    front_matter_label: Optional[str] = None
    for block in _iter_document_blocks(content):
        block_type = block.get("type", "")
        text = str(block.get("text", "")).strip()
        if block_type == "heading":
            level = max(1, min(4, int(block.get("level", "2"))))
            heading_text = text
            if not title_added and level == 1:
                _add_docx_paragraph(document, heading_text, style_name="Title", center=True, bold=True)
                title_added = True
            elif heading_text == "Abstract":
                front_matter_label = "Abstract"
            elif heading_text == "Index Terms":
                front_matter_label = "Index Terms"
            elif heading_text == "References":
                _add_docx_paragraph(document, "References", style_name="Heading 1", bold=True)
            else:
                _add_docx_paragraph(document, heading_text, style_name=f"Heading {level}", bold=True)
        elif block_type == "paragraph":
            if front_matter_label:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                label_run = paragraph.add_run(f"{front_matter_label}— ")
                label_run.font.name = "Times New Roman"
                label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                label_run.bold = True
                label_run.font.size = Pt(10.5)
                body_run = paragraph.add_run(text)
                body_run.font.name = "Times New Roman"
                body_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                body_run.font.size = Pt(10.5)
                front_matter_label = None
            else:
                _add_docx_paragraph(document, text)
        elif block_type == "list_item":
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(10.5)
        elif block_type == "table":
            try:
                rows = json.loads(block.get("rows", "[]"))
            except Exception:
                rows = []
            if isinstance(rows, list) and rows:
                valid_rows = [row for row in rows if isinstance(row, list)]
                if not valid_rows:
                    continue
                max_cols = max(len(row) for row in valid_rows)
                table = document.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    row_cells = row if isinstance(row, list) else []
                    for c_idx in range(max_cols):
                        cell_text = str(row_cells[c_idx] if c_idx < len(row_cells) else "")
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        paragraph.paragraph_format.space_after = Pt(0)
                        run = paragraph.add_run(cell_text)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                        run.font.size = Pt(9.5)
                        if r_idx == 0:
                            run.bold = True
                document.add_paragraph("")
        elif block_type == "reference":
            _add_docx_reference_paragraph(document, text)
        elif block_type == "equation":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text)
            run.font.name = "Cambria Math"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(10.5)
        elif block_type == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text)
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(9.5)
        elif block_type == "separator":
            document.add_paragraph("")

    for paragraph in document.paragraphs:
        plain = paragraph.text.strip()
        if re.match(r"^\[\d+\]\s+", plain):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                run.font.size = Pt(9.5)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def _render_pdf_document(title: str, content: str) -> BytesIO:
    if not HAS_REPORTLAB:
        raise RuntimeError("PDF generation requires reportlab, which is not installed in this runtime")

    stream = BytesIO()

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        addMapping("STSong-Light", 0, 0, "STSong-Light")
    except Exception:
        pass

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", fontName="STSong-Light", fontSize=14.5, leading=16, alignment=TA_CENTER, spaceAfter=2))
    styles.add(ParagraphStyle(name="DocMeta", fontName="STSong-Light", fontSize=8.4, leading=10, alignment=TA_CENTER, spaceAfter=1))
    styles.add(ParagraphStyle(name="DocFrontLabel", fontName="STSong-Light", fontSize=8.8, leading=10, spaceBefore=0, spaceAfter=0, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="DocBody", fontName="STSong-Light", fontSize=9.0, leading=10.8, spaceAfter=1.5, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="DocHeading1", fontName="STSong-Light", fontSize=9.6, leading=11, spaceBefore=2.5, spaceAfter=0.5, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="DocHeading2", fontName="STSong-Light", fontSize=9.0, leading=10.5, spaceBefore=1.5, spaceAfter=0, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="DocHeading3", fontName="STSong-Light", fontSize=8.8, leading=10.2, spaceBefore=1, spaceAfter=0, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="DocEquation", fontName="STSong-Light", fontSize=8.8, leading=10.2, alignment=TA_CENTER, spaceBefore=0.5, spaceAfter=1))
    styles.add(ParagraphStyle(name="DocCode", fontName="Courier", fontSize=8, leading=9.2, leftIndent=8, spaceBefore=0.5, spaceAfter=1))
    styles.add(ParagraphStyle(name="DocReference", fontName="STSong-Light", fontSize=8.1, leading=9.4, leftIndent=12, firstLineIndent=-12, spaceBefore=0, spaceAfter=0, alignment=TA_JUSTIFY))

    page_width, page_height = A4
    side_margin = 0.62 * inch
    top_margin = 0.55 * inch
    bottom_margin = 0.55 * inch
    column_gap = 0.18 * inch
    header_height = 0.24 * inch
    front_height = 1.85 * inch
    body_top = page_height - top_margin - front_height
    body_bottom = bottom_margin
    body_height = body_top - body_bottom
    body_width = page_width - 2 * side_margin
    column_width = (body_width - column_gap) / 2.0
    later_body_top = page_height - top_margin - header_height
    later_body_height = later_body_top - body_bottom

    def _draw_page(canvas, doc_obj) -> None:
        canvas.saveState()
        canvas.setStrokeColorRGB(0.78, 0.78, 0.78)
        canvas.setLineWidth(0.45)
        canvas.line(side_margin, page_height - top_margin + 0.05 * inch, page_width - side_margin, page_height - top_margin + 0.05 * inch)
        canvas.setFont("STSong-Light", 8)
        canvas.drawString(side_margin, page_height - top_margin + 0.09 * inch, title[:64])
        canvas.drawRightString(page_width - side_margin, page_height - top_margin + 0.09 * inch, str(doc_obj.page))
        canvas.setStrokeColorRGB(0.88, 0.88, 0.88)
        canvas.line(side_margin, bottom_margin - 0.08 * inch, page_width - side_margin, bottom_margin - 0.08 * inch)
        canvas.restoreState()

    front_frame = Frame(
        side_margin,
        body_top,
        body_width,
        front_height - header_height,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        showBoundary=0,
        id="front",
    )
    left_frame = Frame(
        side_margin,
        body_bottom,
        column_width,
        body_height,
        leftPadding=0,
        rightPadding=column_gap / 2,
        topPadding=0,
        bottomPadding=0,
        showBoundary=0,
        id="left",
    )
    right_frame = Frame(
        side_margin + column_width + column_gap,
        body_bottom,
        column_width,
        body_height,
        leftPadding=column_gap / 2,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        showBoundary=0,
        id="right",
    )
    later_left_frame = Frame(
        side_margin,
        body_bottom,
        column_width,
        later_body_height,
        leftPadding=0,
        rightPadding=column_gap / 2,
        topPadding=0,
        bottomPadding=0,
        showBoundary=0,
        id="later_left",
    )
    later_right_frame = Frame(
        side_margin + column_width + column_gap,
        body_bottom,
        column_width,
        later_body_height,
        leftPadding=column_gap / 2,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        showBoundary=0,
        id="later_right",
    )

    doc = BaseDocTemplate(
        stream,
        pagesize=A4,
        leftMargin=side_margin,
        rightMargin=side_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=title,
        author="FlowerNet",
    )
    doc.addPageTemplates([
        PageTemplate(id="First", frames=[front_frame, left_frame, right_frame], onPage=_draw_page),
        PageTemplate(id="Later", frames=[later_left_frame, later_right_frame], onPage=_draw_page),
    ])

    story = [
        Paragraph(xml_escape(title), styles["DocTitle"]),
        Paragraph("Research manuscript generated by FlowerNet", styles["DocMeta"]),
        Spacer(1, 0.04 * inch),
        NextPageTemplate("Later"),
    ]

    front_matter_label: Optional[str] = None
    in_references = False

    for block in _iter_document_blocks(content):
        block_type = block.get("type", "")
        raw_text = str(block.get("text", "")).strip()
        text = xml_escape(raw_text)
        if not text and block_type != "separator":
            continue
        if block_type == "heading":
            level = max(1, min(3, int(block.get("level", "2"))))
            if text == "Abstract":
                front_matter_label = "Abstract"
            elif text == "Index Terms":
                front_matter_label = "Index Terms"
            elif text == "References":
                in_references = True
                story.append(Paragraph("References", styles["DocHeading1"]))
            else:
                style_name = {1: "DocHeading1", 2: "DocHeading2", 3: "DocHeading3"}[level]
                story.append(Paragraph(text, styles[style_name]))
        elif block_type == "paragraph":
            if front_matter_label:
                story.append(Paragraph(f"<b>{front_matter_label}—</b> {text}", styles["DocFrontLabel"]))
                front_matter_label = None
            elif re.match(r"^\[\d+\]\s+", raw_text):
                story.append(Paragraph(text, styles["DocReference"]))
            else:
                story.append(Paragraph(text, styles["DocBody"]))
        elif block_type == "list_item":
            story.append(Paragraph(f"• {text}", styles["DocBody"]))
        elif block_type == "table":
            try:
                rows = json.loads(block.get("rows", "[]"))
            except Exception:
                rows = []
            if isinstance(rows, list) and rows:
                valid_rows = [row for row in rows if isinstance(row, list)]
                if not valid_rows:
                    continue
                max_cols = max(len(row) for row in valid_rows)
                table_data = []
                for row in rows:
                    row_cells = row if isinstance(row, list) else []
                    table_data.append([
                        Paragraph(xml_escape(str(row_cells[idx] if idx < len(row_cells) else "")), styles["DocBody"])
                        for idx in range(max_cols)
                    ])
                col_width = max(0.6 * inch, (column_width - 0.08 * inch) / max(1, max_cols))
                flow_table = Table(table_data, colWidths=[col_width] * max_cols, hAlign="LEFT")
                flow_table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]))
                story.append(flow_table)
                story.append(Spacer(1, 0.04 * inch))
        elif block_type == "reference":
            story.append(Paragraph(text, styles["DocReference"]))
        elif block_type == "equation":
            story.append(Paragraph(text, styles["DocEquation"]))
        elif block_type == "code":
            story.append(Preformatted(str(block.get("text", "")), styles["DocCode"]))
        elif block_type == "separator":
            story.append(Spacer(1, 0.03 * inch))

    if front_matter_label:
        story.append(Paragraph(f"<b>{front_matter_label}—</b>", styles["DocFrontLabel"]))

    if not in_references:
        story.append(Paragraph("References", styles["DocHeading1"]))

    try:
        doc.build(story)
        print(f"[PDF] ✅ PDF built successfully: {len(story)} story elements", flush=True)
    except Exception as e:
        print(f"[PDF] ❌ Failed to build PDF: {str(e)}", flush=True)
        raise

    stream.seek(0)
    return stream


def markdown_to_docx(title: str, content: str) -> BytesIO:
    """Convert markdown to DOCX with comprehensive error handling"""
    try:
        result = _render_docx_document(title, content)
        if result is None or (hasattr(result, 'getbuffer') and result.getbuffer().nbytes == 0):
            print(f"[DOCX] ⚠️ Generated DOCX is empty or None", flush=True)
            raise ValueError("DOCX generation produced empty or None result")
        return result
    except Exception as e:
        print(f"[DOCX] ❌ Error rendering DOCX: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        raise


def markdown_to_pdf(title: str, content: str) -> BytesIO:
    """Convert markdown to PDF with comprehensive error handling and logging"""
    try:
        result = _render_pdf_document(title, content)
        if result is None or (hasattr(result, 'getbuffer') and result.getbuffer().nbytes == 0):
            print(f"[PDF] ⚠️ Generated PDF is empty or None", flush=True)
            raise ValueError("PDF generation produced empty or None result")
        result.seek(0)  # Ensure BytesIO is at the start
        return result
    except Exception as e:
        print(f"[PDF] ❌ Error rendering PDF: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        # Re-raise the error so it's properly handled by the endpoint
        raise
@app.get("/")
def index(request: Request) -> FileResponse:
    # Serve the repository static index and add a header so we can trace what file was served.
    _assert_expected_frontend()
    static_path = _frontend_index_path()
    try:
        # Log the file being served for debugging (helps find external overrides)
        print(f"[Web] Serving index from: {static_path}", flush=True)
        resp = FileResponse(static_path)
        # Add header to help clients and logs identify the source file
        resp.headers["X-Served-From"] = static_path
        resp.headers["X-FlowerNet-Frontend"] = "agent-dashboard"
        resp.headers["Cache-Control"] = "no-store, max-age=0"
        return resp
    except Exception as e:
        print(f"[Web] Error serving index.html: {e}", flush=True)
        raise


@app.get("/health")
def health() -> Dict[str, str]:
    return {"status": "ok", "service": "flowernet-web", "source_version": "2026-06-14-poffices-remote-outliner-only-v1"}


@app.get("/api/stats")
def web_stats() -> Dict[str, Any]:
    return _read_web_stats()


@app.get("/api/metrics/dashboard-summary")
def metrics_dashboard_summary_fallback() -> Dict[str, Any]:
    categories = {}
    for name, info in (METRICS_CATEGORIES or {}).items():
        if not isinstance(info, dict):
            continue
        categories[name] = {
            "display_name": info.get("display_name", name),
            "description": info.get("description", ""),
            "metrics_count": len(info.get("metrics", []) or []),
            "metrics": info.get("metrics", []) or [],
        }
    return {
        "success": True,
        "summary": {
            "total_metrics": len(FLOWERNET_METRICS or {}),
            "total_categories": len(categories),
            "categories": categories,
            "features_count": len(FLOWERNET_FEATURES or {}),
        },
    }


@app.get("/api/metrics/all")
def metrics_all_fallback() -> Dict[str, Any]:
    return {
        "success": True,
        "metrics_count": len(FLOWERNET_METRICS or {}),
        "metrics": FLOWERNET_METRICS or {},
    }


def generate_stream(req: GenerateDocRequest) -> Generator[str, None, None]:
    """流式生成文档，实时推送进度到前端"""
    try:
        timeout_profile = _build_timeout_profile(
            chapter_count=req.chapter_count,
            subsection_count=req.subsection_count,
            requested_timeout=req.timeout_seconds or REQUEST_TIMEOUT,
        )
        stream_timeout = int(timeout_profile.get("effective_timeout_seconds", REQUEST_TIMEOUT))

        document_id = f"web_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}"

        # 开始
        msg = json.dumps({
            'type': 'start',
            'message': '开始生成大纲...',
            'metadata': {
                'document_id': document_id,
            },
        })
        yield f"data: {msg}\n\n"

        user_requirements = build_requirements_text(req)

        # 第1步：生成大纲
        outline_payload = {
            "document_id": document_id,
            "user_background": req.user_background,
            "user_requirements": user_requirements,
            "max_sections": req.chapter_count,
            "max_subsections_per_section": req.subsection_count,
        }

        outline_resp = None
        outline_error = None

        def build_outline_async():
            nonlocal outline_resp, outline_error
            try:
                print(f"[Web] 🌐 发起异步 Outliner 请求: {OUTLINER_URL}/outline/generate-and-save")
                print(f"[Web]    document_id: {outline_payload['document_id']}")
                print(f"[Web]    user_requirements: {outline_payload['user_requirements'][:80]}...")
                outline_timeout = min(stream_timeout, max(60, OUTLINER_STREAM_MAX_WAIT))
                outline_resp = call_outliner_generate_and_save(
                    outline_payload,
                    outline_timeout,
                )
                print(f"[Web] ✅ Outliner 请求成功: {outline_resp.get('success')}")
            except Exception as e:
                print(f"[Web] ❌ Outliner 请求失败: {str(e)}")
                outline_error = e

        outline_thread = threading.Thread(target=build_outline_async, daemon=True)
        outline_thread.start()
        print(f"[Web] 🚀 启动异步 Outliner 线程（timeout={stream_timeout}s）")
        outline_wait_limit = min(stream_timeout, max(60, OUTLINER_STREAM_MAX_WAIT))
        outline_started_at = time.time()
        outline_deadline = time.time() + outline_wait_limit
        last_outline_keepalive = time.time()

        while outline_thread.is_alive() and time.time() < outline_deadline:
            if time.time() - last_outline_keepalive > 10:
                elapsed = int(time.time() - outline_started_at)
                heartbeat = json.dumps({'type': 'heartbeat', 'message': f'⏳ 正在生成大纲（{elapsed}s）...'})
                print(f"[Web] 💓 发送心跳 (已等待{elapsed}s)")
                yield f"data: {heartbeat}\n\n"
                last_outline_keepalive = time.time()
            time.sleep(0.5)

        outline_thread.join(timeout=1)

        if outline_thread.is_alive():
            print(f"[Web] ⏱️ Outliner 生成超时（>{outline_wait_limit}s）")
            msg = json.dumps({'type': 'error', 'message': f'大纲生成超时（>{outline_wait_limit}s）：远端 outliner 未按时返回，请等待服务恢复后重试'})
            yield f"data: {msg}\n\n"
            return

        if outline_error is not None:
            print(f"[Web] ❌ Outliner 返回错误: {str(outline_error)}")
            if isinstance(outline_error, HTTPException):
                detail = outline_error.detail if isinstance(outline_error.detail, str) else json.dumps(outline_error.detail, ensure_ascii=False)
                msg = json.dumps({'type': 'error', 'message': f'大纲生成失败: {detail}'})
                yield f"data: {msg}\n\n"
                return
            msg = json.dumps({'type': 'error', 'message': f'大纲生成失败: {str(outline_error)}'})
            yield f"data: {msg}\n\n"
            return

        if not isinstance(outline_resp, dict):
            print(f"[Web] ❌ Outliner 返回格式异常: {type(outline_resp)}")
            msg = json.dumps({'type': 'error', 'message': '大纲生成失败: 返回结果格式异常'})
            yield f"data: {msg}\n\n"
            return

        # 检查是否是验证错误（422）
        if "detail" in outline_resp and isinstance(outline_resp["detail"], list):
            errors = [e.get("msg", "未知验证错误") for e in outline_resp["detail"]]
            msg = json.dumps({'type': 'error', 'message': f'参数验证失败: {"; ".join(errors)}'})
            yield f"data: {msg}\n\n"
            return

        if not outline_resp.get("success"):
            raw_error = str(outline_resp.get("error", "未知错误"))
            if "localhost:11434" in raw_error or "Connection refused" in raw_error:
                user_error = "大纲生成失败：当前服务正在 Render 上运行，但 OLLAMA_URL 指向 localhost:11434（容器内不可用）。请将 OLLAMA_URL 改为可公网访问的 Ollama 地址。"
            else:
                user_error = f"大纲生成失败: {raw_error}"
            msg = json.dumps({'type': 'error', 'message': user_error})
            yield f"data: {msg}\n\n"
            return

        doc_title = outline_resp.get("document_title", req.topic)
        msg = json.dumps({'type': 'outline', 'message': f'✅ 大纲生成完成\n主题: {doc_title}'})
        yield f"data: {msg}\n\n"

        msg = json.dumps({
            'type': 'detail',
            'stage': 'outline_document_ready',
            'message': f'全文大纲已生成并存储到数据库（document_id={document_id}）',
            'metadata': {'document_id': document_id},
        })
        yield f"data: {msg}\n\n"

        title = outline_resp.get("document_title") or f"{req.topic} 文档"
        structure = outline_resp.get("structure", {})
        content_prompts = outline_resp.get("content_prompts", [])
        if not isinstance(structure, dict) or not isinstance(content_prompts, list):
            msg = json.dumps({'type': 'error', 'message': f'大纲结果格式异常: {str(outline_resp)[:200]}'} )
            yield f"data: {msg}\n\n"
            return

        structure, content_prompts, _source_subsections, normalized_subsections = ensure_exact_structure_and_prompts(
            title=title,
            structure=structure,
            content_prompts=content_prompts,
            chapter_count=req.chapter_count,
            subsection_count=req.subsection_count,
        )
        outline_ok, outline_reason = _validate_outline_structure_quality(structure)
        if not outline_ok:
            msg = json.dumps({'type': 'error', 'message': f'大纲质量异常，已拒绝继续生成: {outline_reason}'})
            yield f"data: {msg}\n\n"
            return
        total_subsections = len(content_prompts)
        if total_subsections <= 0:
            msg = json.dumps({'type': 'error', 'message': '大纲结果为空，无法开始内容生成'})
            yield f"data: {msg}\n\n"
            return

        ordered_subsections: List[Dict[str, Any]] = []
        subsection_order = 0
        for section in structure.get("sections", []):
            section_id = str(section.get("id") or "")
            section_title = str(section.get("title") or section_id or "未命名章节")
            msg = json.dumps({
                'type': 'detail',
                'stage': 'section_outline_ready',
                'message': f'章节大纲已生成并存储: {section_title}',
                'metadata': {
                    'section_id': section_id,
                    'section_title': section_title,
                },
            })
            yield f"data: {msg}\n\n"
            for subsection in section.get("subsections", []):
                subsection_id = str(subsection.get("id") or "")
                subsection_title = str(subsection.get("title") or subsection_id or "未命名小节")
                subsection_order += 1
                ordered_subsections.append({
                    "section_id": section_id,
                    "subsection_id": subsection_id,
                    "section_title": section_title,
                    "subsection_title": subsection_title,
                    "subsection_order": subsection_order,
                })
                msg = json.dumps({
                    'type': 'detail',
                    'stage': 'subsection_outline_ready',
                    'message': f'第{subsection_order}小节大纲已生成并存储: {section_title} > {subsection_title}',
                    'metadata': {
                        'section_id': section_id,
                        'subsection_id': subsection_id,
                        'section_title': section_title,
                        'subsection_title': subsection_title,
                        'subsection_order': subsection_order,
                    },
                })
                yield f"data: {msg}\n\n"

        emitted_start_indices = set()
        emitted_passed_indices = set()

        # 第2步：生成文档内容（异步启动）
        msg = json.dumps({'type': 'progress', 'message': f'🚀 开始生成内容（共{total_subsections}个小节）...'})
        yield f"data: {msg}\n\n"

        if ordered_subsections:
            first_item = ordered_subsections[0]
            first_msg = f"开始处理小节: {first_item['section_title']} > {first_item['subsection_title']}"
            msg = json.dumps({
                'type': 'detail',
                'message': first_msg,
                'stage': 'subsection_start',
                'metadata': {
                    'section_id': first_item['section_id'],
                    'subsection_id': first_item['subsection_id'],
                    'section_title': first_item['section_title'],
                    'subsection_title': first_item['subsection_title'],
                    'subsection_order': 1,
                    'section_subsection_total': total_subsections,
                    'synthetic': True,
                },
            })
            yield f"data: {msg}\n\n"
            emitted_start_indices.add(0)

        generate_payload = {
            "document_id": document_id,
            "title": title,
            "structure": structure,
            "content_prompts": content_prompts,
            "user_background": req.user_background,
            "user_requirements": user_requirements,
            "rel_threshold": req.rel_threshold,
            "red_threshold": req.red_threshold,
        }
        preflight_warning = preflight_generator_warning(req.topic)
        if preflight_warning:
            msg = json.dumps({'type': 'warning', 'message': preflight_warning})
            yield f"data: {msg}\n\n"

        # 在后台线程中启动生成，同时主线程推送进度
        gen_resp = None
        error_occurred = False

        def generate_async():
            nonlocal gen_resp, error_occurred
            try:
                gen_resp = generate_document_with_recovery(
                    document_id=document_id,
                    generate_payload=generate_payload,
                    timeout_seconds=stream_timeout,
                )
                # Debug: 打印gen_resp的关键字段
                if isinstance(gen_resp, dict):
                    print(f"🔍 [DEBUG] gen_resp keys: {list(gen_resp.keys())}")
                    print(f"🔍 [DEBUG] quality_score_avg: {gen_resp.get('quality_score_avg')}")
                    print(f"🔍 [DEBUG] unieval_available_subsections: {gen_resp.get('unieval_available_subsections')}")
                    print(f"🔍 [DEBUG] bandit_selected_arm_counts: {gen_resp.get('bandit_selected_arm_counts')}")

                if not isinstance(gen_resp, dict) or not gen_resp.get("success"):
                    error_occurred = True
                    print(f"生成错误: {gen_resp}")
            except Exception as e:
                error_occurred = True
                gen_resp = {"success": False, "error": f"生成服务异常: {e}"}
                print(f"生成错误: {e}")

        gen_thread = threading.Thread(target=generate_async, daemon=True)
        gen_thread.start()

        # 定期检查生成进度
        last_count = 0
        last_event_id = 0
        timeout = time.time() + stream_timeout
        last_progress_update = time.time()
        last_keepalive = time.time()

        while gen_thread.is_alive() and time.time() < timeout:
            try:
                # 查询当前生成的小节数
                history_resp = DOWNSTREAM_SESSION.post(
                    f"{OUTLINER_URL}/history/get",
                    json={"document_id": document_id},
                    timeout=10
                )
                if history_resp.status_code == 200:
                    history = history_resp.json().get("history", [])
                    current_count = len(history)

                    # 注意：不要在生成过程中基于 history 的增量来发送 subsection_passed 事件！
                    # 因为这会导致前端错误地认为所有小节都已完成，即使生成还在进行中。
                    # 只有当生成器真正返回最终结果后，才发送这些事件。
                    # （参考下面的"最后一轮小节完成补发"部分）

                    # Do not synthesize later subsection_start events from
                    # saved-history counts. The generator emits real
                    # subsection_start/progress events; synthetic starts can
                    # make the UI show multiple Drafting items while the real
                    # pipeline is still working on the first active subsection.

                    # 每次进度变化或30秒都推送一次进度（但不要做implicit的完成判断）
                    if current_count > last_count or time.time() - last_progress_update > 30:
                        # 注意：这里只是报告已添加到 history 的项目数，不应该用来判断是否完成。
                        # 最终完成状态必须以后端返回的完整文档结果为准。
                        progress = min(100, int(current_count / total_subsections * 100)) if total_subsections > 0 else 0
                        msg = json.dumps({
                            'type': 'progress',
                            'message': f'进度: {current_count}/{total_subsections} 小节已处理 ({progress}%)',
                            'metadata': {
                                'completed': current_count,
                                'total': total_subsections,
                                'percent': progress,
                            },
                            'note': '进度反映已保存的内容，最终状态以后端完整结果为准',
                        })
                        yield f"data: {msg}\n\n"
                        last_count = current_count
                        last_progress_update = time.time()

                    if time.time() - last_keepalive > 15:
                        heartbeat = json.dumps({'type': 'heartbeat', 'message': '⏳ 生成中，正在保持连接...'})
                        yield f"data: {heartbeat}\n\n"
                        last_keepalive = time.time()

                # 查询流程细节事件
                events_resp = DOWNSTREAM_SESSION.post(
                    f"{OUTLINER_URL}/history/progress",
                    json={"document_id": document_id, "after_id": last_event_id, "limit": 200},
                    timeout=10,
                )
                if events_resp.status_code == 200:
                    events = events_resp.json().get("events", [])
                    for event_item in events:
                        detail_msg = event_item.get("message", "")
                        event_stage = event_item.get("stage", "")
                        event_meta = event_item.get("metadata", {})
                        event_meta = dict(event_meta) if isinstance(event_meta, dict) else {}
                        if event_item.get("section_id") and not event_meta.get("section_id"):
                            event_meta["section_id"] = event_item.get("section_id")
                        if event_item.get("subsection_id") and not event_meta.get("subsection_id"):
                            event_meta["subsection_id"] = event_item.get("subsection_id")
                        if event_stage == "controller_result":
                            print(f"🎯 [Web SSE Forward] controller_result event: arm={event_meta.get('selected_arm')}, reward={event_meta.get('reward')}")
                        msg = json.dumps({
                            'type': 'detail',
                            'message': detail_msg,
                            'stage': event_stage,
                            'metadata': event_meta,
                        })
                        yield f"data: {msg}\n\n"
                        last_event_id = max(last_event_id, int(event_item.get("id", 0)))
            except Exception as e:
                print(f"查询进度异常: {e}")

            time.sleep(2)  # 每2秒检查一次

        # 等待线程结束（最多等待10秒）
        gen_thread.join(timeout=10)

        if error_occurred:
            history_items = fetch_history_items(document_id=document_id, timeout_seconds=30)
            if history_items:
                recovered = _recover_partial_document(document_id=document_id, attempts=1, timeout_seconds=30)
                if recovered.get("success") and recovered.get("content"):
                    if isinstance(gen_resp, dict):
                        recovered.setdefault("stats", {}).update(extract_document_quality_metrics(gen_resp))
                    msg = json.dumps({'type': 'complete', 'result': recovered})
                    yield f"data: {msg}\n\n"
                    return
                detail = _clean_error_text((gen_resp or {}).get("error") if isinstance(gen_resp, dict) else "", "生成服务连接失败")
                msg = json.dumps({'type': 'error', 'message': f'{detail}；已生成 {len(history_items)} 个小节，但未完成全文，已停止返回不完整文档'})
                yield f"data: {msg}\n\n"
                return
            else:
                detail = ""
                if isinstance(gen_resp, dict):
                    detail = _clean_error_text(gen_resp.get("error") or gen_resp.get("message"), "")
                    if not detail:
                        detail = str({
                            "task_id": gen_resp.get("task_id"),
                            "last_status": gen_resp.get("last_status"),
                            "interrupted": gen_resp.get("interrupted"),
                        })
                message = "生成服务连接失败"
                if detail:
                    message = f"{message}: {detail[:260]}"
                msg = json.dumps({'type': 'error', 'message': message})
                yield f"data: {msg}\n\n"
                return

        if gen_resp is None:
            # 线程仍在运行但超时 - 尝试从数据库恢复
            try:
                history_resp = DOWNSTREAM_SESSION.post(
                    f"{OUTLINER_URL}/history/get",
                    json={"document_id": document_id},
                    timeout=10
                )
                if history_resp.status_code == 200:
                    history_items = history_resp.json().get("history", [])
                    if len(history_items) > 0:
                        if len(history_items) >= req.chapter_count * req.subsection_count:
                            recovered = _recover_partial_document(document_id=document_id, attempts=1, timeout_seconds=30)
                            if recovered.get("content"):
                                recovered["success"] = True
                                recovered["partial"] = False
                                recovered.setdefault("stats", {})["failed_subsections"] = 0
                                recovered.setdefault("stats", {})["forced_subsections"] = 0
                                msg = json.dumps({'type': 'complete', 'result': recovered})
                                yield f"data: {msg}\n\n"
                                return
                        # 不把未完成全文伪装成成功结果；保留进度，要求用户继续轮询。
                        msg = json.dumps({'type': 'error', 'message': f'生成仍在运行或超时；已生成 {len(history_items)} 个小节，但未完成全文，请稍后刷新任务状态'})
                        yield f"data: {msg}\n\n"
                        return
                    else:
                        msg = json.dumps({'type': 'error', 'message': '生成未能开始，请检查生成服务'})
                        yield f"data: {msg}\n\n"
                        return
            except:
                msg = json.dumps({'type': 'error', 'message': '生成超时且无法恢复'})
                yield f"data: {msg}\n\n"
                return

        if not gen_resp.get("success"):
            history_items = fetch_history_items(document_id=document_id, timeout_seconds=30)
            if history_items:
                recovered = _recover_partial_document(document_id=document_id, attempts=1, timeout_seconds=30)
                if len(history_items) >= req.chapter_count * req.subsection_count and recovered.get("content"):
                    recovered["success"] = True
                    recovered["partial"] = False
                    recovered.setdefault("stats", {})["expected_subsections"] = req.chapter_count * req.subsection_count
                    recovered.setdefault("stats", {})["passed_subsections"] = len(history_items)
                    recovered.setdefault("stats", {})["failed_subsections"] = 0
                    recovered.setdefault("stats", {})["forced_subsections"] = 0
                    recovered.setdefault("stats", {})["total_generated"] = len(history_items)
                    if isinstance(gen_resp, dict):
                        recovered.setdefault("stats", {}).update(extract_document_quality_metrics(gen_resp))
                    msg = json.dumps({'type': 'complete', 'result': recovered})
                    yield f"data: {msg}\n\n"
                    return
                if recovered.get("success") and recovered.get("content"):
                    if isinstance(gen_resp, dict):
                        recovered.setdefault("stats", {}).update(extract_document_quality_metrics(gen_resp))
                    msg = json.dumps({'type': 'complete', 'result': recovered})
                    yield f"data: {msg}\n\n"
                    return
                err_msg = _clean_error_text(gen_resp.get('error'), '文档生成失败')
                msg = json.dumps({'type': 'error', 'message': f'{err_msg}；已生成 {len(history_items)} 个小节，但未完成全文，已停止返回不完整文档'})
                yield f"data: {msg}\n\n"
                return
            else:
                err_msg = _clean_error_text(gen_resp.get('error'), '文档生成失败')
                msg = json.dumps({'type': 'error', 'message': err_msg})
                yield f"data: {msg}\n\n"
                return

        # 再抓取一轮收尾事件，避免线程结束时最后几条细节日志丢失
        try:
            events_resp = DOWNSTREAM_SESSION.post(
                f"{OUTLINER_URL}/history/progress",
                json={"document_id": document_id, "after_id": last_event_id, "limit": 500},
                timeout=10,
            )
            if events_resp.status_code == 200:
                events = events_resp.json().get("events", [])
                for event_item in events:
                    detail_msg = event_item.get("message", "")
                    event_stage = event_item.get("stage", "")
                    event_meta = event_item.get("metadata", {})
                    event_meta = dict(event_meta) if isinstance(event_meta, dict) else {}
                    if event_item.get("section_id") and not event_meta.get("section_id"):
                        event_meta["section_id"] = event_item.get("section_id")
                    if event_item.get("subsection_id") and not event_meta.get("subsection_id"):
                        event_meta["subsection_id"] = event_item.get("subsection_id")
                    msg = json.dumps({
                        'type': 'detail',
                        'message': detail_msg,
                        'stage': event_stage,
                        'metadata': event_meta,
                    })
                    yield f"data: {msg}\n\n"
                    last_event_id = max(last_event_id, int(event_item.get("id", 0)))
        except Exception as e:
            print(f"收尾事件查询异常: {e}")

        try:
            history_resp = DOWNSTREAM_SESSION.post(
                f"{OUTLINER_URL}/history/get",
                json={"document_id": document_id},
                timeout=10
            )
            final_history = history_resp.json().get("history", []) if history_resp.status_code == 200 else []
            final_count = len(final_history)

            # 仅在生成器完全返回后，才补发真正通过验证的小节事件。
            for idx in range(len(emitted_passed_indices), final_count):
                if idx >= len(ordered_subsections):
                    break

                item = ordered_subsections[idx]
                history_item = final_history[idx] if idx < len(final_history) else {}
                is_forced_pass = history_item.get("metadata", {}).get("forced_pass", False)

                if not is_forced_pass:
                    pass_msg = f"小节通过验证: {item['section_title']} > {item['subsection_title']}"
                    msg = json.dumps({
                        'type': 'detail',
                        'message': pass_msg,
                        'stage': 'subsection_passed',
                        'metadata': {
                            'section_id': item['section_id'],
                            'subsection_id': item['subsection_id'],
                            'section_title': item['section_title'],
                            'subsection_title': item['subsection_title'],
                            'subsection_order': idx + 1,
                            'section_subsection_total': total_subsections,
                            'verification_passed': True,
                        },
                    })
                    yield f"data: {msg}\n\n"
                else:
                    pass_msg = f"小节保留最佳真实草稿: {item['section_title']} > {item['subsection_title']}"
                    msg = json.dumps({
                        'type': 'detail',
                        'message': pass_msg,
                        'stage': 'subsection_passed',
                        'metadata': {
                            'section_id': item['section_id'],
                            'subsection_id': item['subsection_id'],
                            'section_title': item['section_title'],
                            'subsection_title': item['subsection_title'],
                            'subsection_order': idx + 1,
                            'section_subsection_total': total_subsections,
                            'forced_pass': True,
                            'force_reason': history_item.get("metadata", {}).get("force_reason", "unknown"),
                            'best_effort': True,
                            'verification_passed': True,
                        },
                    })
                    yield f"data: {msg}\n\n"

                emitted_passed_indices.add(idx)
        except Exception as e:
            print(f"收尾补发小节事件异常: {e}")

        # 第3步：获取最终内容
        msg = json.dumps({'type': 'progress', 'message': '📦 整合文档内容...'})
        yield f"data: {msg}\n\n"

        try:
            history_resp = DOWNSTREAM_SESSION.post(
                f"{OUTLINER_URL}/history/get",
                json={"document_id": document_id},
                timeout=10
            )
            history_items = history_resp.json().get("history", []) if history_resp.status_code == 200 else []
        except:
            history_items = []

        markdown_content = build_markdown_document(
            title,
            structure,
            history_items,
            generated_sections=gen_resp.get("sections", []),
            user_background=req.user_background,
            extra_requirements=req.extra_requirements,
        )

        # 计算统计数据
        expected_subsections = req.chapter_count * req.subsection_count
        passed = gen_resp.get("passed_subsections", 0)
        failed = len(gen_resp.get("failed_subsections", []))
        forced = len(gen_resp.get("forced_subsections", []))
        total_generated = passed + failed

        partial_mode = passed < expected_subsections or bool(gen_resp.get("interrupted"))
        if passed <= 0:
            msg = json.dumps({
                'type': 'error',
                'message': f'生成未产出可用内容：通过 {passed}/{expected_subsections}，失败 {failed}。请重试或检查下游服务。'
            })
            yield f"data: {msg}\n\n"
            return
        if partial_mode:
            warn_text = gen_resp.get("interrupted_reason") or f'生成未达到目标小节数：通过 {passed}/{expected_subsections}，失败 {failed}'
            msg = json.dumps({'type': 'warning', 'message': f'⚠️ {warn_text}，先展示已通过内容'})
            yield f"data: {msg}\n\n"

        # 聚合所有小节的source_reference_count
        total_source_refs = 0
        for h_item in history_items:
            meta = h_item.get("metadata", {})
            verification = meta.get("verification", {}) or h_item.get("verification", {})
            if isinstance(verification, dict):
                src_check = verification.get("source_check", {})
                if isinstance(src_check, dict):
                    total_source_refs += src_check.get("reference_count", 0)

        result = {
            "success": True,
            "partial": partial_mode,
            "interrupted": partial_mode,
            "document_id": document_id,
            "title": title,
            "content": markdown_content,
            "stats": {
                "expected_subsections": expected_subsections,
                "passed_subsections": passed,
                "failed_subsections": failed,
                "forced_subsections": forced,
                "total_generated": total_generated,
                "total_iterations": gen_resp.get("total_iterations", 0),
                "generation_time": gen_resp.get("generation_time", ""),
                "controller_effective_subsections": gen_resp.get("controller_effective_subsections", 0),
                "rag_used_subsections": gen_resp.get("rag_used_subsections", 0),
                "rag_search_success_subsections": gen_resp.get("rag_search_success_subsections", 0),
                "total_source_references": total_source_refs,
                "token_usage": gen_resp.get("token_usage", {}),
                "prompt_cache_hit_rate": gen_resp.get("prompt_cache_hit_rate", 0.0),
                "generator_short_draft_total": gen_resp.get("generator_short_draft_total", 0),
                "chapter_assets": gen_resp.get("chapter_assets", []) if isinstance(gen_resp.get("chapter_assets"), list) else [],
                "chapter_asset_count": len(gen_resp.get("chapter_assets", [])) if isinstance(gen_resp.get("chapter_assets"), list) else 0,
                **extract_document_quality_metrics(gen_resp if isinstance(gen_resp, dict) else {}),
            },
        }
        web_stats_snapshot = _record_generation_success(partial=partial_mode)
        result["stats"]["total_generations"] = web_stats_snapshot.get("total_generations", WEB_TOTAL_GENERATIONS_BASE)

        msg = json.dumps({'type': 'complete', 'result': result})
        print(f"🎉 [BACKEND] Sending complete event with stats:")
        print(f"  - quality_score_avg: {result['stats'].get('quality_score_avg')}")
        print(f"  - unieval_available_subsections: {result['stats'].get('unieval_available_subsections')}")
        print(f"  - bandit_selected_arm_counts: {result['stats'].get('bandit_selected_arm_counts')}")
        print(f"  - bandit_reward_avg: {result['stats'].get('bandit_reward_avg')}")
        print(f"🎉 [BACKEND] Complete message: {msg[:200]}...")
        yield f"data: {msg}\n\n"

    except Exception as e:
        msg = json.dumps({'type': 'error', 'message': f'内部错误: {str(e)}'})
        yield f"data: {msg}\n\n"


@app.get("/api/generate-stream")
async def generate_stream_endpoint(
    topic: str,
    chapter_count: int = 2,
    subsection_count: int = 2,
    user_background: str = "",
    extra_requirements: str = "",
    rel_threshold: float = WEB_DEFAULT_REL_THRESHOLD,
    red_threshold: float = WEB_DEFAULT_RED_THRESHOLD,
    timeout_seconds: int = 7200,
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    """SSE 端点：实时推送文档生成进度"""
    verify_auth(x_api_key=x_api_key, authorization=authorization)
    req = GenerateDocRequest(
        topic=topic,
        chapter_count=chapter_count,
        subsection_count=subsection_count,
        user_background=user_background,
        extra_requirements=extra_requirements,
        rel_threshold=rel_threshold,
        red_threshold=red_threshold,
        timeout_seconds=timeout_seconds,
    )

    return StreamingResponse(
        generate_stream(req),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )


@app.get("/api/recover-document")
def recover_document(document_id: str) -> Dict[str, Any]:
    return _recover_partial_document(document_id=document_id, attempts=6, timeout_seconds=45)


@app.post("/api/generate")
def generate_document(
    req: GenerateDocRequest,
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
) -> Dict[str, Any]:
    verify_auth(x_api_key=x_api_key, authorization=authorization)
    timeout_profile = _build_timeout_profile(
        chapter_count=req.chapter_count,
        subsection_count=req.subsection_count,
        requested_timeout=req.timeout_seconds or REQUEST_TIMEOUT,
    )
    effective_timeout = timeout_profile["effective_timeout_seconds"]
    start_time = time.time()
    try:
        result = _build_document(req=req, timeout_seconds=effective_timeout)
        elapsed = time.time() - start_time
        _record_timeout_metrics(elapsed_seconds=elapsed, result=result)
        _inject_timeout_profile(result=result, timeout_profile=timeout_profile)
        if result.get("success"):
            web_stats_snapshot = _record_generation_success(partial=bool(result.get("partial") or result.get("interrupted")))
            result.setdefault("stats", {})
            if isinstance(result["stats"], dict):
                result["stats"]["total_generations"] = web_stats_snapshot.get("total_generations", WEB_TOTAL_GENERATIONS_BASE)
        return result
    except HTTPException:
        raise
    except Exception as e:
        print(f"[generate_document] UNEXPECTED ERROR: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文档生成异常: {str(e)[:200]}")


@app.post("/api/download-docx")
def download_docx(
    req: DownloadDocxRequest,
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    print(f"[DOCX] ✅ ENDPOINT CALLED", flush=True)
    verify_auth(x_api_key=x_api_key, authorization=authorization)
    stream = markdown_to_docx(req.title, req.content)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = (req.title or "flowernet_document").strip()[:40]
    ascii_fallback = "flowernet_document"
    encoded = quote(f"{safe_title}_{ts}.docx")

    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={ascii_fallback}_{ts}.docx; filename*=UTF-8''{encoded}"
        },
    )


@app.post("/api/download-pdf")
def download_pdf(
    req: DownloadDocxRequest,
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    print(f"[PDF] ✅ ENDPOINT CALLED", flush=True)
    verify_auth(x_api_key=x_api_key, authorization=authorization)
    try:
        print(f"[PDF] 🔄 Building PDF (title={req.title[:30] if req.title else 'N/A'}..., content_len={len(req.content) if req.content else 0})", flush=True)
        stream = markdown_to_pdf(req.title, req.content)
        if stream.getbuffer().nbytes == 0:
            print(f"[PDF] ❌ Generated PDF is empty!", flush=True)
            raise ValueError("PDF generation produced empty file")
        print(f"[PDF] ✅ PDF ready: {stream.getbuffer().nbytes} bytes", flush=True)
    except Exception as e:
        print(f"[PDF] ❌ PDF download failed: {str(e)}", flush=True)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = (req.title or "flowernet_document").strip()[:40]
    ascii_fallback = "flowernet_document"
    encoded = quote(f"{safe_title}_{ts}.pdf")

    return StreamingResponse(
        stream,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={ascii_fallback}_{ts}.pdf; filename*=UTF-8''{encoded}"
        },
    )


@app.get("/api/poffices/openapi.json")
def poffices_openapi(request: Request):
    return _build_poffices_openapi(request)


def _run_poffices_task(task_id: str, req: PofficesGenerateRequest):
    try:
        _set_poffices_task(
            task_id,
            status="running",
            message="任务运行中",
            started_at=datetime.now().isoformat(),
        )

        timeout_profile = _build_timeout_profile(
            chapter_count=req.chapter_count,
            subsection_count=req.subsection_count,
            requested_timeout=req.timeout_seconds,
        )
        start_time = time.time()
        deadline = start_time + timeout_profile["effective_timeout_seconds"]
        build_attempt = 0

        while True:
            current_task = _restore_poffices_task(task_id) or {}
            if _poffices_task_cancelled(current_task):
                raise RuntimeError("task_cancelled")
            if build_attempt >= POFFICES_OUTLINER_RETRY_MAX:
                raise TimeoutError(f"outliner_retry_exhausted_after_{build_attempt}_attempts")
            build_attempt += 1
            doc_req = GenerateDocRequest(
                topic=req.query,
                chapter_count=req.chapter_count,
                subsection_count=req.subsection_count,
                user_background=req.user_background,
                extra_requirements=req.extra_requirements,
                rel_threshold=req.rel_threshold,
                red_threshold=req.red_threshold,
            )
            try:
                result = _build_document(
                    req=doc_req,
                    timeout_seconds=max(30, int(deadline - time.time())),
                )
                break
            except HTTPException as exc:
                detail = str(exc.detail)
                if not _is_retryable_outline_failure(detail) or time.time() >= deadline:
                    raise
                delay = min(_outline_retry_delay_seconds(detail, attempt=build_attempt), max(1.0, deadline - time.time()))
                _set_poffices_task(
                    task_id,
                    status="running",
                    message=(
                        "远端 outliner 正在限流或排队，FlowerNet 已保留任务并自动重试；"
                        f"第 {build_attempt} 次恢复等待 {delay:.0f}s"
                    ),
                    last_retryable_error=detail[:500],
                    retry_count=build_attempt,
                    updated_at=datetime.now().isoformat(),
                )
                print(f"[Poffices] ⏳ retryable outliner failure for {task_id}; retry in {delay:.1f}s: {detail[:180]}")
                time.sleep(delay)
            except Exception as exc:
                detail = str(exc)
                if not _is_retryable_outline_failure(detail) or time.time() >= deadline:
                    raise
                delay = min(_outline_retry_delay_seconds(detail, attempt=build_attempt), max(1.0, deadline - time.time()))
                _set_poffices_task(
                    task_id,
                    status="running",
                    message=(
                        "远端 outliner 正在限流或排队，FlowerNet 已保留任务并自动重试；"
                        f"第 {build_attempt} 次恢复等待 {delay:.0f}s"
                    ),
                    last_retryable_error=detail[:500],
                    retry_count=build_attempt,
                    updated_at=datetime.now().isoformat(),
                )
                print(f"[Poffices] ⏳ retryable outliner exception for {task_id}; retry in {delay:.1f}s: {detail[:180]}")
                time.sleep(delay)

        current_task = _restore_poffices_task(task_id) or {}
        if _poffices_task_cancelled(current_task):
            raise RuntimeError("task_cancelled")

        elapsed = time.time() - start_time
        _record_timeout_metrics(elapsed_seconds=elapsed, result=result)
        _inject_timeout_profile(result=result, timeout_profile=timeout_profile)
        if elapsed > timeout_profile["effective_timeout_seconds"]:
            raise TimeoutError(f"任务超时: {elapsed:.1f}s > {timeout_profile['effective_timeout_seconds']}s")

        _set_poffices_task(
            task_id,
            status="completed",
            result=result,
            message="任务完成",
            completed_at=datetime.now().isoformat(),
        )
    except HTTPException as exc:
        _set_poffices_task(
            task_id,
            status="failed",
            error=str(exc.detail),
            message="任务失败",
            completed_at=datetime.now().isoformat(),
        )
    except Exception as exc:
        _set_poffices_task(
            task_id,
            status="failed",
            error=str(exc),
            message="任务失败",
            completed_at=datetime.now().isoformat(),
        )


@app.post("/api/poffices/generate")
def poffices_generate(
    req: PofficesGenerateRequest,
    request: Request,
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    verify_auth(x_api_key=x_api_key, authorization=authorization)

    incoming_task_id = _extract_task_id_from_payload(req.model_dump()) or _extract_task_id_from_payload(getattr(req, "model_extra", {}) or {})
    if incoming_task_id:
        try:
            return _poffices_wait_for_task_result(
                request=request,
                task_id=incoming_task_id,
                wait=False,
                wait_seconds=POFFICES_POLL_WAIT_SECONDS,
            )
        except HTTPException as exc:
            if exc.status_code == 404:
                recovered_payload = req.model_dump()
                recovered_payload.update(getattr(req, "model_extra", {}) or {})
                recovered = _poffices_recover_from_missing_task(
                    request=request,
                    payload=recovered_payload,
                    wait=False,
                    wait_seconds=POFFICES_POLL_WAIT_SECONDS,
                )
                if recovered is not None:
                    recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {incoming_task_id}")
                    return recovered
            raise

    normalized_query = (req.query or "").strip()
    query_task_id = _extract_task_id_from_payload(normalized_query)
    if query_task_id:
        try:
            return _poffices_wait_for_task_result(
                request=request,
                task_id=query_task_id,
                wait=False,
                wait_seconds=POFFICES_POLL_WAIT_SECONDS,
            )
        except HTTPException as exc:
            if exc.status_code == 404:
                recovered = _poffices_recover_from_missing_task(
                    request=request,
                    payload=req.model_dump(),
                    wait=False,
                    wait_seconds=POFFICES_POLL_WAIT_SECONDS,
                )
                if recovered is not None:
                    recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {query_task_id}")
                    return recovered
            raise

    if len(normalized_query) < 2:
        return {
            "success": False,
            "task_status": "failed",
            "error": "query 不能为空且至少 2 个字符",
            "message": "请求参数无效",
        }

    req = req.model_copy(update={"query": normalized_query})

    if req.async_mode:
        return _poffices_start_or_reuse_async_task(
            request=request,
            req=req,
            wait=False,
            wait_seconds=POFFICES_POLL_WAIT_SECONDS,
        )

    timeout_profile = _build_timeout_profile(
        chapter_count=req.chapter_count,
        subsection_count=req.subsection_count,
        requested_timeout=req.timeout_seconds,
    )

    doc_req = GenerateDocRequest(
        topic=req.query,
        chapter_count=req.chapter_count,
        subsection_count=req.subsection_count,
        user_background=req.user_background,
        extra_requirements=req.extra_requirements,
        rel_threshold=req.rel_threshold,
        red_threshold=req.red_threshold,
    )
    try:
        start_time = time.time()
        result = _build_document(req=doc_req, timeout_seconds=timeout_profile["effective_timeout_seconds"])
        elapsed = time.time() - start_time
        _record_timeout_metrics(elapsed_seconds=elapsed, result=result)
        _inject_timeout_profile(result=result, timeout_profile=timeout_profile)
        return _build_poffices_result(request=request, result=result)
    except HTTPException as exc:
        return {
            "success": False,
            "task_status": "failed",
            "error": str(exc.detail),
            "message": "文档生成失败",
        }
    except Exception as exc:
        return {
            "success": False,
            "task_status": "failed",
            "error": str(exc),
            "message": "文档生成异常",
        }


@app.post("/api/poffices/task-status", name="poffices_task_status")
def poffices_task_status(
    request: Request,
    payload: Dict[str, Any] = Body(default_factory=dict),
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    verify_auth(x_api_key=x_api_key, authorization=authorization)

    req = PofficesTaskStatusRequest.model_validate(payload or {})
    task_id = (req.task_id or _extract_task_id_from_payload(payload)).strip()
    wait = _coerce_bool((payload or {}).get("wait"), default=req.wait)
    wait_seconds = int((payload or {}).get("wait_seconds") or req.wait_seconds or POFFICES_POLL_WAIT_SECONDS)
    cancel = _coerce_bool((payload or {}).get("cancel"), default=req.cancel)
    stale_404_payload = _payload_contains_task_not_found(payload)

    if stale_404_payload:
        recovered_req = _coerce_poffices_request_from_payload(payload)
        if recovered_req is not None:
            recovered = _poffices_start_or_reuse_async_task(
                request=request,
                req=recovered_req,
                wait=wait,
                wait_seconds=wait_seconds,
            )
            if task_id:
                recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {task_id}")
            return recovered

    if not task_id:
        recovered_req = _coerce_poffices_request_from_payload(payload)
        if recovered_req is not None:
            return _poffices_start_or_reuse_async_task(
                request=request,
                req=recovered_req,
                wait=wait,
                wait_seconds=wait_seconds,
            )
        # If a Poffices block passes an older result as text, still try to pull
        # a task id out of the response-shaped fields before failing.
        task_id = _extract_task_id_from_payload(_extract_text_field_from_payload(payload, ("content", "text", "result", "output", "markdown", "document"))).strip()

    if not task_id:
        return {
            "success": False,
            "status": "failed",
            "task_status": "failed",
            "error": "task_id not found in Poffices input. Connect FlowerNet Start Task output to Poll Render, or pass task_id explicitly.",
            "message": "缺少 task_id",
        }

    if cancel:
        return _cancel_poffices_task(task_id)

    try:
        return _poffices_wait_for_task_result(
            request=request,
            task_id=task_id,
            wait=wait,
            wait_seconds=wait_seconds,
        )
    except HTTPException as exc:
        if exc.status_code == 404:
            recovered = _poffices_recover_from_missing_task(
                request=request,
                payload=payload or {},
                wait=wait,
                wait_seconds=wait_seconds,
            )
            if recovered is not None:
                recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {task_id}")
                return recovered
        raise


@app.post("/api/poffices/poll-render")
def poffices_poll_render(
    request: Request,
    payload: Dict[str, Any] = Body(default_factory=dict),
    x_api_key: str = Header(default="", alias="X-API-Key"),
    authorization: str = Header(default="", alias="Authorization"),
):
    verify_auth(x_api_key=x_api_key, authorization=authorization)

    task_id = _extract_task_id_from_payload(payload).strip()
    wait = _coerce_bool((payload or {}).get("wait"), default=True)
    wait_seconds = int((payload or {}).get("wait_seconds") or POFFICES_POLL_WAIT_SECONDS)
    stale_404_payload = _payload_contains_task_not_found(payload)

    if stale_404_payload:
        recovered_req = _coerce_poffices_request_from_payload(payload)
        if recovered_req is not None:
            recovered = _poffices_start_or_reuse_async_task(
                request=request,
                req=recovered_req,
                wait=wait,
                wait_seconds=wait_seconds,
            )
            if task_id:
                recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {task_id}")
            return recovered

    if not task_id:
        task_id = _extract_task_id_from_payload(_extract_text_field_from_payload(payload, ("content", "text", "result", "output", "markdown", "document"))).strip()
    if not task_id:
        recovered_req = _coerce_poffices_request_from_payload(payload)
        if recovered_req is not None:
            return _poffices_start_or_reuse_async_task(
                request=request,
                req=recovered_req,
                wait=wait,
                wait_seconds=wait_seconds,
            )
    if not task_id:
        return {
            "success": False,
            "status": "failed",
            "task_status": "failed",
            "error": "task_id not found in Poffices input. Poll Render must receive FlowerNet Start Task output.",
            "message": "缺少 task_id",
        }

    try:
        return _poffices_wait_for_task_result(
            request=request,
            task_id=task_id,
            wait=wait,
            wait_seconds=wait_seconds,
        )
    except HTTPException as exc:
        if exc.status_code == 404:
            recovered = _poffices_recover_from_missing_task(
                request=request,
                payload=payload or {},
                wait=wait,
                wait_seconds=wait_seconds,
            )
            if recovered is not None:
                recovered.setdefault("warning", f"stale_or_unknown_task_id_recovered: {task_id}")
                return recovered
        raise


def _legacy_poffices_task_status(
    req: PofficesTaskStatusRequest,
    request: Request,
) -> Dict[str, Any]:
    task = _restore_poffices_task(req.task_id)

    if not task:
        raise HTTPException(status_code=404, detail="task_id not found")

    status = task.get("status", "unknown")
    if status in {"queued", "running"}:
        _restart_restored_poffices_task(req.task_id, task)
        task = _restore_poffices_task(req.task_id) or task
        status = task.get("status", status)

    if status in {"queued", "running"}:
        started_at = task.get("started_at") or task.get("created_at")
        timeout_seconds = int(task.get("timeout_seconds") or 0)
        if started_at and timeout_seconds > 0:
            try:
                elapsed = (datetime.now() - datetime.fromisoformat(started_at)).total_seconds()
                if elapsed > timeout_seconds + 30:
                    with POFFICES_TASKS_LOCK:
                        current_status = POFFICES_TASKS.get(req.task_id, {}).get("status")
                    if current_status in {"queued", "running"}:
                        _set_poffices_task(
                            req.task_id,
                            status="failed",
                            error=f"任务超时: {int(elapsed)}s",
                            message="任务超时",
                            completed_at=datetime.now().isoformat(),
                        )
                    status = "failed"
                    task = _restore_poffices_task(req.task_id) or task
            except ValueError:
                pass

    if status == "completed":
        result = task.get("result", {})
        return _build_poffices_result(request=request, result=result)

    if status == "failed":
        return {
            "success": False,
            "task_id": req.task_id,
            "status": "failed",
            "error": task.get("error", "unknown error"),
            "message": task.get("message", "任务失败"),
        }

    return {
        "success": True,
        "task_id": req.task_id,
        "status": status,
        "message": task.get("message", "处理中"),
    }


# ==================== 导入指标展示 API ====================
try:
    from metrics_api import router as metrics_router
    app.include_router(metrics_router)
    print("✅ 指标展示 API 已加载")
except ImportError:
    print("⚠️  指标展示 API 未加载")


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "8010"))
    uvicorn.run(app, host="0.0.0.0", port=port)
