#!/usr/bin/env python3
"""Build a visual DOCX report for the Week 2 benchmark."""

from __future__ import annotations

import json
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "results" / "week2" / "week2_benchmark_outputs.json"
SUMMARY = ROOT / "results" / "week2" / "week2_summary.json"
OUTDIR = ROOT / "reports" / "week2"
ASSET_DIR = OUTDIR / "assets"
REPORT = OUTDIR / "FlowerNet_Week2_15Topic_Benchmark_Report.docx"


def load_rows() -> List[Dict[str, Any]]:
    if not RESULTS.exists():
        raise FileNotFoundError(f"Missing benchmark outputs: {RESULTS}")
    data = json.loads(RESULTS.read_text(encoding="utf-8"))
    return data if isinstance(data, list) else []


def group_summary(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    by_system: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for row in rows:
        by_system[str(row.get("system"))].append(row)
    out = []
    for system, items in by_system.items():
        n = len(items)
        out.append(
            {
                "system": system,
                "n": n,
                "ok": sum(1 for x in items if x.get("status") == "ok"),
                "avg_score": round(sum(float(x.get("week2_score", 0) or 0) for x in items) / max(1, n), 4),
                "avg_chars": round(sum(float((x.get("metrics") or {}).get("chars", 0) or 0) for x in items) / max(1, n), 1),
                "avg_citations": round(sum(float((x.get("metrics") or {}).get("citation_marker_count", 0) or 0) for x in items) / max(1, n), 1),
                "avg_repeat": round(sum(float((x.get("metrics") or {}).get("repeat_3gram_ratio", 0) or 0) for x in items) / max(1, n), 4),
                "avg_calls": round(sum(float(x.get("llm_calls", 0) or 0) for x in items) / max(1, n), 2),
                "controller_calls": sum(int(x.get("controller_calls", 0) or 0) for x in items),
                "forced_pass": sum(int(x.get("forced_pass_subsections", 0) or 0) for x in items),
                "time": round(sum(float(x.get("elapsed_seconds", 0) or 0) for x in items), 1),
            }
        )
    return sorted(out, key=lambda x: x["avg_score"], reverse=True)


def font(size: int = 18, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_bar_chart(summary: List[Dict[str, Any]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1500, 820
    img = Image.new("RGB", (w, h), "#f8fafc")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(22)
    small_f = font(18)
    d.text((55, 40), "Week 2 Main Results: Average Quality Score", fill="#0f172a", font=title_f)
    d.text((55, 95), "Scores are computed from real outputs: length, structure, citations, verifier pass, tables, and repetition penalty.", fill="#475569", font=label_f)
    chart_x, chart_y = 90, 170
    chart_w, chart_h = 1320, 500
    max_score = max([float(x["avg_score"]) for x in summary] + [1.0])
    colors = ["#0f766e", "#2563eb", "#7c3aed", "#ea580c", "#dc2626", "#0891b2", "#65a30d", "#9333ea", "#475569"]
    bar_gap = 18
    bar_w = max(70, int((chart_w - bar_gap * (len(summary) - 1)) / max(1, len(summary))))
    for i in range(6):
        y = chart_y + chart_h - i * chart_h / 5
        d.line((chart_x, y, chart_x + chart_w, y), fill="#dbe4ee", width=2)
        d.text((30, y - 12), f"{i * max_score / 5:.2f}", fill="#64748b", font=small_f)
    for idx, row in enumerate(summary):
        score = float(row["avg_score"])
        x = chart_x + idx * (bar_w + bar_gap)
        bh = int(chart_h * score / max_score)
        y = chart_y + chart_h - bh
        d.rounded_rectangle((x, y, x + bar_w, chart_y + chart_h), radius=12, fill=colors[idx % len(colors)])
        d.text((x + 4, y - 30), f"{score:.3f}", fill="#0f172a", font=small_f)
        label = row["system"].replace("flowernet_", "fn_").replace("_adapter", "_ad")
        d.text((x, chart_y + chart_h + 18), label[:18], fill="#334155", font=small_f)
    img.save(path)


def draw_ablation_chart(summary: List[Dict[str, Any]], path: Path) -> None:
    ablations = [x for x in summary if x["system"].startswith("flowernet")]
    if not ablations:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1300, 620
    img = Image.new("RGB", (w, h), "#fff7ed")
    d = ImageDraw.Draw(img)
    title_f = font(38, True)
    label_f = font(22)
    small_f = font(18)
    d.text((50, 35), "FlowerNet Ablation Study", fill="#172554", font=title_f)
    d.text((50, 85), "Full system should dominate ablations when verifier, NLI fusion, and bandit control are useful.", fill="#475569", font=label_f)
    base_x, base_y = 70, 160
    max_score = max(float(x["avg_score"]) for x in ablations + [{"avg_score": 1.0}])
    for idx, row in enumerate(sorted(ablations, key=lambda r: r["system"])):
        y = base_y + idx * 90
        score = float(row["avg_score"])
        d.text((base_x, y + 10), row["system"], fill="#0f172a", font=label_f)
        d.rounded_rectangle((base_x + 360, y, base_x + 1100, y + 46), radius=20, fill="#fed7aa")
        d.rounded_rectangle((base_x + 360, y, base_x + 360 + int(740 * score / max_score), y + 46), radius=20, fill="#f97316")
        d.text((base_x + 1120, y + 8), f"{score:.3f}", fill="#0f172a", font=small_f)
    img.save(path)


def draw_bandit_chart(rows: List[Dict[str, Any]], path: Path) -> None:
    counts: Counter[str] = Counter()
    for row in rows:
        if row.get("system") == "flowernet_full":
            for sub in row.get("subsections") or []:
                v = sub.get("verification") if isinstance(sub, dict) else {}
                if isinstance(v, dict):
                    dims = v.get("semantic_dimensions") or v.get("quality_dimensions") or {}
                    if isinstance(dims, dict) and dims:
                        worst = min(dims.items(), key=lambda kv: float((kv[1] or {}).get("score", kv[1]) if not isinstance(kv[1], (int, float)) else kv[1]))[0]
                        counts[str(worst)] += 1
    if not counts:
        counts.update({"controller_calls": sum(int(r.get("controller_calls", 0) or 0) for r in rows if str(r.get("system", "")).startswith("flowernet"))})
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1000, 520
    img = Image.new("RGB", (w, h), "#0f172a")
    d = ImageDraw.Draw(img)
    title_f = font(34, True)
    label_f = font(20)
    d.text((42, 35), "Controller / Verifier Diagnostic Distribution", fill="#e2e8f0", font=title_f)
    total = max(1, sum(counts.values()))
    y = 120
    palette = ["#22c55e", "#f97316", "#3b82f6", "#a855f7", "#ef4444", "#06b6d4"]
    for idx, (name, count) in enumerate(counts.most_common()):
        width = int(760 * count / total)
        d.text((55, y + 8), name[:24], fill="#cbd5e1", font=label_f)
        d.rounded_rectangle((330, y, 330 + 760, y + 42), radius=18, fill="#1e293b")
        d.rounded_rectangle((330, y, 330 + width, y + 42), radius=18, fill=palette[idx % len(palette)])
        d.text((1110, y + 8), str(count), fill="#e2e8f0", font=label_f)
        y += 62
    img.save(path)


def add_table(doc: Document, rows: List[Dict[str, Any]]) -> None:
    headers = ["System", "N", "OK", "Avg score", "Chars", "Cites", "Repeat", "LLM calls", "Controller", "Forced", "Time(s)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        vals = [
            row["system"], row["n"], row["ok"], row["avg_score"], row["avg_chars"],
            row["avg_citations"], row["avg_repeat"], row["avg_calls"],
            row["controller_calls"], row["forced_pass"], row["time"],
        ]
        for i, val in enumerate(vals):
            cells[i].text = str(val)


def build() -> Path:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    rows = load_rows()
    summary = group_summary(rows)
    if not summary:
        raise RuntimeError("No rows to report.")

    score_png = ASSET_DIR / "week2_scores.png"
    ablation_png = ASSET_DIR / "week2_ablation.png"
    bandit_png = ASSET_DIR / "week2_bandit.png"
    draw_bar_chart(summary, score_png)
    draw_ablation_chart(summary, ablation_png)
    draw_bandit_chart(rows, bandit_png)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("FlowerNet Week 2 Benchmark Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("15-topic long-document benchmark with baselines, ARISE/LongWriter adapters, and three FlowerNet ablations.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", level=1)
    best = summary[0]
    flowernet = next((x for x in summary if x["system"] == "flowernet_full"), None)
    doc.add_paragraph(
        f"The current benchmark contains {len(rows)} executed rows across {len(set(r.get('topic_id') for r in rows))} topics. "
        f"The leading system is {best['system']} with average score {best['avg_score']}. "
        + (f"FlowerNet full score is {flowernet['avg_score']}, with {flowernet['controller_calls']} controller calls and {flowernet['forced_pass']} best-real passes. " if flowernet else "")
        + "LongWriter GGUF and ARISE are included as executable local/adapted baselines where the original full runtime is unavailable on this Mac."
    )

    doc.add_picture(str(score_png), width=Inches(6.6))
    doc.add_heading("Main Results", level=1)
    add_table(doc, summary)

    doc.add_heading("Ablation Study", level=1)
    if ablation_png.exists():
        doc.add_picture(str(ablation_png), width=Inches(6.4))
    doc.add_paragraph(
        "Ablations isolate the effects of contextual bandit control, NLI/UniEval fusion, and multidimensional quality gates. "
        "Forced best-real passes are counted separately from strict verifier passes, so the report does not hide quality failures."
    )

    doc.add_heading("Controller and Quality Diagnostics", level=1)
    doc.add_picture(str(bandit_png), width=Inches(6.2))
    doc.add_paragraph(
        "Controller calls and verifier failures are extracted from the generated row metadata. "
        "If a subsection reaches the retry cap, the runner can accept the best real draft, but it is marked as forced_best_real_pass rather than strict pass."
    )

    doc.add_heading("Experimental Setup", level=1)
    doc.add_paragraph(
        "Topic set: experiments/topics_week1.json, consisting of 10 FreshWiki-2024-style public knowledge prompts and 5 education-domain prompts, seed=42. "
        "All executable systems use the same topic prompts. Default generation scale is 2 chapters x 2 subsections for comparability."
    )
    doc.add_paragraph(
        "Systems: Vanilla LLM, Self-Refine, CogWriter adapter, LongWriter GGUF via Ollama, ARISE-compatible rubric-guided survey adapter, FlowerNet full, w/o Bandit, w/o NLI, and w/o multidim. "
        "Metrics combine output completeness, structure, citation markers, table evidence, verifier pass rate, forced pass rate, and repeated 3-gram penalty."
    )

    doc.add_heading("Important Reproducibility Notes", level=1)
    notes = [
        "ARISE original GitHub clone failed through the current network path, so this run uses an explicitly labeled ARISE-compatible adapter with rubric critique and revision.",
        "LongWriter uses the Mac-feasible bartowski LongWriter-llama3.1-8b GGUF Q4_K_M model through Ollama. Invalid repeated or blank outputs are detected and scored as failed_invalid_output.",
        "No numbers are fabricated: failed systems remain failed, and partial/forced passes are separated from strict verifier passes.",
    ]
    for note in notes:
        doc.add_paragraph(note, style=None)

    doc.add_heading("Selected Raw Evidence", level=1)
    for row in rows[:8]:
        doc.add_paragraph(
            f"{row.get('system')} / {row.get('topic_id')}: status={row.get('status')}, score={row.get('week2_score')}, "
            f"chars={(row.get('metrics') or {}).get('chars', 0)}, error={str(row.get('error', ''))[:120]}"
        )

    doc.save(REPORT)
    SUMMARY.write_text(json.dumps({"summary": summary, "report": str(REPORT), "assets": [str(score_png), str(ablation_png), str(bandit_png)]}, ensure_ascii=False, indent=2), encoding="utf-8")
    return REPORT


if __name__ == "__main__":
    print(build())
