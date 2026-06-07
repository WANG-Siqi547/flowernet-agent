#!/usr/bin/env python3
"""Build a DOCX action report for Week 2 feedback.

The report is intentionally conservative: adapter results, official-code
availability, and hardware/API blockers are labeled separately so the document
can be shared with collaborators without overclaiming.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUTDIR = ROOT / "reports" / "week2"
ASSET_DIR = OUTDIR / "assets"
REPORT = OUTDIR / "FlowerNet_Week2_Feedback_Action_Report.docx"

MAIN_SUMMARY = ROOT / "results" / "week2" / "week2_summary_full.json"
MAIN_OUTPUTS = ROOT / "results" / "week2" / "week2_benchmark_outputs_full.json"
NO_VC_SUMMARY = ROOT / "results" / "week2" / "no_vc_15topic_summary.json"
NO_VC_SMOKE_SUMMARY = ROOT / "results" / "week2" / "no_vc_smoke_summary.json"
JOURNAL_METRICS = ROOT / "results" / "week2" / "journal_metrics_rouge_bertscore_15topic.json"

EXTERNAL_REPOS = {
    "LongWriter": {
        "repo": "https://github.com/THUDM/LongWriter",
        "local": ROOT / "external_baselines" / "LongWriter",
        "mac_status": "Official model path requires 8B/9B HF weights or vLLM/CUDA for practical 10k+ word generation. Mac can run only quantized/adapter attempts; the previous Ollama/GGUF attempt did not find the model locally.",
        "server_action": "Run official THUDM LongWriter with transformers/vLLM on GPU, record model checkpoint, commit hash, max_new_tokens, and judge settings.",
    },
    "ARISE": {
        "repo": "https://github.com/ziwang11112/ARISE",
        "local": ROOT / "external_baselines" / "ARISE",
        "mac_status": "Source is available, but the full pipeline needs Serper/OpenAI/Gemini/Anthropic keys, paper retrieval, and LaTeX/BibTeX. It is runnable on Mac only after API and toolchain configuration; current Week 2 data used an adapter, not the full original pipeline.",
        "server_action": "Configure .env keys, paper count, topic, and LaTeX; run ARISE_Source_Code/run_all.py per topic and evaluate generated PDFs/Markdown.",
    },
    "CogWriter": {
        "repo": "https://github.com/KaiyangWan/CogWriter",
        "local": ROOT / "external_baselines" / "CogWriter",
        "mac_status": "Source is available. Official README recommends Llama-3.3-70B via vLLM; closed-source model branch has a placeholder API key. True official reproduction needs a vLLM endpoint or a clean API-key patch.",
        "server_action": "Run CogWriter main.py with the original model/backend configuration or patch llms/llms.py to read environment keys, then run the 15-topic dataset.",
    },
}


def load_json(path: Path) -> Any:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def font(size: int = 20, bold: bool = False):
    for path in [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def summary_rows(summary: Any) -> List[Dict[str, Any]]:
    if isinstance(summary, dict):
        rows = summary.get("summary")
        if isinstance(rows, list):
            return rows
    return []


def merge_summaries() -> List[Dict[str, Any]]:
    main = summary_rows(load_json(MAIN_SUMMARY))
    no_vc = summary_rows(load_json(NO_VC_SUMMARY))
    if not no_vc:
        no_vc = summary_rows(load_json(NO_VC_SMOKE_SUMMARY))
        for row in no_vc:
            row["system"] = f"{row.get('system')} (smoke only)"
            row["note"] = "Only one-topic smoke is complete at report build time."
    by_name: Dict[str, Dict[str, Any]] = {}
    for row in main + no_vc:
        by_name[str(row.get("system"))] = row
    return sorted(by_name.values(), key=lambda r: float(r.get("avg_score", 0) or 0), reverse=True)


def draw_score_chart(rows: List[Dict[str, Any]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    rows = rows[:12]
    w, h = 1500, 820
    img = Image.new("RGB", (w, h), "#f8fafc")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    body = font(22)
    small = font(18)
    d.text((55, 35), "Week 2 Feedback Check: Quality vs. Call-Budget Baselines", fill="#0f172a", font=title)
    d.text((55, 92), "No-VC baselines isolate whether verifier/controller add value beyond extra calls; external metrics flag what must be improved and rerun.", fill="#475569", font=body)
    chart_x, chart_y, chart_w, chart_h = 90, 170, 1320, 500
    max_score = max([float(r.get("avg_score", 0) or 0) for r in rows] + [1.0])
    for i in range(6):
        y = chart_y + chart_h - i * chart_h / 5
        d.line((chart_x, y, chart_x + chart_w, y), fill="#dbe4ee", width=2)
        d.text((28, y - 12), f"{i * max_score / 5:.2f}", fill="#64748b", font=small)
    gap = 16
    bw = max(72, int((chart_w - gap * (len(rows) - 1)) / max(1, len(rows))))
    palette = {
        "flowernet_full": "#0f766e",
        "flowernet_no_vc_budget20": "#f97316",
        "flowernet_no_vc_direct": "#fb923c",
    }
    fallback = ["#2563eb", "#7c3aed", "#dc2626", "#0891b2", "#65a30d", "#9333ea", "#475569"]
    for i, row in enumerate(rows):
        system = str(row.get("system", ""))
        score = float(row.get("avg_score", 0) or 0)
        x = chart_x + i * (bw + gap)
        y = chart_y + chart_h - int(chart_h * score / max_score)
        color = palette.get(system.replace(" (smoke only)", ""), fallback[i % len(fallback)])
        d.rounded_rectangle((x, y, x + bw, chart_y + chart_h), radius=12, fill=color)
        d.text((x, y - 30), f"{score:.3f}", fill="#0f172a", font=small)
        label = system.replace("flowernet_", "fn_").replace("_adapter", "_ad")
        d.text((x - 3, chart_y + chart_h + 18), label[:18], fill="#334155", font=small)
    img.save(path)


def draw_pipeline(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1500, 520
    img = Image.new("RGB", (w, h), "#0f172a")
    d = ImageDraw.Draw(img)
    title = font(38, True)
    body = font(22)
    small = font(18)
    d.text((50, 35), "Confirmed FlowerNet Full Workflow", fill="#e2e8f0", font=title)
    boxes = [
        ("1. Outliner", "Plan chapters and subsections"),
        ("2. Generator", "Draft each subsection from outline"),
        ("3. Verifier", "Check evidence, relevance, redundancy, structure"),
        ("4. Controller", "Choose repair arm if verification fails"),
        ("5. Re-check", "Loop until pass or max 5 attempts"),
    ]
    x, y = 55, 145
    for idx, (head, desc) in enumerate(boxes):
        d.rounded_rectangle((x, y, x + 235, y + 160), radius=20, fill="#1e293b", outline="#38bdf8", width=3)
        d.text((x + 22, y + 30), head, fill="#f8fafc", font=body)
        d.text((x + 22, y + 76), desc, fill="#cbd5e1", font=small)
        if idx < len(boxes) - 1:
            d.line((x + 245, y + 80, x + 305, y + 80), fill="#f97316", width=5)
            d.polygon([(x + 305, y + 80), (x + 286, y + 68), (x + 286, y + 92)], fill="#f97316")
        x += 285
    d.text((65, 365), "New no-VC baseline removes steps 3-5: outline + generate only. Budget20 uses the same 20 downstream generation attempts but no verifier/controller.", fill="#e2e8f0", font=body)
    img.save(path)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(15, 23, 42)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[Any]]) -> None:
    headers = list(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(str(h))
        run.bold = True
    for vals in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(vals):
            cells[idx].text = str(val)


def repo_commit(path: Path) -> str:
    head = path / ".git" / "HEAD"
    if not head.exists():
        return "not cloned"
    text = head.read_text(encoding="utf-8").strip()
    if text.startswith("ref:"):
        ref = path / ".git" / text.split(" ", 1)[1]
        return ref.read_text(encoding="utf-8").strip()[:12] if ref.exists() else "cloned"
    return text[:12]


def build() -> Path:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    rows = merge_summaries()
    score_png = ASSET_DIR / "week2_feedback_scores.png"
    pipe_png = ASSET_DIR / "week2_feedback_pipeline.png"
    draw_score_chart(rows, score_png)
    draw_pipeline(pipe_png)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FlowerNet Week 2 Feedback Action Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(15, 23, 42)
    subtitle = doc.add_paragraph("Call-budget baseline, official-baseline feasibility, and journal-grade evaluation plan")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Direct Answers to the Feedback", 1)
    add_para(doc, "Yes, the current FlowerNet full architecture is: (a) Outliner generates the outline; (b) Generator drafts each subsection from the outline; (c) Verifier checks every subsection; (d) if a check fails, Controller selects a repair strategy and the system loops back to verification, up to 5 attempts. After max attempts, FlowerNet uses the best real draft rather than a written template fallback.")
    add_para(doc, "Implemented: Added two no-verifier/no-controller baselines. flowernet_no_vc_direct runs only Outliner + Generator and directly accepts each first draft. flowernet_no_vc_budget20 keeps the same 20 downstream generation attempts as FlowerNet full, but removes Verifier and Controller; it selects candidates only by a static surface heuristic. This isolates whether FlowerNet's quality comes from the Modifier/Controller loop rather than just more LLM calls.")
    add_para(doc, "Non-negotiable analysis rule: if another system beats FlowerNet full on a real metric, the correct response is to inspect the failure mode, improve FlowerNet full, and rerun. The report must not change the evaluator to favor FlowerNet or invent data.")

    doc.add_picture(str(pipe_png), width=Inches(7.0))

    add_heading(doc, "2. Current Results Including the New Baseline", 1)
    doc.add_picture(str(score_png), width=Inches(7.0))
    add_table(
        doc,
        ["System", "N", "OK", "Avg score", "Avg chars", "Avg citations", "Avg LLM calls", "Controller calls", "Forced/best passes", "Time(s)"],
        [
            [
                r.get("system"),
                r.get("n"),
                r.get("ok"),
                r.get("avg_score"),
                r.get("avg_chars"),
                r.get("avg_citations"),
                r.get("avg_llm_calls"),
                r.get("avg_controller_calls"),
                r.get("forced_pass_total"),
                r.get("elapsed_total_seconds"),
            ]
            for r in rows
        ],
    )
    add_para(doc, "Interpretation: the no-VC baseline is the missing fairness control requested in the feedback. If its full 15-topic result remains below FlowerNet full, the claim becomes stronger: extra calls alone are not enough; verification and controller-guided repair add measurable value. If it beats FlowerNet full on any robust metric, FlowerNet full must be improved and rerun before making a paper claim.")

    add_heading(doc, "3. Official Baseline Reproduction Status", 1)
    add_table(
        doc,
        ["System", "Official repo", "Local commit", "MacBook status", "Next action"],
        [
            [name, info["repo"], repo_commit(info["local"]), info["mac_status"], info["server_action"]]
            for name, info in EXTERNAL_REPOS.items()
        ],
    )
    add_para(doc, "Important distinction: the existing Week 2 report used executable adapters for LongWriter/ARISE/CogWriter where official reproduction was not yet available. The official repositories have now been cloned locally, but authentic reproduction requires the constraints above. Adapter scores should be labeled as adapter baselines, not original-system scores.")

    add_heading(doc, "4. Journal-Grade Three-Layer Evaluation Plan", 1)
    add_para(doc, "Layer 1 - Automatic metrics: add ROUGE and BERTScore against external references. ROUGE follows Lin (2004), and BERTScore follows Zhang et al. (2020). FreshWiki-style topics should use curated multi-page Wikipedia reference packs, not a single arbitrary page; education topics should use published survey/review papers as references.")
    add_para(doc, "Layer 2 - LLM-as-a-Judge: use GPT-4o-style pairwise and rubric judging, with randomized order, swapped-order bias checks, and blind system names. Judge dimensions should include coverage, factual grounding, citation faithfulness, coherence, non-redundancy, novelty, and usefulness.")
    add_para(doc, "Layer 3 - Human evaluation: optional for KBS/ESWA but valuable. Recommended protocol: 3-5 blinded raters, pairwise comparison plus Likert dimensions, report agreement such as Krippendorff's alpha or Fleiss' kappa, and include examples of disagreements.")

    external = load_json(JOURNAL_METRICS)
    if isinstance(external, dict) and external.get("summary_by_system"):
        add_heading(doc, "4.1 Added External ROUGE/BERTScore Results", 2)
        add_para(doc, str(external.get("metric_note", "")))
        add_para(doc, "Reference protocol: FreshWiki-style topics use Wikipedia reference packs; education topics use published review/survey papers retrieved from Semantic Scholar/Crossref. The current reference file is saved under references/week2_reference_sets.json for auditability.")
        add_table(
            doc,
            ["System", "N", "BERTScore F1", "ROUGE-L", "ROUGE-2", "ROUGE-1"],
            [
                [
                    r.get("system"),
                    r.get("n"),
                    f"{float(r.get('bertscore_f1', 0) or 0):.4f}",
                    f"{float(r.get('rougeL', 0) or 0):.4f}",
                    f"{float(r.get('rouge2', 0) or 0):.4f}",
                    f"{float(r.get('rouge1', 0) or 0):.4f}",
                ]
                for r in external.get("summary_by_system", [])
            ],
        )
        systems = {str(r.get("system")): r for r in external.get("summary_by_system", [])}
        full = systems.get("flowernet_full", {})
        wo_bandit = systems.get("flowernet_wo_bandit", {})
        no_vc_budget = systems.get("flowernet_no_vc_budget20", {})
        add_para(
            doc,
            "Current external-metric finding: FlowerNet full is above the no-verifier/no-controller budget20 baseline on ROUGE-2 and ROUGE-L, which supports a real contribution from verification/control beyond call count. However, flowernet_wo_bandit currently leads on BERTScore/ROUGE, so the bandit/controller policy should be inspected for over-repair or reference drift before claiming FlowerNet full is the final best system.",
        )
        add_table(
            doc,
            ["Comparison", "BERTScore F1", "ROUGE-L", "Interpretation"],
            [
                [
                    "Full vs no-VC budget20",
                    f"{float(full.get('bertscore_f1', 0) or 0):.4f} vs {float(no_vc_budget.get('bertscore_f1', 0) or 0):.4f}",
                    f"{float(full.get('rougeL', 0) or 0):.4f} vs {float(no_vc_budget.get('rougeL', 0) or 0):.4f}",
                    "Full improves lexical overlap but BERTScore gap is small; judge/human layers are needed.",
                ],
                [
                    "Full vs wo-bandit",
                    f"{float(full.get('bertscore_f1', 0) or 0):.4f} vs {float(wo_bandit.get('bertscore_f1', 0) or 0):.4f}",
                    f"{float(full.get('rougeL', 0) or 0):.4f} vs {float(wo_bandit.get('rougeL', 0) or 0):.4f}",
                    "wo-bandit is currently stronger on these external overlap metrics; diagnose bandit strategy before main paper claims.",
                ],
            ],
        )

    add_heading(doc, "5. Generation Scale Upgrade", 1)
    add_para(doc, "The current 2x2 setting is useful for engineering iteration but is too short for a long-document paper claim. The next benchmark should use 3x3 or 5x3. A 5x3 setup has 15 subsections; with max 5 attempts this can require up to 75 downstream generation attempts plus outlining, verification, controller, table/reference assembly, and judge calls. That is a compute/cost decision, not a code-only issue.")
    add_para(doc, "Recommended path: first run a 3-topic 3x3 pilot to stabilize timeout, citation density, and controller rate; then run full 15-topic 3x3; only then move to 5x3 if the system remains stable and the external metrics support the claim.")

    add_heading(doc, "6. What Still Needs Compute or Manual Inputs", 1)
    add_para(doc, "LongWriter original: needs GPU server or a carefully documented non-official GGUF route. The MacBook can test wrappers, but it cannot fairly reproduce the official long-output claim at full context and speed.")
    add_para(doc, "External references: ROUGE/BERTScore cannot be final until reference packs are built. FreshWiki references require Wikipedia page selection; education references require published survey/review papers.")
    add_para(doc, "LLM-as-a-Judge and human evaluation: GPT-4o judge requires API budget and a locked rubric; human evaluation requires rater recruitment and blind annotation forms.")

    add_heading(doc, "7. Recommended Next Run Matrix", 1)
    add_table(
        doc,
        ["Run", "Purpose", "Systems", "Scale", "Evaluator"],
        [
            ["Fairness control", "Prove VC contribution beyond call count", "FlowerNet full vs no_vc_budget20", "15 topics, 2x2", "Internal + external metrics"],
            ["Long-form pilot", "Test true long document behavior", "FlowerNet full + strongest baselines", "3 topics, 3x3", "ROUGE/BERTScore + judge"],
            ["Official baselines", "Remove adapter objection", "LongWriter/CogWriter/ARISE original code", "15 topics", "Same evaluator"],
            ["Paper-grade main", "Final claim", "All systems + ablations", "15 topics, 3x3 or 5x3", "Auto + judge + human eval"],
        ],
    )

    doc.add_paragraph("Generated from real local files; incomplete or blocked items are explicitly labeled above.")
    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    print(build())
